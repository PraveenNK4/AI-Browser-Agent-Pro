from __future__ import annotations

import asyncio
import base64
import io
import json
import logging
import os
import pathlib
import platform
import shutil
import time
import time
import uuid
import builtins
from datetime import datetime
from src.utils.utils import slugify
from src.utils.chrome_frame import apply_chrome_frame, build_chrome_frame_html, get_page_metadata
from typing import Any, AsyncGenerator, Dict, Optional

from browser_use.agent.views import AgentHistoryList, AgentOutput
from browser_use.browser.browser import BrowserConfig
from browser_use.browser.context import BrowserContext, BrowserContextConfig
from browser_use.browser.views import BrowserState
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Inches
from gradio.components import Component
from langchain_core.language_models.chat_models import BaseChatModel
from PIL import Image

# NOTE: Security wait bypass monkeypatch removed for compatibility with browser-use v0.1.48+

from src.agent.browser_use.browser_use_agent import BrowserUseAgent
from src.orchestrator.process_orchestrator import run_all_processes
from src.orchestrator.process_parser import parse_processes
from src.browser.custom_browser import CustomBrowser
from src.controller.custom_controller import CustomController
from src.utils import llm_provider
from src.utils.dom_snapshot import set_run_id
from src.utils.strip_and_enrich_history import strip_and_enrich
from src.utils.llm_script_generator import generate_script as generate_llm_script
from src.webui.webui_manager import WebuiManager
import gradio as gr

logger = logging.getLogger(__name__)


def _detect_browser():
    """Detect available browsers and prefer Chrome for automation."""
    system = platform.system()
    if system == "Windows":
        chrome_paths = [
            r"C:\Program Files\Google\Chrome\Application\chrome.exe",
            r"C:\Program Files (x86)\Google\Chrome\Application\chrome.exe",
        ]
        edge_paths = [
            r"C:\Program Files\Microsoft\Edge\Application\msedge.exe",
            r"C:\Program Files (x86)\Microsoft\Edge\Application\msedge.exe"
        ]
    elif system == "Darwin":
        chrome_paths = ["/Applications/Google Chrome.app/Contents/MacOS/Google Chrome"]
        edge_paths = ["/Applications/Microsoft Edge.app/Contents/MacOS/Microsoft Edge"]
    else:
        chrome_paths = ["/usr/bin/google-chrome", "/usr/bin/google-chrome-stable"]
        edge_paths = ["/usr/bin/microsoft-edge", "/usr/bin/microsoft-edge-stable"]

    for path in edge_paths + chrome_paths:
        if os.path.exists(path):
            logger.info(f"Detected browser: {path}")
            return path
    logger.warning("No supported browser found.")
    return None


# Filter to suppress browser-use security warnings when secrets are used without domain locking
class _SensitiveDataSecurityFilter(logging.Filter):
    def filter(self, record):
        msg = record.getMessage()
        # Silence all security-related warnings and "Waiting..." messages
        if any(term in msg for term in ["not locked down", "sensitive_data", "malicious website", "Waiting 10 seconds", "insecure settings"]):
            return False
        return True

# Apply the filter to multiple loggers to ensure total silence
logging.getLogger("browser_use.agent.service").addFilter(_SensitiveDataSecurityFilter())
logging.getLogger("browser_use.agent").addFilter(_SensitiveDataSecurityFilter())
logging.getLogger("browser_use").addFilter(_SensitiveDataSecurityFilter())


async def _initialize_llm(
    provider: Optional[str],
    model_name: Optional[str],
    temperature: float,
    base_url: Optional[str],
    api_key: Optional[str],
    num_ctx: Optional[int] = None,
) -> Optional[BaseChatModel]:
    """Initialize LLM based on settings."""
    if not provider or not model_name:
        logger.info("LLM Provider/Model not specified.")
        return None
    try:
        logger.info(f"Initializing LLM: {provider}/{model_name} (temp={temperature})")
        llm = llm_provider.get_llm_model(
            provider=provider,
            model_name=model_name,
            temperature=temperature,
            base_url=base_url or None,
            api_key=api_key or None,
            num_ctx=num_ctx if provider == "ollama" else None,
        )
        return llm
    except Exception as e:
        logger.error(f"Failed to initialize LLM: {e}", exc_info=True)
        gr.Warning(f"Failed to initialize LLM '{model_name}': {e}")
        return None


def _get_config_value(
    webui_manager: WebuiManager,
    comp_dict: Dict[Component, Any],
    comp_id_suffix: str,
    default: Any = None,
) -> Any:
    """Get value from component dictionary using ID suffix."""
    for prefix in ["browser_use_agent", "agent_settings", "browser_settings"]:
        try:
            comp_id = f"{prefix}.{comp_id_suffix}"
            comp = webui_manager.get_component_by_id(comp_id)
            return comp_dict.get(comp, default)
        except KeyError:
            continue
    logger.warning(f"Component '{comp_id_suffix}' not found.")
    return default


def _format_agent_output(model_output: AgentOutput) -> str:
    """Format AgentOutput for chatbot display."""
    content = ""
    if model_output:
        try:
            # LLM reasoning
            llm_response = ""
            if hasattr(model_output, 'response'):
                response = getattr(model_output, 'response', None)
                if response:
                    llm_response = str(response)
            elif hasattr(model_output, 'text'):
                text = getattr(model_output, 'text', None)
                if text:
                    llm_response = str(text)

            if llm_response:
                content += f"**LLM Reasoning:**\n{llm_response}\n\n"

            # Actions JSON
            action_dump = [action.model_dump(exclude_none=True) for action in model_output.action]
            state_dump = model_output.current_state.model_dump(exclude_none=True)
            output_dump = {"current_state": state_dump, "action": action_dump}
            json_str = json.dumps(output_dump, indent=4, ensure_ascii=False)
            content += f"**Actions:**\n<pre><code class='language-json'>{json_str}</code></pre>"
        except Exception as e:
            logger.error(f"Error formatting output: {e}", exc_info=True)
            content = f"<pre><code>Error formatting output: {e}\nRaw: {str(model_output)}</code></pre>"
    return content.strip()


async def _handle_new_step(
    webui_manager: WebuiManager, state: BrowserState, output: AgentOutput, step_num: int
):
    """Callback for each agent step."""
    if not hasattr(webui_manager, "bu_chat_history"):
        webui_manager.bu_chat_history = []
    
    step_num -= 1
    step_end_time = time.perf_counter()
    step_duration = 0.0
    if webui_manager.bu_step_start_time:
        step_duration = round(step_end_time - webui_manager.bu_step_start_time, 3)
    
    logger.info(f"Step {step_num} completed.")

    # Progression tracking
    if webui_manager.check_step_progression(step_num):
        logger.error(f"Step {step_num}: Excessive backward progression.")
    
    # Failure tracking
    step_failed = False
    if hasattr(output, 'error'):
        error = getattr(output, 'error', None)
        if error:
            step_failed = True
            logger.warning(f"Step {step_num} failed: {error}")
    
    if webui_manager.check_step_failure(step_num, step_failed):
        logger.error(f"Step {step_num} exceeded 3 failures. Stopping.")
        webui_manager.bu_should_stop_agent = True
        if webui_manager.bu_agent:
            webui_manager.bu_agent.state.stopped = True
    
    # Hallucination tracking
    current_action = None
    if hasattr(output, 'action') and output.action:
        current_action = str(output.action[0]) if isinstance(output.action, (list, tuple)) else str(output.action)
    
    if webui_manager.check_hallucination(current_action):
        logger.error(f"HALLUCINATION: Action repeated {webui_manager.bu_repeated_action_count} times. Stopping.")
        webui_manager.bu_should_stop_agent = True
        webui_manager.bu_hallucination_triggered = True
        if webui_manager.bu_agent:
            webui_manager.bu_agent.state.stopped = True

    # Token usage
    current_tokens = 0
    if webui_manager.bu_agent:
        current_tokens = webui_manager.bu_agent.state.history.total_input_tokens()
    if not hasattr(webui_manager, 'bu_previous_tokens'):
        webui_manager.bu_previous_tokens = 0
    step_tokens = current_tokens - webui_manager.bu_previous_tokens
    webui_manager.bu_previous_tokens = current_tokens

    # Screenshot — apply Chrome frame when headless
    screenshot_html = ""
    screenshot_data = getattr(state, "screenshot", None)
    if screenshot_data and isinstance(screenshot_data, str) and len(screenshot_data) > 100:
        # Apply Chrome frame if headless
        _headless = os.getenv("HEADLESS", "false").lower() == "true"
        if _headless and webui_manager.bu_browser_context:
            try:
                _page = await webui_manager.bu_browser_context.get_current_page()
                _url = _page.url or ""
                # Skip framing if page is blank or internal
                if _url and not _url.startswith("about:") and not _url.startswith("chrome:"):
                    _title = await _page.title() or _url
                    _vp = _page.viewport_size or {}
                    _pw_ctx = webui_manager.bu_browser_context.session.context
                    screenshot_data = await apply_chrome_frame(
                        browser=_pw_ctx,
                        screenshot_b64=screenshot_data,
                        url=_url,
                        title=_title,
                        is_secure=_url.startswith("https://"),
                        page_width=_vp.get("width", 1280),
                        page_height=_vp.get("height", 900),
                    )
            except Exception as _e:
                logger.debug(f"Chrome frame failed for step screenshot: {_e}")
        img_tag = f'<img src="data:image/png;base64,{screenshot_data}" style="max-width:800px;max-height:600px;object-fit:contain;" />'
        screenshot_html = img_tag + "<br/>"

    # Format output
    formatted_output = _format_agent_output(output)

    # Warnings
    failure_warning = ""
    if step_failed:
        failure_warning += "âš ï¸ **STEP FAILED** - Retrying...<br/>"
    if webui_manager.bu_should_stop_agent:
        if webui_manager.bu_hallucination_triggered:
            failure_warning += "🔴 **HALLUCINATION DETECTED** - Stopped.<br/>"
        else:
            failure_warning += "🔴 **MAX FAILURES** - Stopped.<br/>"

    # Combine message
    step_header = f"--- **Step {step_num}** (Tokens: {step_tokens}, Time: {step_duration:.3f}s) ---"
    final_content = step_header + "<br/>" + screenshot_html + failure_warning + formatted_output

    webui_manager.bu_chat_history.append({"role": "assistant", "content": final_content.strip()})
    webui_manager.bu_step_start_time = time.perf_counter()
    await asyncio.sleep(0.05)


def generate_docx_report(webui_manager: WebuiManager, history: AgentHistoryList, task: str):
    """Generate comprehensive DOCX report with screenshots and findings."""
    try:
        logger.info("Generating DOCX report...")
        if not webui_manager or not hasattr(webui_manager, 'bu_agent_task_id'):
            logger.error("Missing webui_manager or task_id")
            return None
        
        doc = Document()
        doc.add_heading('AI Browser Agent Execution Report', 0)

        # Task summary
        doc.add_heading('Task Summary', level=1)
        task_para = doc.add_paragraph()
        task_para.add_run('Task: ').bold = True
        task_para.add_run(task)

        duration = history.total_duration_seconds()
        tokens = history.total_input_tokens()
        final_result = history.final_result()

        summary_table = doc.add_table(rows=5, cols=2)
        summary_table.style = 'Table Grid'
        summary_table.rows[0].cells[0].text = 'Duration'
        summary_table.rows[0].cells[1].text = f"{duration:.2f}s"
        summary_table.rows[1].cells[0].text = 'Tokens'
        summary_table.rows[1].cells[1].text = str(tokens)
        summary_table.rows[2].cells[0].text = 'Steps'
        summary_table.rows[2].cells[1].text = str(len(history.history))
        summary_table.rows[3].cells[0].text = 'Status'
        summary_table.rows[3].cells[1].text = 'Completed' if final_result else 'Failed'
        summary_table.rows[4].cells[0].text = 'Task ID'
        summary_table.rows[4].cells[1].text = webui_manager.bu_agent_task_id or 'Unknown'

        # Final result
        doc.add_heading('Final Result', level=1)
        if final_result:
            result_para = doc.add_paragraph()
            result_para.add_run('Outcome: ').bold = True
            result_para.add_run(final_result)
        
        step_failures = getattr(webui_manager, 'bu_step_failures', {})
        total_failures = sum(step_failures.values())
        status_para = doc.add_paragraph()
        status_para.add_run('Status: ').bold = True
        if total_failures > 0:
            status_para.add_run(f"Completed with {total_failures} failures")
        else:
            status_para.add_run("All steps successful")

        # Step details
        doc.add_heading('Step-by-Step Execution', level=1)
        for i, step in enumerate(history.history, 1):
            step_status = "PASS" if step_failures.get(i, 0) == 0 else f"FAIL ({step_failures.get(i, 0)}x)"
            doc.add_heading(f'Step {i} - {step_status}', level=2)

            # Goal
            goal_para = doc.add_paragraph()
            goal_para.add_run('Goal: ').bold = True
            goal_text = "Execute action"
            if hasattr(step, 'model_output') and step.model_output:
                if hasattr(step.model_output, 'current_state') and step.model_output.current_state:
                    if hasattr(step.model_output.current_state, 'next_goal') and step.model_output.current_state.next_goal:
                        goal_text = str(step.model_output.current_state.next_goal)
            goal_para.add_run(goal_text)

            # Action
            action_para = doc.add_paragraph()
            action_para.add_run('Action: ').bold = True
            action_desc = "No action"
            if hasattr(step, 'action'):
                action = getattr(step, 'action', None)
                if action:
                    action_desc = str(action[0]) if isinstance(action, (list, tuple)) else str(action)
            action_para.add_run(action_desc)

            # Screenshot
            if hasattr(step, 'state') and hasattr(step.state, 'screenshot') and step.state.screenshot:
                try:
                    screenshot_data = step.state.screenshot
                    if screenshot_data.startswith('data:image'):
                        screenshot_data = screenshot_data.split(',')[1]
                    image_data = base64.b64decode(screenshot_data)
                    image = Image.open(io.BytesIO(image_data))
                    
                    if image.width > 900:
                        ratio = 900 / image.width
                        image = image.resize((int(image.width * ratio), int(image.height * ratio)), Image.Resampling.LANCZOS)
                    
                    if image.mode in ("RGBA", "LA", "P"):
                        image = image.convert("RGB")
                    
                    img_buffer = io.BytesIO()
                    image.save(img_buffer, format='JPEG', quality=95)
                    img_buffer.seek(0)
                    
                    doc.add_picture(img_buffer, width=Inches(6))
                    caption = doc.add_paragraph(f'Step {i} screenshot')
                    caption.alignment = WD_ALIGN_PARAGRAPH.CENTER
                except Exception as e:
                    logger.warning(f"Failed to add screenshot for step {i}: {e}")

            # Result
            result_para = doc.add_paragraph()
            result_para.add_run('Result: ').bold = True
            result_text = "No result"
            if hasattr(step, 'result') and step.result:
                # result is list[ActionResult], take first if available
                if isinstance(step.result, list) and len(step.result) > 0:
                    first_result = step.result[0]
                    if hasattr(first_result, 'extracted_content'):
                        extracted = getattr(first_result, 'extracted_content', None)
                        if extracted:
                            result_text = str(extracted)
                        else:
                            result_text = str(first_result)
                    else:
                        result_text = str(first_result)
                else:
                    result_text = str(step.result)
            result_para.add_run(result_text)

            doc.add_paragraph()

        # Errors
        errors = history.errors()
        if errors and any(errors):
            doc.add_heading('Errors', level=1)
            for error in errors:
                if error:
                    doc.add_paragraph(f'Error: {str(error)}')

        # Save
        task_id = webui_manager.bu_agent_task_id
        if not task_id:
            logger.error("Task ID is None")
            return None
        agent_history_path = os.getenv("AGENT_HISTORY_PATH", "./tmp/agent_history")
        task_dir = os.path.join(agent_history_path, task_id)
        os.makedirs(task_dir, exist_ok=True)
        
        # Sanitize task_id for filename (it may contain slashes if it's a nested process)
        safe_filename = task_id.replace("/", "_").replace("\\", "_")
        docx_path = os.path.join(task_dir, f"{safe_filename}_report.docx")
        doc.save(docx_path)
        logger.info(f"DOCX report saved: {docx_path}")
        return docx_path
    except Exception as e:
        logger.error(f"DOCX generation failed: {e}", exc_info=True)
        return None


def _handle_done(webui_manager: WebuiManager, history: AgentHistoryList):
    """Callback when agent finishes - generate reports."""
    logger.info("=" * 80)
    logger.info("REPORT GENERATION STARTED")
    logger.info("=" * 80)
    
    try:
        logger.info(f"Task finished. Duration: {history.total_duration_seconds():.2f}s, Tokens: {history.total_input_tokens()}")

        # Generate DOCX
        task = getattr(webui_manager.bu_agent, 'task', 'Unknown') if webui_manager.bu_agent else 'Unknown'
        docx_path = generate_docx_report(webui_manager, history, task)
        if docx_path:
            webui_manager.bu_docx_path = docx_path
            logger.info(f"DOCX: {docx_path}")

        final_summary = "**Task Completed**\n"
        final_summary += f"- Duration: {history.total_duration_seconds():.2f}s\n"
        final_summary += f"- Tokens: {history.total_input_tokens()}\n"
        
        final_result = history.final_result()
        if final_result:
            final_summary += f"- Result: {final_result}\n"
        
        errors = history.errors()
        if errors and any(errors):
            final_summary += f"- Errors: {len([e for e in errors if e])}\n"
        
        if docx_path:
            final_summary += f"- 📄 Report: {os.path.basename(docx_path)}\n"

        if not any("**Task Completed**" in str(msg.get("content", "")) for msg in webui_manager.bu_chat_history if msg.get("role") == "assistant"):
            webui_manager.bu_chat_history.append({"role": "assistant", "content": final_summary})
        
        logger.info("Report generation complete")
    except Exception as e:
        logger.error(f"Error in done callback: {e}", exc_info=True)


async def _ask_assistant_callback(
    webui_manager: WebuiManager, query: str, browser_context: BrowserContext
) -> Dict[str, Any]:
    """Callback for agent assistance requests."""
    logger.info("Agent needs assistance.")
    
    webui_manager.bu_chat_history.append({
        "role": "assistant",
        "content": f"**Need Help:** {query}\nPlease respond below and click 'Submit Response'."
    })

    webui_manager.bu_response_event = asyncio.Event()
    webui_manager.bu_user_help_response = None

    try:
        from src.utils.config import WEBUI_RESPONSE_TIMEOUT_S
        await asyncio.wait_for(webui_manager.bu_response_event.wait(), timeout=WEBUI_RESPONSE_TIMEOUT_S)
    except asyncio.TimeoutError:
        logger.warning("Timeout waiting for user.")
        webui_manager.bu_chat_history.append({"role": "assistant", "content": "**Timeout:** No response."})
        webui_manager.bu_response_event = None
        return {"response": "Timeout: No response."}

    response = webui_manager.bu_user_help_response
    webui_manager.bu_chat_history.append({"role": "user", "content": response})
    webui_manager.bu_response_event = None
    return {"response": response}


async def run_agent_task(
    webui_manager: WebuiManager, components: Dict[Component, Any]
) -> AsyncGenerator[Dict[Component, Any], None]:
    """Main agent execution loop."""
    
    # Get components
    user_input_comp = webui_manager.get_component_by_id("browser_use_agent.user_input")
    run_button_comp = webui_manager.get_component_by_id("browser_use_agent.run_button")
    stop_button_comp = webui_manager.get_component_by_id("browser_use_agent.stop_button")
    pause_resume_button_comp = webui_manager.get_component_by_id("browser_use_agent.pause_resume_button")
    clear_button_comp = webui_manager.get_component_by_id("browser_use_agent.clear_button")
    chatbot_comp = webui_manager.get_component_by_id("browser_use_agent.chatbot")
    history_file_comp = webui_manager.get_component_by_id("browser_use_agent.agent_history_file")
    gif_comp = webui_manager.get_component_by_id("browser_use_agent.recording_gif")
    playwright_script_comp = webui_manager.get_component_by_id("browser_use_agent.playwright_script")
    docx_report_comp = webui_manager.get_component_by_id("browser_use_agent.docx_report")
    browser_view_comp = webui_manager.get_component_by_id("browser_use_agent.browser_view")

    # Browser settings from env
    browser_binary_path = os.getenv("BROWSER_PATH") or None
    browser_user_data_dir = os.getenv("BROWSER_USER_DATA") or None
    use_own_browser = os.getenv("USE_OWN_BROWSER", "true").lower() == "true"
    keep_browser_open = os.getenv("KEEP_BROWSER_OPEN", "false").lower() == "true"
    should_close_browser_on_finish = not keep_browser_open
    logger.info(f"Settings: Use Own Browser={use_own_browser}, Keep Open={keep_browser_open}, Auto-Close={should_close_browser_on_finish}")
    headless = os.getenv("HEADLESS", "false").lower() == "true"
    disable_security = os.getenv("DISABLE_SECURITY", "false").lower() == "true"
    window_w = int(os.getenv("WINDOW_WIDTH", "1280"))
    window_h = int(os.getenv("WINDOW_HEIGHT", "1100"))
    save_agent_history_path = os.getenv("AGENT_HISTORY_PATH", "./tmp/agent_history")
    save_download_path = os.getenv("DOWNLOADS_PATH", "./tmp/downloads")
    save_recording_path = os.getenv("RECORDING_PATH") or None
    save_trace_path = os.getenv("TRACE_PATH") or None

    os.makedirs(save_agent_history_path, exist_ok=True)
    if save_download_path:
        os.makedirs(save_download_path, exist_ok=True)

    # Get task
    task = components.get(user_input_comp, "").strip()
    if not task:
        gr.Warning("Please enter a task.")
        yield {run_button_comp: gr.update(interactive=True)}
        return

    # Generate task ID (Readable Format)
    timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
    short_uuid = str(uuid.uuid4())[:8]
    task_slug = slugify(task)[:50] # Limit length
    webui_manager.bu_agent_task_id = f"{timestamp}_{task_slug}_{short_uuid}"
    set_run_id(webui_manager.bu_agent_task_id)
    webui_manager.reset_failure_tracking()
    
    # Reset chat history and controller for fresh run
    webui_manager.bu_chat_history = []
    webui_manager.bu_controller = None

    # Process context files
    context_files_comp = webui_manager.get_component_by_id("browser_use_agent.context_files")
    uploaded_files = components.get(context_files_comp, [])
    context_info = ""
    available_file_paths = []
    
    if uploaded_files:
        if not isinstance(uploaded_files, list):
            uploaded_files = [uploaded_files]
            
        context_info = "\n\nContext files:"
        for file_obj in uploaded_files:
            # Gradio files can be objects with .name or strings
            fpath = getattr(file_obj, 'name', str(file_obj))
            if fpath and os.path.exists(fpath):
                filename = os.path.basename(fpath)
                context_info += f"\n- {filename}"
                temp_dir = os.path.join(save_agent_history_path, webui_manager.bu_agent_task_id, "context_files")
                os.makedirs(temp_dir, exist_ok=True)
                temp_path = os.path.join(temp_dir, filename)
                try:
                    shutil.copy2(fpath, temp_path)
                    available_file_paths.append(fpath)
                    available_file_paths.append(temp_path)
                    logger.info(f"📁 Added file to available paths: {fpath}")
                except Exception as e:
                    logger.error(f"Failed to save {filename}: {e}")
    task += context_info

    # ── Cert task detection & rewrite ────────────────────────────────────────
    import re as _re
    _CERT_RE = r'\b(ssl|tls|certificate|https cert|cert detail|cert valid|lock icon|connection secure)\b'
    _obj_text = task.lower()
    if _re.search(_CERT_RE, _obj_text, _re.IGNORECASE):
        _url_match = _re.search(r'(https?://[^\s\]\)\'"]+)', task)
        _url_hint  = _url_match.group(1).rstrip('.,;)') if _url_match else None

        _nav_step = f'go_to_url("{_url_hint}")' if _url_hint else "go_to_url to navigate to the target URL"
        task = (
            f"SSL Certificate Check Task\n"
            f"{'='*50}\n"
            f"{'URL: ' + _url_hint if _url_hint else 'Navigate to the URL in the original task.'}\n"
            f"{'='*50}\n\n"
            f"EXACT STEPS — execute in order, nothing else:\n"
            f"Step 1: {_nav_step}\n"
            f"Step 2: check_ssl_certificate()\n"
            f"Step 3: done(success=true, text='Certificate check complete')\n\n"
            f"RULES:\n"
            f"- check_ssl_certificate() is a registered action — USE IT\n"
            f"- Do NOT click anything in the browser\n"
            f"- Do NOT use extract_content\n"
            f"- Do NOT try to access the lock icon\n"
            f"- After check_ssl_certificate() returns, immediately call done()"
        )
        logger.info(f"[task-rewrite] Cert task → replaced with 3-step check_ssl_certificate flow")

    webui_manager.bu_chat_history.append({"role": "user", "content": task})

    yield {
        user_input_comp: gr.Textbox(value="", interactive=False, placeholder="Agent running..."),
        run_button_comp: gr.Button(value="⏳ Running...", interactive=False),
        stop_button_comp: gr.Button(interactive=True),
        pause_resume_button_comp: gr.Button(value="⸏ Pause", interactive=True),
        clear_button_comp: gr.Button(interactive=False),
        chatbot_comp: gr.update(value=webui_manager.bu_chat_history),
        history_file_comp: gr.update(value=None),
        gif_comp: gr.update(value=None),
    }

    # Agent settings (Ollama only, low temp for strict mode)
    llm_provider_name = "ollama"
    llm_model_name = os.getenv("LLM_MODEL_NAME", "qwen2.5:7b")
    llm_temperature = float(os.getenv("LLM_TEMPERATURE", "0.3"))
    use_vision = os.getenv("USE_VISION", "false").lower() == "true"
    ollama_num_ctx = int(os.getenv("OLLAMA_NUM_CTX", "16000"))
    llm_base_url = os.getenv("LLM_BASE_URL")
    max_steps = int(os.getenv("MAX_STEPS", "25"))
    max_actions = int(os.getenv("MAX_ACTIONS", "3"))
    max_input_tokens = int(os.getenv("MAX_INPUT_TOKENS", "128000"))
    
    # Planner model settings (optional)
    planner_provider = os.getenv("PLANNER_LLM_PROVIDER")
    planner_model_name = os.getenv("PLANNER_LLM_MODEL_NAME")
    planner_temperature = float(os.getenv("PLANNER_LLM_TEMPERATURE", "1.0"))
    
    # Extraction model settings (optional)
    extraction_provider = os.getenv("EXTRACTION_LLM_PROVIDER")
    extraction_model_name = os.getenv("EXTRACTION_LLM_MODEL_NAME")
    extraction_temperature = float(os.getenv("EXTRACTION_LLM_TEMPERATURE", "0.0"))

    from src.utils.config import VAULT_CREDENTIAL_PREFIX as _VPX
    # Planner system message - explicit instructions for file uploads
    extend_planner_system_message = f"""
━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━
🎯 ACTION STRATEGY - REDUCE TURNS & INCREASE SPEED
━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━

1. 🔐 LOGIN (CRITICAL SPEED BOOST):
   - NEVER use separate input_text for username and password. 
   - ALWAYS use: smart_login(username, password, button_text)
   - This performs UN/PW/Click in ONE atomic turn.
   - Example: smart_login(username="{{{{{_VPX}_USERNAME}}}}", password="{{{{{_VPX}_PASSWORD}}}}", button_text="Sign in")
   - ⚠️ CRITICAL: Check the page URL. If you are on a login page, use smart_login immediately.

2. 🔗 URL INTEGRITY:
   - Ensure URLs are correct. NEVER use double dots (e.g., http://host..domain).
   - Use the exact URL provided in the task.

2. 🎯 HOVER DECISION TREE - READ THIS FIRST BEFORE ANY HOVER ACTION!
━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━

QUESTION: What are you trying to hover?

❓ The indexed element ITSELF (text, button, link)?
   → Use: hover_element(index=N)
   → Example: Hover "Submit" button at index 10
   → Action: hover_element(index=10)

❓ An icon/badge POSITIONED NEAR an indexed element?
   → Use: hover_with_offset(index=N, direction="left/right/up/down")
   → Example: Hover status icon LEFT of "Enterprise Update Distributor" at index 18
   → Action: hover_with_offset(index=18, direction="left", distance=30)
   → ⚠️  CRITICAL: If task says "hover icon left of text" → USE hover_with_offset!

❓ Need to capture DOM state changes after hover?
   → Use: hover_capture(index=N)
   → Example: Hover element and capture what dynamically appears
   → Action: hover_capture(index=15)

━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━
🚨 COMMON MISTAKE TO AVOID:
━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━

❌ WRONG: Task says "hover status icon left of Enterprise Update Distributor"
          Action: hover_element(index=18)  ← This hovers the TEXT, not the icon!

✅ RIGHT: Task says "hover status icon left of Enterprise Update Distributor" 
          Action: hover_with_offset(index=18, direction="left", distance=30)
          ↑ This hovers 30px LEFT of the text, where the icon is!

━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━

CRITICAL RULES:
1. Execute ONLY the exact task requested. No extra steps.
2. Use retrieve_value_by_element to extract data, then validate_value to check it.
3. For clicks: prefer click_element_by_text(text) over indices.
4. Use scroll_page() to access content below fold.

FILE UPLOAD — HARD RULE:
- if any of the keywords “Choose File”, “Browse”, "upload", "upload file", or file selector is mentioned.
- You MUST use upload_file(index, path)
- Paths must come from available context files only

RETRIEVAL PATTERN:
Step 1: retrieve_value_by_element(index=N) → extracts plain text
Step 2: validate_value(actual=text, expected=value, operator="equals")

NEVER pass HTML to validate_value. Always retrieve first.

━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━
📊 TABLE DATA EXTRACTION - MANDATORY RULES
━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━
When extracting data from tables (e.g., agent statuses, server lists):
1. PREFERRED TOOL: Use find_table_data_rows(header_text="...") first. It returns the CONTENT of every row.
2. NON-INTERACTIVE TABLES: This tool works even if the table has no indices. Use its "rows" output to read data.
3. IDENTIFY HEADERS: If not using the tool, manually find the row containing column names (e.g., "Status").
4. EXCLUDE HEADERS: If retrieve_value_by_element(index=N) matches a header name, SKIP IT.
5. SKIP LAYOUT NOISE: Ignore rows that are empty, only contain spacing, or have fewer cells.
6. STALE INDICES: If you get a KEY ERROR, call get_clickable_elements or scroll.
━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━

━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━
🏁 TASK COMPLETION - MANDATORY
━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━
1. Once ALL points in the user's prompt are satisfied, you MUST call done(success=true).
2. Do NOT idle, scroll unnecessarily, or "think" about extra steps once the objective is met.
3. If you have extracted the required data or performed the required clicks, END THE SESSION.
4. Efficiency is priority. End the task immediately upon fulfillment.
━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━
"""

    # Simplified system prompt (no more giant walls of text)
    extend_system_prompt = """
━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━
🎯 HOVER DECISION TREE - READ THIS FIRST BEFORE ANY HOVER ACTION!
━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━

QUESTION: What are you trying to hover?

❓ The indexed element ITSELF (text, button, link)?
   → Use: hover_element(index=N)
   → Example: Hover "Submit" button at index 10
   → Action: hover_element(index=10)

❓ An icon/badge POSITIONED NEAR an indexed element?
   → Use: hover_with_offset(index=N, direction="left/right/up/down")
   → Example: Hover status icon LEFT of "Enterprise Update Distributor" at index 18
   → Action: hover_with_offset(index=18, direction="left", distance=30)
   → ⚠️  CRITICAL: If task says "hover icon left of text" → USE hover_with_offset!

❓ Need to capture DOM state changes after hover?
   → Use: hover_capture(index=N)
   → Example: Hover element and capture what dynamically appears
   → Action: hover_capture(index=15)

━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━
🚨 COMMON MISTAKE TO AVOID:
━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━

❌ WRONG: Task says "hover status icon left of Enterprise Update Distributor"
          Action: hover_element(index=18)  ← This hovers the TEXT, not the icon!

✅ RIGHT: Task says "hover status icon left of Enterprise Update Distributor" 
          Action: hover_with_offset(index=18, direction="left", distance=30)
          ↑ This hovers 30px LEFT of the text, where the icon is!

━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━

CORE RULES:
1. Execute ONLY the exact task requested. No extra steps.
2. Use retrieve_value_by_element to extract data, then validate_value to check it.
3. For clicks: prefer click_element_by_text(text) over indices.
4. Use scroll_page() to access content below fold.
5. Use go_back(), go_forward(), and refresh_page() for browser navigation.

━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━
🔐 LOGIN RULES - CRITICAL!
━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━
1. When you see a login form (username + password fields), use smart_login action.
2. After login submission: WAIT for page to transition.
3. LOGIN IS COMPLETE when:
   - Password field is no longer visible
   - Page title changed (e.g., "Login" → "Enterprise")  
   - You see authenticated content (menus, dashboards)
4. Once logged in: Call done() with success=true, DO NOT try to fill credentials again.
5. If you're on a dashboard/workspace page with no login fields → LOGIN ALREADY COMPLETE!
━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━

FILE UPLOAD — HARD RULE:
- If the element is “Choose File”, “Browse”, or file selector
- The element is <input type="file">
- input_text is FORBIDDEN
- You MUST use upload_file(index, path)
- Paths must come from available context files only

RETRIEVAL PATTERN:
Step 1: retrieve_value_by_element(index=N) → extracts plain text
Step 2: validate_value(actual=text, expected=value, operator="equals")

NEVER pass HTML to validate_value. Always retrieve first.

━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━
🔄 ANTI-LOOPING RULE — CRITICAL
━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━
If you tried the SAME action and it FAILED in a previous step:
  1. Do NOT retry the same action with the same parameters
  2. Try a DIFFERENT approach (different index, different action)
  3. If you've failed 2+ times on the same goal → call done() with
     whatever information you have so far
  
NEVER repeat a failed action more than once. Move on or finish.
━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━

━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━
🏁 TASK COMPLETION - MANDATORY
━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━
1. Once ALL points in the user's prompt are satisfied, you MUST call done(success=true).
2. Do NOT idle, scroll unnecessarily, or "think" about extra steps once the objective is met.
3. If you have extracted the required data or performed the required clicks, END THE SESSION.
4. Efficiency is priority. End the task immediately upon fulfillment.
━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━

━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━
🚫 extract_content IS BANNED — DO NOT USE IT EVER
━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━
extract_content is FORBIDDEN in ALL tasks. It dumps the entire page as
raw HTML/text and is never the right tool. It is always wrong.

INSTEAD — use the correct targeted action:
  ❌ extract_content  →  ✅ retrieve_value_by_element(index=N)
  ❌ extract_content  →  ✅ done(text="...") with data already visible on screen

If you are tempted to call extract_content, STOP.
Read what is already visible. Use retrieve_value_by_element for specific values.
Call done() when you have the information.

━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━
📊 TABLE DATA EXTRACTION - MANDATORY RULES
━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━
When extracting data from tables (e.g., agent statuses, server lists):
1. PREFERRED TOOL: Use find_table_data_rows(header_text="...") first. It returns the CONTENT of every row.
2. NON-INTERACTIVE TABLES: This tool works even if the table has no indices. Use its "rows" output to read data.
3. IDENTIFY HEADERS: If not using the tool, manually find the row containing column names (e.g., "Status").
4. EXCLUDE HEADERS: If retrieve_value_by_element(index=N) matches a header name, SKIP IT.
5. SKIP LAYOUT NOISE: Ignore rows that are empty, only contain spacing, or have fewer cells.
6. STALE INDICES: If you get a KEY ERROR, call get_clickable_elements or scroll.
━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━

━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━
🔒 SSL / CERTIFICATE TASKS — MANDATORY APPROACH
━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━
When the task mentions: "SSL certificate", "certificate details",
"certificate validity", "lock icon", "connection secure", "TLS", "HTTPS cert"

EXACT SEQUENCE — nothing else:
  1. go_to_url("the target URL")
  2. check_ssl_certificate()    ← registered action, use it
  3. done(success=true)

YOU MUST NOT:
  ❌ Click the lock icon (it is in the browser chrome — unreachable)
  ❌ Use extract_content for cert data
  ❌ Look for cert info in the DOM
  ❌ Try to open any security popup
  ❌ Use get_site_security_info (replaced by check_ssl_certificate)

"""


    stream_vh = int(70 * window_h // window_w)

    # Initialize LLM
    main_llm = await _initialize_llm(
        llm_provider_name, llm_model_name, llm_temperature, 
        llm_base_url, None, ollama_num_ctx
    )
    
    # Initialize planner model (optional)
    planner_llm = None
    # Planner disabled - only using main agent
    
    # Initialize extraction model (optional) - used for page_extraction_llm
    page_extraction_llm = None
    if extraction_provider and extraction_model_name:
        page_extraction_llm = await _initialize_llm(
            extraction_provider, extraction_model_name, extraction_temperature,
            llm_base_url, None, ollama_num_ctx if extraction_provider == "ollama" else None
        )

    async def ask_callback_wrapper(query: str, browser_context: BrowserContext) -> Dict[str, Any]:
        return await _ask_assistant_callback(webui_manager, query, browser_context)

    if not webui_manager.bu_controller:
        webui_manager.bu_controller = CustomController(ask_assistant_callback=ask_callback_wrapper)
        await webui_manager.bu_controller.setup_mcp_client(None)

    should_close_browser_on_finish = not keep_browser_open

    try:
        # Clean up old resources - always create fresh browser
        if webui_manager.bu_browser_context:
            await webui_manager.bu_browser_context.close()
            webui_manager.bu_browser_context = None
        if webui_manager.bu_browser:
            await webui_manager.bu_browser.close()
            webui_manager.bu_browser = None
        await asyncio.sleep(0.5)

        # Create fresh browser (always, not conditional)
        if use_own_browser:
            browser_binary_path = browser_binary_path or _detect_browser()
        webui_manager.bu_browser = CustomBrowser(config=BrowserConfig(
                headless=headless,
                disable_security=disable_security,
                browser_binary_path=browser_binary_path if use_own_browser else None,
                extra_browser_args=["--start-maximized"], # Correct field name for custom browser
            ))

        # Vault credential handling (MOVED UP)
        import re
        from src.utils.vault import vault
        vault_matches = re.findall(r'@vault\.([a-zA-Z0-9_]+)', task)
        sensitive_data = {}

        # Capture the vault prefix for script generation (first @vault.XXX in task wins).
        # This is passed to generate_llm_script so generated scripts call
        # maybe_login(page, vault_prefix="XXX") with the correct vault key,
        # instead of always falling back to VAULT_CREDENTIAL_PREFIX from config.
        _task_vault_prefix = vault_matches[0].upper() if vault_matches else None
        
        for vault_key in vault_matches:
            creds = vault.get_credentials(vault_key)
            if creds:
                replacement_parts = []
                for k, v in creds.items():
                    # CRITICAL: browser-use library expects <secret>KEY</secret> format
                    # NOT {{KEY}} format - the library's _replace_sensitive_data looks for XML tags
                    key_name = f"{vault_key.upper()}_{k.upper()}"
                    placeholder_for_task = f"<secret>{key_name}</secret>"
                    replacement_parts.append(f"{k.upper()}: {placeholder_for_task}")
                    sensitive_data[key_name] = v
                    
                    # Update global redaction filter
                    if hasattr(builtins, 'redacting_filter_values'):
                        builtins.redacting_filter_values.add(v)
                        
                replacement = ", ".join(replacement_parts)
                task = task.replace(f'@vault.{vault_key}', replacement)

        # Create context
        if not webui_manager.bu_browser_context:
            allowed_domains_str = os.getenv("ALLOWED_DOMAINS", "")
            allowed_domains = [d.strip() for d in allowed_domains_str.split(",") if d.strip()]
            
            context_config = BrowserContextConfig(
                trace_path=save_trace_path,
                save_recording_path=save_recording_path,
                save_downloads_path=save_download_path,
                allowed_domains=allowed_domains if allowed_domains else None,
                no_viewport=True, # Required for --start-maximized to work
            )
            if allowed_domains:
                logger.info(f"🔒 Browser restricted to domains: {allowed_domains}")
            webui_manager.bu_browser_context = await webui_manager.bu_browser.new_context(config=context_config)

        # Paths
        os.makedirs(os.path.join(save_agent_history_path, webui_manager.bu_agent_task_id), exist_ok=True)
        history_file = os.path.join(save_agent_history_path, webui_manager.bu_agent_task_id, f"{webui_manager.bu_agent_task_id}.json")
        gif_path = os.path.join(save_agent_history_path, webui_manager.bu_agent_task_id, f"{webui_manager.bu_agent_task_id}.gif")
        playwright_script_path = os.path.join(save_agent_history_path, webui_manager.bu_agent_task_id, f"{webui_manager.bu_agent_task_id}_playwright.py")

        webui_manager.bu_previous_tokens = 0
        report_generated = False

        async def step_callback_wrapper(state: BrowserState, output: AgentOutput, step_num: int):
            await _handle_new_step(webui_manager, state, output, step_num)

        async def done_callback_wrapper(history: AgentHistoryList):
            nonlocal report_generated
            _handle_done(webui_manager, history)
            report_generated = True

        # Create/update agent
        if not webui_manager.bu_agent:
            # Check if memory is enabled in environment
            enable_memory = os.getenv('ENABLE_MEMORY', 'false').lower() in ('true', '1', 'yes')
            
            webui_manager.bu_agent = BrowserUseAgent(
                task=task,
                llm=main_llm,
                planner_interval=0,  # Disabled - no planner
                page_extraction_llm=page_extraction_llm,
                available_file_paths=available_file_paths,
                include_attributes=['title', 'type', 'name', 'role', 'aria-label', 'placeholder', 'value', 'aria-expanded'],
                browser=webui_manager.bu_browser,
                browser_context=webui_manager.bu_browser_context,
                controller=webui_manager.bu_controller,
                register_new_step_callback=step_callback_wrapper,
                register_done_callback=done_callback_wrapper,
                use_vision=use_vision,
                max_input_tokens=max_input_tokens,
                max_actions_per_step=max_actions,
                enable_memory=enable_memory,
                sensitive_data=sensitive_data,
                extend_system_message=extend_system_prompt,  # appended to system prompt by prompts.py
            )
            # ALSO update the controller's sensitive data directly, 
            # because CustomController.input_text uses self.sensitive_data
            if webui_manager.bu_controller:
                webui_manager.bu_controller.sensitive_data = sensitive_data

            webui_manager.bu_agent.state.agent_id = webui_manager.bu_agent_task_id
            webui_manager.bu_agent.settings.generate_gif = gif_path
            # Note: save_playwright_script_path not supported in current browser-use version
        else:
            webui_manager.bu_agent.add_new_task(task)
            webui_manager.bu_agent.state.agent_id = webui_manager.bu_agent_task_id
            webui_manager.bu_agent.settings.generate_gif = gif_path
            # Note: save_playwright_script_path not supported in current browser-use version

        # Inject credentials and available file paths
        if webui_manager.bu_controller:
            webui_manager.bu_controller.sensitive_data = sensitive_data
            webui_manager.bu_controller.available_file_paths = available_file_paths
            logger.info(f"Set {len(available_file_paths)} available file paths for upload")

        webui_manager.bu_step_start_time = time.perf_counter()
        
        # Check for multi-process mode
        is_multi_process = any(keyword in task for keyword in ["Process 1:", "Process 2:", "Mandatory Process:", "\n---"])
        
        if is_multi_process:
             # Parse processes using the new orchestrator parser
             processes = parse_processes(task)
             logger.info(f"Parsing tasks for orchestration: {len(processes)} processes found")
             
             # Orchestration executor
             async def _run_orchestrated():
                 try:
                     # Add specialized callbacks for orchestrator
                     orchestrator_kwargs = {
                         'agent_history_path': save_agent_history_path,
                         'page_extraction_llm': page_extraction_llm,
                         'available_file_paths': available_file_paths,
                         'use_vision': use_vision,
                         'max_input_tokens': max_input_tokens,
                         'max_actions_per_step': max_actions,
                         'max_steps': max_steps,
                         'ask_callback': ask_callback_wrapper,
                         'step_callback': step_callback_wrapper,
                         'webui_manager': webui_manager,
                         'sensitive_data': sensitive_data,
                         'browser_context_config': context_config,
                         'extend_system_prompt': extend_system_prompt,
                     }
                     
                     results, output_dir = await run_all_processes(
                         processes=processes,
                         browser=webui_manager.bu_browser,
                         llm=main_llm,
                         **orchestrator_kwargs
                     )
                     
                     # Final summary in chat
                     from src.orchestrator.report_generator import generate_combined_playwright_script, generate_process_summary
                     
                     # Generate combined script
                     combined_script_path = os.path.join(output_dir, "combined_playwright_script.py")
                     generate_combined_playwright_script(results, combined_script_path)
                     
                     # Generate and show summary
                     summary_md = generate_process_summary(results)
                     webui_manager.bu_chat_history.append({
                         "role": "assistant", 
                         "content": f"## 🏁 Orchestration Complete\n\n{summary_md}\n\nAll artifacts saved to: `{output_dir}`"
                     })
                     
                     # Update UI with combined script if successful
                     if os.path.exists(combined_script_path):
                         try:
                             with open(combined_script_path, 'r', encoding='utf-8') as f:
                                 script_content = f.read()
                             webui_manager.bu_playwright_script = script_content
                         except: pass
                         
                 except Exception as e:
                     logger.error(f"Orchestration failed: {e}", exc_info=True)
                     webui_manager.bu_chat_history.append({
                         "role": "assistant", 
                         "content": f"❌ **Orchestration Failed**: {str(e)}"
                     })
                     
             task_func = _run_orchestrated
        else:
             # Standard single agent flow
             webui_manager.bu_agent.state.agent_id = webui_manager.bu_agent_task_id
             webui_manager.bu_agent.settings.generate_gif = gif_path
             task_func = lambda: webui_manager.bu_agent.run(max_steps=max_steps)

        # Launch task
        agent_task = asyncio.create_task(task_func())
        webui_manager.bu_current_task = agent_task

        last_chat_len = len(webui_manager.bu_chat_history)

        while not agent_task.done():
            is_paused = webui_manager.bu_agent.state.paused
            is_stopped = webui_manager.bu_agent.state.stopped

            if is_paused:
                yield {pause_resume_button_comp: gr.update(value="▶️ Resume", interactive=True)}
                while is_paused and not agent_task.done():
                    is_paused = webui_manager.bu_agent.state.paused
                    is_stopped = webui_manager.bu_agent.state.stopped
                    if is_stopped:
                        break
                    await asyncio.sleep(0.2)
                if agent_task.done() or is_stopped:
                    break
                yield {pause_resume_button_comp: gr.update(value="⸏ Pause", interactive=True)}

            if is_stopped:
                logger.info("Agent stopped.")
                break

            update_dict = {}
            if webui_manager.bu_response_event:
                update_dict = {
                    user_input_comp: gr.update(placeholder="Enter response...", interactive=True),
                    run_button_comp: gr.update(value="✔️ Submit Response", interactive=True),
                    chatbot_comp: gr.update(value=webui_manager.bu_chat_history),
                }
                last_chat_len = len(webui_manager.bu_chat_history)
                yield update_dict
                await webui_manager.bu_response_event.wait()
                if not agent_task.done():
                    yield {
                        user_input_comp: gr.update(placeholder="Running...", interactive=False),
                        run_button_comp: gr.update(value="⏳ Running...", interactive=False),
                    }

            if len(webui_manager.bu_chat_history) > last_chat_len:
                update_dict[chatbot_comp] = gr.update(value=webui_manager.bu_chat_history)
                last_chat_len = len(webui_manager.bu_chat_history)

            if headless and webui_manager.bu_browser_context:
                try:
                    screenshot_b64 = await webui_manager.bu_browser_context.take_screenshot()
                    if screenshot_b64:
                        # Apply Chrome frame to live streaming view
                        try:
                            _page = await webui_manager.bu_browser_context.get_current_page()
                            _url = _page.url or ""
                            if _url and not _url.startswith("about:") and not _url.startswith("chrome:"):
                                _title = await _page.title() or _url
                                _vp = _page.viewport_size or {}
                                _pw_ctx = webui_manager.bu_browser_context.session.context
                                screenshot_b64 = await apply_chrome_frame(
                                    browser=_pw_ctx,
                                    screenshot_b64=screenshot_b64,
                                    url=_url,
                                    title=_title,
                                    is_secure=_url.startswith("https://"),
                                    page_width=_vp.get("width", 1280),
                                    page_height=_vp.get("height", 900),
                                )
                        except Exception:
                            pass  # fallback: show raw screenshot without frame
                        html_content = f'<img src="data:image/png;base64,{screenshot_b64}" style="width:{stream_vw}vw;height:{stream_vh}vh;border:1px solid #ccc;">'
                        update_dict[browser_view_comp] = gr.update(value=html_content, visible=True)
                except Exception:
                    pass

            if update_dict:
                yield update_dict
            await asyncio.sleep(0.5)

        # Finalize
        webui_manager.bu_agent.state.paused = False
        webui_manager.bu_agent.state.stopped = False
        final_update = {}
        
        try:
            if not agent_task.done():
                await agent_task
            elif agent_task.exception():
                agent_task.result()
            
            # Force report generation
            if not report_generated and webui_manager.bu_agent and webui_manager.bu_agent.state.history:
                _handle_done(webui_manager, webui_manager.bu_agent.state.history)
                report_generated = True

            # Save history
            webui_manager.bu_agent.save_history(history_file)
            if os.path.exists(history_file):
                final_update[history_file_comp] = gr.File(value=history_file)
            
            # Enrich history with DOM context and generate Playwright script
            try:
                logger.info("Enriching history with DOM context...")
                strip_and_enrich(pathlib.Path(history_file))
                logger.info("History enrichment complete")
                
                logger.info("Generating Playwright script via LLM...")
                history_path = pathlib.Path(history_file)
                from src.utils.config import SCRIPT_GEN_MODEL, SCRIPT_GEN_PROVIDER
                logger.info(f"Script gen model: {SCRIPT_GEN_PROVIDER}/{SCRIPT_GEN_MODEL} (agent model: {llm_provider_name}/{llm_model_name})")
                logger.info(f"Script vault prefix: {_task_vault_prefix or '(using config default)'}")
                script_path, script_content = generate_llm_script(
                    str(history_path),
                    model_name=SCRIPT_GEN_MODEL,
                    provider=SCRIPT_GEN_PROVIDER,
                    objective=task,
                    vault_prefix=_task_vault_prefix,
                )
                logger.info(f"Playwright script saved to {script_path}")
                logger.info(f"Script is ready to run: python {os.path.basename(script_path)}")
            except Exception as e:
                logger.warning(f"Could not auto-generate enrichment or script: {e}", exc_info=True)

            # GIF
            if gif_path and os.path.exists(gif_path):
                final_update[gif_comp] = gr.Image(value=gif_path)

            # Playwright script - Use the content directly from the generator
            if 'script_content' in locals() and script_content:
                final_update[playwright_script_comp] = gr.Code(value=script_content, language="python")
            else:
                # Fallback to file search if variable is missing
                custom_script_path = history_file.replace('.json', '_LLM.py')
                if not os.path.exists(custom_script_path):
                    custom_script_path = history_file.replace('.json', '.py')
                
                if os.path.exists(custom_script_path):
                    with open(custom_script_path, 'r', encoding='utf-8') as f:
                        script_content = f.read()
                    final_update[playwright_script_comp] = gr.Code(value=script_content, language="python")
                else:
                    final_update[playwright_script_comp] = gr.Code(value="# Script generation not available or disabled", language="python")

            # DOCX
            docx_path = getattr(webui_manager, 'bu_docx_path', None)
            if docx_path and os.path.exists(docx_path):
                final_update[docx_report_comp] = gr.File(value=docx_path)

        except asyncio.CancelledError:
            logger.info("Task cancelled.")
            final_update[chatbot_comp] = gr.update(value=webui_manager.bu_chat_history)
        except Exception as e:
            logger.error(f"Execution error: {e}", exc_info=True)
            webui_manager.bu_chat_history.append({"role": "assistant", "content": f"**Error:** {e}"})
            final_update[chatbot_comp] = gr.update(value=webui_manager.bu_chat_history)
        finally:
            webui_manager.bu_current_task = None
            
            # Always clean up agent to prevent state pollution
            webui_manager.bu_agent = None
            
            if should_close_browser_on_finish:
                logger.info("Closing browser and context...")
                if webui_manager.bu_browser_context:
                    await webui_manager.bu_browser_context.close()
                    webui_manager.bu_browser_context = None
                if webui_manager.bu_browser:
                    # Force kill browser to ensure it closes even if connected via CDP
                    await webui_manager.bu_browser.close(force=True)
                    webui_manager.bu_browser = None
                logger.info("Browser closed successfully")

            final_update.update({
                user_input_comp: gr.update(value="", interactive=True, placeholder="Enter next task..."),
                run_button_comp: gr.update(value="▶️ Submit Task", interactive=True),
                stop_button_comp: gr.update(interactive=False),
                pause_resume_button_comp: gr.update(interactive=False),
                clear_button_comp: gr.update(interactive=True),
                chatbot_comp: gr.update(value=webui_manager.bu_chat_history),
            })
            yield final_update

    except Exception as e:
        logger.error(f"Setup error: {e}", exc_info=True)
        yield {
            user_input_comp: gr.update(interactive=True),
            run_button_comp: gr.update(value="▶️ Submit Task", interactive=True),
            chatbot_comp: gr.update(value=webui_manager.bu_chat_history + [{"role": "assistant", "content": f"**Error:** {e}"}])
        }

async def handle_submit(webui_manager: WebuiManager, components: Dict[Component, Any]):
    """Handle submit button clicks."""
    user_input_comp = webui_manager.get_component_by_id("browser_use_agent.user_input")
    user_input_value = components.get(user_input_comp, "").strip()

    # Check if waiting for assistance
    if webui_manager.bu_response_event and not webui_manager.bu_response_event.is_set():
        logger.info(f"User submitted assistance: {user_input_value}")
        webui_manager.bu_user_help_response = user_input_value if user_input_value else "No response."
        webui_manager.bu_response_event.set()
        yield {
            user_input_comp: gr.update(value="", interactive=False, placeholder="Waiting..."),
            webui_manager.get_component_by_id("browser_use_agent.run_button"): gr.update(value="⏳ Running...", interactive=False),
        }
    elif webui_manager.bu_current_task and not webui_manager.bu_current_task.done():
        logger.warning("Submit clicked while running.")
        gr.Info("Agent already running.")
        yield {}
    else:
        # New task
        async for update in run_agent_task(webui_manager, components):
            yield update


async def handle_stop(webui_manager: WebuiManager):
    """Handle stop button."""
    logger.info("Stop clicked.")
    agent = webui_manager.bu_agent
    task = webui_manager.bu_current_task

    if agent and task and not task.done():
        agent.state.stopped = True
        agent.state.paused = False
        return {
            webui_manager.get_component_by_id("browser_use_agent.stop_button"): gr.update(interactive=False, value="⏹️ Stopping..."),
            webui_manager.get_component_by_id("browser_use_agent.pause_resume_button"): gr.update(interactive=False),
        }
    return {
        webui_manager.get_component_by_id("browser_use_agent.run_button"): gr.update(interactive=True),
        webui_manager.get_component_by_id("browser_use_agent.stop_button"): gr.update(interactive=False),
    }


async def handle_pause_resume(webui_manager: WebuiManager):
    """Handle pause/resume."""
    agent = webui_manager.bu_agent
    task = webui_manager.bu_current_task

    if agent and task and not task.done():
        if agent.state.paused:
            logger.info("Resume.")
            agent.resume()
            return {webui_manager.get_component_by_id("browser_use_agent.pause_resume_button"): gr.update(value="⸏ Pause")}
        else:
            logger.info("Pause.")
            agent.pause()
            return {webui_manager.get_component_by_id("browser_use_agent.pause_resume_button"): gr.update(value="▶️ Resume")}
    return {}


async def handle_clear(webui_manager: WebuiManager):
    """Handle clear button."""
    logger.info("Clear clicked.")
    
    task = webui_manager.bu_current_task
    if task and not task.done():
        if webui_manager.bu_agent:
            webui_manager.bu_agent.stop()
        task.cancel()
        try:
            from src.utils.config import TASK_CANCEL_TIMEOUT_S
            await asyncio.wait_for(task, timeout=TASK_CANCEL_TIMEOUT_S)
        except Exception:
            pass
    
    webui_manager.bu_current_task = None
    if webui_manager.bu_controller:
        await webui_manager.bu_controller.close_mcp_client()
        webui_manager.bu_controller = None
    webui_manager.bu_agent = None
    webui_manager.bu_chat_history = []
    webui_manager.bu_response_event = None
    webui_manager.bu_user_help_response = None
    webui_manager.bu_agent_task_id = None

    return {
        webui_manager.get_component_by_id("browser_use_agent.chatbot"): gr.update(value=[]),
        webui_manager.get_component_by_id("browser_use_agent.user_input"): gr.update(value="", placeholder="Enter task..."),
        webui_manager.get_component_by_id("browser_use_agent.agent_history_file"): gr.update(value=None),
        webui_manager.get_component_by_id("browser_use_agent.recording_gif"): gr.update(value=None),
        webui_manager.get_component_by_id("browser_use_agent.browser_view"): gr.update(value="<div>Cleared</div>"),
        webui_manager.get_component_by_id("browser_use_agent.run_button"): gr.update(value="▶️ Submit Task", interactive=True),
        webui_manager.get_component_by_id("browser_use_agent.stop_button"): gr.update(interactive=False),
        webui_manager.get_component_by_id("browser_use_agent.pause_resume_button"): gr.update(interactive=False),
        webui_manager.get_component_by_id("browser_use_agent.clear_button"): gr.update(interactive=True),
    }


def create_browser_use_agent_tab(webui_manager: WebuiManager):
    """Create the run agent tab."""
    webui_manager.init_browser_use_agent()

    tab_components = {}
    with gr.Column():
        user_input = gr.Textbox(
            label="Task or Response",
            placeholder="Enter your task here...",
            lines=3,
            interactive=True,
            elem_id="user_input",
        )

        context_files = gr.File(
            label="Upload Context Files",
            file_types=[".pdf", ".txt", ".docx", ".md", ".csv", ".json"],
            file_count="multiple",
            interactive=True,
            elem_id="context_files",
        )

        with gr.Row():
            stop_button = gr.Button("⏹️ Stop", interactive=False, variant="stop", scale=2)
            pause_resume_button = gr.Button("⸏ Pause", interactive=False, variant="secondary", scale=2)
            clear_button = gr.Button("🗑️ Clear", interactive=True, variant="secondary", scale=2)
            run_button = gr.Button("▶️ Submit Task", variant="primary", scale=3)

        chatbot = gr.Chatbot(
            lambda: webui_manager.bu_chat_history,
            elem_id="browser_use_chatbot",
            label="Agent Interaction",
            height=600,
            type="messages",
        )

        browser_view = gr.HTML(
            value="<div style='width:100%;height:50vh;display:flex;justify-content:center;align-items:center;border:1px solid #ccc;background:#f0f0f0;'><p>Browser View (Headless mode)</p></div>",
            label="Browser Live View",
            visible=False,
        )
        
        with gr.Column():
            gr.Markdown("### Task Outputs")
            agent_history_file = gr.File(label="Agent History JSON", interactive=False)
            recording_gif = gr.Image(label="Task Recording GIF", interactive=False)
            playwright_script = gr.Code(label="Playwright Script", language="python", interactive=False, lines=20)
            docx_report = gr.File(label="DOCX Report", interactive=False)

    tab_components.update(dict(
        chatbot=chatbot,
        user_input=user_input,
        context_files=context_files,
        clear_button=clear_button,
        run_button=run_button,
        stop_button=stop_button,
        pause_resume_button=pause_resume_button,
        agent_history_file=agent_history_file,
        recording_gif=recording_gif,
        browser_view=browser_view,
        playwright_script=playwright_script,
        docx_report=docx_report,
    ))
    
    webui_manager.add_components("browser_use_agent", tab_components)
    all_managed_components = list(webui_manager.get_components())
    run_tab_outputs = list(tab_components.values())

    async def submit_wrapper(*args) -> AsyncGenerator[Dict[Component, Any], None]:
        # Reconstruct the components dict from unpacked arguments
        components_dict = {component: value for component, value in zip(all_managed_components, args)}
        async for update in handle_submit(webui_manager, components_dict):
            yield update

    async def stop_wrapper() -> AsyncGenerator[Dict[Component, Any], None]:
        update_dict = await handle_stop(webui_manager)
        yield update_dict

    async def pause_resume_wrapper() -> AsyncGenerator[Dict[Component, Any], None]:
        update_dict = await handle_pause_resume(webui_manager)
        yield update_dict

    async def clear_wrapper() -> AsyncGenerator[Dict[Component, Any], None]:
        update_dict = await handle_clear(webui_manager)
        yield update_dict

    run_button.click(fn=submit_wrapper, inputs=all_managed_components, outputs=run_tab_outputs, trigger_mode="multiple", concurrency_limit=1)
    user_input.submit(fn=submit_wrapper, inputs=all_managed_components, outputs=run_tab_outputs)
    stop_button.click(fn=stop_wrapper, inputs=None, outputs=run_tab_outputs)
    pause_resume_button.click(fn=pause_resume_wrapper, inputs=None, outputs=run_tab_outputs)
    clear_button.click(fn=clear_wrapper, inputs=None, outputs=run_tab_outputs)

    logger.info("✅ Browser Use Agent tab created")