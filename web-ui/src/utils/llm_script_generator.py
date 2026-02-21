"""
LLM Script Generator for Browser Agent
Uses an LLM (via llm_provider) to generate robust Playwright scripts from agent history.
Key Features:
- Redundancy Elimination (Smart Login)
- Context-Aware Code Generation
- Professional Report Integration
"""

import sys
import os
import json
import logging
import argparse
import re
from pathlib import Path
from typing import Dict, Any

# Ensure we can import from src
sys.path.append(str(Path(__file__).parent.parent.parent))

from src.utils.llm_provider import get_llm_model
from langchain_core.messages import SystemMessage, HumanMessage

# Configure logging
logging.basicConfig(level=logging.INFO, format='%(asctime)s - %(levelname)s - %(message)s')
logger = logging.getLogger(__name__)

SYSTEM_PROMPT = r'''
You are an expert Playwright automation engineer. Convert execution history into a ROBUST and FAST Python script.

### 🛡️ GOLDEN RULES
1.  **Vault**: Use `get_secret(key)`. Never hardcode. Prefix is dynamic (e.g., 'OTCS_USER' -> key 'OTCS').
2.  **Atomic Login**: 
    - **Custom Action**: If the history contains `smart_login`, generate a clean block that calls `get_secret` and uses `resilient_fill` for both UN/PW followed by a click.
    - **Detection**: Even if `smart_login` is missing but UN/PW fields are seen sequentially, **Consolidate** them. Combine username, password, and the login click into ONE block. 
    - **Wait**: Always `await page.wait_for_selector('input[type="password"]', timeout=5000)` before filling credentials to handle redirects.
    - **Skip Redundancy**: Completely ignore manual login steps or trial-and-error attempts found in the history once logged in.
    - **Login Hook**: After navigating to a protected URL, call `await maybe_login(page)` before data extraction.
3.  **Efficiency & Accuracy**: 
    - Use `resilient_fill` and `resilient_click`.
    - **URL Integrity**: Ensure URLs do NOT have double dots (e.g., `host..com`) or common path typos (e.g., `ccs.exe` for `cs.exe`).
    - **Selector Priority**: Use strictly **text** or **title** if it matches the goal (e.g., "Goal Name"). DO NOT use internal system labels (like "Secondary Label") if a clearer label exists in memory.
    - **Pruning**: Do NOT include failed actions, typos, or trial-and-error clicks. If the agent failed a click then succeeded, generate ONLY the successful one. 
    - **Uploads**: MUST use `click_and_upload_file` with a robust timeout (at least 30s).
    - **Exact Match**: Use regex anchors `re.compile(f"^{target}$", re.I)` for ambiguous menus.
4.  **Browser**: Start maximized with `no_viewport=True`.
5.  **Universal Resilience**:
    - **Action Elements**: For interactions (e.g., clicking icons or toolbar items), prioritize `[aria-label="Name"]` or `[title="Name"]` over tag-specific selectors like `button:text()`. 
    - **Tag Agnosticism**: In modern SPAs, actions can be `<a>`, `button`, or `div`. Generate code that handles this by focusing on descriptive attributes and labels found in the DOM context.
    - **Dynamic Discovery**: Never hardcode site-specific logic. Apply these principles dynamically based on the provided history and element metadata.
7.  **TABLE EXTRACTION - CRITICAL**:
    - **Resilient Scraper**: ALWAYS use `await resilient_extract_table(page, table_selector, ["Header1", "Header2"])` for data scraping.
    - **Strict Row Validation & Noise Elimination**: 
        - **Header/Nav Filtering**: Skip rows if any cell contains common boilerplate like "Default", "Name", "Description", "Enterprise", "Personal", "Tools", "Status", or "Header".
        - **Data Signature Validation**: Only include a row if it matches a valid data pattern (e.g., status is in a list of known states like 'Active', 'Standby', or the metric/error column is numeric).
        - **Targeted Row Check**: If possible, verify the first cell contains expected keywords for the resource (e.g., 'Server', 'Partition', 'Agent').
    - **Cell Text Extraction**: ALWAYS use `(await cell.inner_text() or "").strip()` instead of `text_content()`.
    - **Structured Output**: Format as clean report lines and pass to `generate_report(..., output="\n".join(extracted_lines))`.
    - **STOP ON COMPLETION**: If the history shows successful data extraction matching the goal, END the script immediately.
    - **STRICT LOOP LOGIC**: Do NOT use `async for row in locator`. ALWAYS use: `elements = await locator.all()` then `for element in elements:`.
    - **STRICT IMPORTS**: You MUST include `from docx import Document` at the top of the script.
    - **DOM CONTEXT MANDATORY**: If `dom_elements` or `selectors` are provided in the cleaned history, you MUST scope extraction to those selectors. Do NOT use global `page.locator("tr")` or `page.locator("table")` without scoping to the provided selector.
    - **REQUIRED COLUMNS**: If `required_columns_hint` is present, you MUST extract those columns and include them in the output.
    - **TABLE SCOPING**: Always start from the most stable selector in `dom_elements.selectors` (prefer `preferred` or `css`) and then locate rows within that table only.
    - **WAIT FOR TABLE**: Prefer `await page.wait_for_selector(table_selector, timeout=10000)` over fixed sleep when possible.

### 🛡️ GOLDEN RULES (REVISITED)
1.  **Vault**: Use `get_secret(key)`. Never hardcode. 
2.  **Uploads**: MUST use `click_and_upload_file` with a robust timeout (at least 30s).
3.  **URL Integrity**: Ensure URLs do NOT have double dots or common path typos.

### 📝 STRUCTURE
Generate a script following this boilerplate exactly. Ensure ALL imports (sys, re, io, time) are present.

```python
import asyncio
import io
import sys
import os
import re
import time
from playwright.async_api import async_playwright
from docx import Document

# Setup path to find 'src'
current_dir = os.getcwd()
if "web-ui" not in current_dir and os.path.exists(os.path.join(current_dir, "web-ui")):
    sys.path.append(os.path.join(current_dir, "web-ui"))
else:
    sys.path.append(current_dir)

from src.utils.vault import vault

async def get_secret(key):
    parts = key.split("_")
    v_key = parts[0] if len(parts) > 1 else key
    creds = vault.get_credentials(v_key) or vault.get_credentials(v_key.lower())
    if not creds: return None
    if any(k in key.upper() for k in ["USERNAME", "USER"]): return creds.get("username")
    if any(k in key.upper() for k in ["PASSWORD", "PWD"]): return creds.get("password")
    return None

async def find_in_frames(page, selector, timeout=5000):
    """Search for a selector in the main page and all its iframes."""
    # Try main page first
    try:
        el = page.locator(selector).first
        if await el.count() > 0 and await el.is_visible(timeout=1000):
            return el
    except: pass

    # Try all frames
    for frame in page.frames:
        try:
            el = frame.locator(selector).first
            if await el.count() > 0 and await el.is_visible(timeout=1000):
                return el
        except: continue
    return None

async def resilient_fill(page, selector, value):
    if value is None: return
    selectors = selector if isinstance(selector, list) else [selector]
    for sel in selectors:
        try:
            print(f"[*] Filling {sel}...")
            el = await find_in_frames(page, sel)
            if not el:
                await page.wait_for_selector(sel, timeout=3000)
                el = page.locator(sel).first
            await el.fill(str(value))
            await el.dispatch_event("input")
            await el.dispatch_event("change")
            return True
        except Exception:
            continue
    print(f"[-] Fill failed for {selectors}")
    return False

async def resilient_click(page, selectors, desc="element"):
    if isinstance(selectors, str): selectors = [selectors]
    for sel in selectors:
        try:
            print(f"[*] Attempting click on {desc} via {sel}...")
            el = await find_in_frames(page, sel)
            if el:
                await el.click(timeout=3000)
                return True
        except: continue
    return False

async def click_and_upload_file(page, target_name, file_path):
    print(f"[*] Starting robust upload for: {target_name}")
    try:
        if not os.path.exists(file_path): 
            print(f"[-] File not found: {file_path}")
            return False
        # Strategy 1: Native Chooser via Strict Match
        try:
            target = page.locator("a, button, [role='button']").filter(has_text=re.compile(f"^{target_name}$", re.I)).first
            async with page.expect_file_chooser(timeout=7000) as fc_info:
                await target.click(timeout=5000)
            file_chooser = await fc_info.value
            await file_chooser.set_files(file_path)
            return True
        except: pass
        # Strategy 2: Hidden File Input
        try:
            await page.locator('input[type="file"]').first.set_input_files(file_path)
            return True
        except: pass
        # Strategy 3: Global Text Search
        try:
            target = page.get_by_text(target_name, exact=True).first
            async with page.expect_file_chooser(timeout=5000) as fc_info:
                await target.click()
            file_chooser = await fc_info.value
            await file_chooser.set_files(file_path)
            return True
        except: pass
        return False
    except Exception as e:
        print(f"[-] Upload attempt failed: {e}")
        return False

async def maybe_login(page):
    """Attempt login if a login form is detected."""
    username = await get_secret("OTCS_USERNAME")
    password = await get_secret("OTCS_PASSWORD")
    if not username or not password:
        return

    user_selectors = [
        "input[name='username']",
        "input#username",
        "input[name='user']",
        "input[type='text']"
    ]
    pass_selectors = [
        "input[name='password']",
        "input#password",
        "input[type='password']"
    ]
    submit_selectors = [
        "button:has-text('Sign in')",
        "button:has-text('Sign In')",
        "input[type='submit']",
        "button[type='submit']"
    ]

    # Only attempt if a password field is visible somewhere
    pass_el = None
    for sel in pass_selectors:
        pass_el = await find_in_frames(page, sel)
        if pass_el:
            break
    if not pass_el:
        return

    await resilient_fill(page, user_selectors, username)
    await resilient_fill(page, pass_selectors, password)
    await resilient_click(page, submit_selectors, desc="login submit")

async def resilient_extract_table(page, table_selector, header_keywords, data_row_keywords=None, skip_boilerplate=None):
    """Robustly extract table data using dynamic header resolution and strict filtering."""
    # Scope to the real table container from dom_context
    table = await find_in_frames(page, table_selector)
    if not table:
        raise RuntimeError(f"Table selector not found: {table_selector}")

    rows = await table.locator("tr").all()
    
    # If no rows in main page, check frames using the same table selector
    if len(rows) < 3:
        for frame in page.frames:
            try:
                f_table = frame.locator(table_selector).first
                if await f_table.count() > 0:
                    f_rows = await f_table.locator("tr").all()
                    if len(f_rows) > len(rows):
                        rows = f_rows
            except Exception:
                continue
    
    print(f"[*] Analyzing {len(rows)} rows for table extraction...")
    
    # 1. Resolve Headers
    col_map = {kw: -1 for kw in header_keywords}
    for i in range(min(20, len(rows))):
        cells = await rows[i].locator("th, td").all()
        if len(cells) < len(header_keywords): continue
        h_texts = [(await c.inner_text() or "").strip() for c in cells]
        
        for kw in header_keywords:
            if col_map[kw] != -1: continue
            idx = next((j for j, txt in enumerate(h_texts) if kw.lower() in txt.lower()), -1)
            if idx != -1: col_map[kw] = idx
            
        if all(v != -1 for v in col_map.values()):
            print(f"[*] Resolved headers: {col_map}")
            break

    # 2. Extract Data
    results = []
    boilerplate = skip_boilerplate or ["Default", "Name", "Description", "Enterprise", "Personal", "Tools", "Status", "Header"]
    
    for row in rows:
        cells = await row.locator("td").all()
        if len(cells) <= max(col_map.values()): continue
        
        row_data = {}
        is_valid = True
        all_text = " ".join([(await c.inner_text() or "").strip() for c in cells])
        
        # Skip header/nav rows
        if any(b in all_text for b in boilerplate): continue
        
        for kw, idx in col_map.items():
            cell = cells[idx]
            # Check for image titles (status icons)
            img = cell.locator("img[title]").first
            text = await img.get_attribute("title") if await img.count() > 0 else (await cell.inner_text() or "").strip()
            
            if not text or text == "null": is_valid = False; break
            row_data[kw] = text
            
        if is_valid:
            results.append(row_data)
            
    return results

async def generate_report(scenario, status, screenshot=None, output=None):
    try:
        # Print extracted data to terminal
        print("\n" + "="*60)
        print(f"📋 EXTRACTED DATA: {scenario}")
        print("="*60)
        if output:
            print(output)
        print("="*60 + "\n")
        
        # Generate timestamped filename in script's folder
        from datetime import datetime
        timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
        script_dir = os.path.dirname(os.path.abspath(__file__))
        report_path = os.path.join(script_dir, f"report_{timestamp}.docx")
        
        doc = Document()
        doc.add_heading(f"Report: {scenario}", 0)
        doc.add_paragraph(f"Generated: {datetime.now().strftime('%Y-%m-%d %H:%M:%S')}")
        doc.add_paragraph(f"Status: {'SUCCESS' if status else 'FAILED'}")
        if output: doc.add_paragraph(f"Output:\n{output}")
        if screenshot: doc.add_picture(io.BytesIO(screenshot))
        doc.save(report_path)
        print(f"[+] Report generated: {report_path}")
    except Exception as e:
        print(f"[-] Report generation failed: {e}")

async def run():
    async with async_playwright() as p:
        # Launch browser in MAXIMIZED mode
        browser = await p.chromium.launch(headless=False, args=["--start-maximized"])
        context = await browser.new_context(no_viewport=True)
        page = await context.new_page()
        
        # --- Task Logic ---
        
        await browser.close()

if __name__ == "__main__":
    asyncio.run(run())
'''


def clean_history_json(history_data: Dict[str, Any], objective: str = None) -> str:
    """Optimize JSON payload by pruning redundant actions, failed attempts, and duplicate navigations."""
    def columns_from_results(results):
        if not results:
            return []
        for res in results:
            if not isinstance(res, dict):
                continue
            content = res.get("extracted_content") or ""
            m = re.search(r'```json\s*(.*?)\s*```', content, re.DOTALL)
            if not m:
                continue
            try:
                data = json.loads(m.group(1))
            except Exception:
                continue
            cols = set()
            def walk(obj):
                if isinstance(obj, dict):
                    if obj and all(not isinstance(v, dict) for v in obj.values()):
                        for k in obj.keys():
                            cols.add(k)
                    else:
                        for v in obj.values():
                            walk(v)
                elif isinstance(obj, list):
                    for v in obj:
                        walk(v)
            walk(data)
            if cols:
                return list(cols)
        return []
    simplified = []
    seen_urls = []

    def infer_required_columns(text: str) -> list[str]:
        """
        Infer required columns from the user's objective/goal without hardcoded domain terms.
        Strategy:
          1) Extract quoted phrases (single or double quotes).
          2) Extract list after keywords like "columns", "fields", "extract".
          3) Extract Title Case multi-word phrases as a fallback.
        """
        if not text:
            return []

        candidates = []

        # 1) Quoted phrases: "Status", 'Errors', etc.
        candidates += re.findall(r'"([^"]{2,})"', text)
        candidates += re.findall(r"'([^']{2,})'", text)

    # 2) List after delimiters (colon or dash), derived from the text itself
        for line in re.split(r'[\r\n]+', text):
            if ":" in line or " - " in line:
                tail = re.split(r'[:\-]\s*', line, maxsplit=1)[-1]
                parts = re.split(r',| and ', tail, flags=re.IGNORECASE)
                candidates += [p.strip() for p in parts if p.strip()]

        # 3) Title Case phrases (e.g., "Enterprise Update Distributor")
        candidates += re.findall(r'([A-Z][a-z]+(?:\s+[A-Z][a-z]+)*)', text)

        # Normalize and filter noise using dynamic stopwords from the text
        tokens = re.findall(r'[A-Za-z]+', text.lower())
        freq = {}
        for t in tokens:
            freq[t] = freq.get(t, 0) + 1
        common = sorted(freq, key=freq.get, reverse=True)[:15]
        stopwords = set(common)
        cleaned = []
        seen = set()
        for c in candidates:
            c = c.strip().strip(".:")
            if not c or c in stopwords:
                continue
            if len(c) < 2:
                continue
            if c not in seen:
                seen.add(c)
                cleaned.append(c)

        return cleaned
    
    for item in history_data.get('history', []):
        state = item.get('state', {})
        model_out = item.get('model_output', {})
        action = model_out.get('action', [])
        evaluation = model_out.get('current_state', {}).get('evaluation_previous_goal', '')
        
        current_url = state.get('url', '')
        if current_url:
            # Simple URL deduplication: if the URL is logically the same as the last one, mark as redundant
            # (ignoring minor query param diffs or typos that redirect)
            if seen_urls and seen_urls[-1].split('?')[0] == current_url.split('?')[0]:
                is_redundant_nav = True
            else:
                is_redundant_nav = False
                seen_urls.append(current_url)
        else:
            is_redundant_nav = False

        step_info = {
            "url": current_url,
            "title": state.get('title'),
            "evaluation": evaluation,
            "goal": model_out.get('current_state', {}).get('next_goal'),
            "actions": [],
            "dom_elements": []
        }
        
        # Extract actions and their context
        for act_idx, act in enumerate(action):
            # PRUNING: Only skip if there was an actual error AND it wasn't a core interaction
            # Sometimes 'timeouts' happen but the action actually worked (SPA redirects)
            # We keep 'click', 'input', 'upload' unless they are followed by a successful retry
            results = item.get('result', [])
            if act_idx < len(results):
                res = results[act_idx]
                if isinstance(res, dict) and res.get('error'):
                    error_msg = res.get('error', '').lower()
                    action_type = next(iter(act.keys())) if act else "unknown"
                    
                    # If it's a timeout for a core action, keep it - the LLM can decide
                    if "timeout" in error_msg and action_type in ['click_element', 'input_text', 'upload_file', 'click_element_by_text']:
                        logger.info(f"Keeping core action {action_type} despite timeout: {error_msg}")
                    else:
                        logger.info(f"Pruning action {act_idx} due to error: {res.get('error')}")
                        continue

            # Skip redundant navigations if they were just retries of the same URL
            if is_redundant_nav and "go_to_url" in act:
                continue

            # Create a shallow copy of action to exclude dom_context from the 'actions' list
            act_copy = {k: v for k, v in act.items() if k != 'dom_context'}
            step_info['actions'].append(act_copy)
            
            if 'dom_context' in act:
                ctx = act['dom_context']
                selectors = ctx.get('selectors', {})
                attrs = ctx.get('attributes', {})
                comp = ctx.get('comprehensive', {})
                
                # ENRICHED ELEMENT DATA: Pass all critical selectors to the LLM
                # Also look into 'comprehensive' data which often has better labels
                comp_attrs = comp.get('all_attributes', {})
                comp_selectors = {s['type']: s['selector'] for s in comp.get('all_selectors', [])}
                
                element_data = {
                    "tag": ctx.get('tagName') or comp.get('tag'),
                    "text": selectors.get('text', '') or attrs.get('textContent', '') or comp_attrs.get('title', '') or comp_attrs.get('textContent', ''),
                    "title": attrs.get('title', '') or comp_attrs.get('title', ''),
                    "id": attrs.get('id', '') or comp_attrs.get('id', ''),
                    "name": attrs.get('name', '') or comp_attrs.get('name', ''),
                    "type": attrs.get('type', '') or comp_attrs.get('type', ''),
                    "role": attrs.get('role', '') or comp_attrs.get('role', ''),
                    "aria-label": attrs.get('aria-label', '') or comp_attrs.get('aria-label', ''),
                    "selectors": {**selectors, **comp_selectors}, # Merge basic and comprehensive selectors
                }
                
                # Clean up empty selectors
                element_data["selectors"] = {k: v for k, v in element_data["selectors"].items() if v}

                # Include recommended selector if available
                rec_sel = comp.get('recommended_selector', {}).get('selector', '')
                if rec_sel:
                    element_data["recommended_selector"] = rec_sel
                    
                step_info['dom_elements'].append(element_data)

            if isinstance(act, dict) and "extract_content" in act:
                goal = act.get("extract_content", {}).get("goal", "") if isinstance(act.get("extract_content", {}), dict) else ""
                cols = columns_from_results(results)
                if cols:
                    goal_text = f"{objective or ''} {goal}".lower()
                    filtered = []
                    for c in cols:
                        c_l = c.lower()
                        c_singular = c_l[:-1] if c_l.endswith("s") else c_l
                        if c_l in goal_text or c_singular in goal_text:
                            filtered.append(c)
                    required = filtered or cols
                    if "server" in goal_text and not any("server" in c.lower() for c in required):
                        required.append("Server")
                else:
                    required = infer_required_columns(f"{objective or ''} {goal}")
                if required:
                    step_info["required_columns_hint"] = required
                 
        # Final Pruning: Skip steps that contain no successful actions unless it's a Terminal step
        if not step_info['actions'] and "done" not in str(step_info['goal']).lower():
            continue
                
        simplified.append(step_info)
    
    result_data = simplified
    if objective:
        result_data = {"objective": objective, "steps": simplified}
        
    return json.dumps(result_data, indent=2)

def _extract_table_hints(cleaned_history_json: str):
    try:
        data = json.loads(cleaned_history_json)
    except Exception:
        return None, None

    steps = data.get("steps") if isinstance(data, dict) else data
    if not isinstance(steps, list):
        return None, None

    for step in steps:
        if not isinstance(step, dict):
            continue
        required = step.get("required_columns_hint")
        selector = None
        for el in step.get("dom_elements", []) or []:
            if not isinstance(el, dict):
                continue
            selectors = el.get("selectors") or {}
            for key in ["preferred", "css", "xpath", "aria", "text"]:
                val = selectors.get(key)
                if val:
                    selector = val
                    break
            if selector:
                break
        if selector or required:
            return selector, required
    return None, None

def _postprocess_generated_code(code: str, cleaned_history_json: str) -> str:
    selector, required = _extract_table_hints(cleaned_history_json)
    if selector:
        sel_literal = json.dumps(selector)
        code = re.sub(r'page\.wait_for_selector\(\s*["\']table["\']', f'page.wait_for_selector({sel_literal}', code)
        code = re.sub(r'resilient_extract_table\(\s*page\s*,\s*["\']table["\']', f'resilient_extract_table(page, {sel_literal}', code)
        code = re.sub(r'table_selector\s*=\s*["\']table["\']', f'table_selector = {sel_literal}', code)

    if required:
        req_literal = json.dumps(required)
        code = re.sub(r'(resilient_extract_table\(\s*page\s*,\s*[^,]+,\s*)(\[[^\]]*\])',
                      rf'\1{req_literal}', code)

    if "async def maybe_login" not in code:
        maybe_login_block = """

async def maybe_login(page):
    \"\"\"Attempt login if a login form is detected.\"\"\"
    username = await get_secret("OTCS_USERNAME")
    password = await get_secret("OTCS_PASSWORD")
    if not username or not password:
        return

    user_selectors = [
        "input[name='username']",
        "input#username",
        "input[name='user']",
        "input[type='text']"
    ]
    pass_selectors = [
        "input[name='password']",
        "input#password",
        "input[type='password']"
    ]
    submit_selectors = [
        "button:has-text('Sign in')",
        "button:has-text('Sign In')",
        "input[type='submit']",
        "button[type='submit']"
    ]

    pass_el = None
    for sel in pass_selectors:
        pass_el = await find_in_frames(page, sel)
        if pass_el:
            break
    if not pass_el:
        return

    await resilient_fill(page, user_selectors, username)
    await resilient_fill(page, pass_selectors, password)
    await resilient_click(page, submit_selectors, desc="login submit")
"""
        insert_before = "async def resilient_extract_table"
        if insert_before in code:
            code = code.replace(insert_before, maybe_login_block + "\n" + insert_before, 1)

    if "maybe_login(page)" not in code:
        code = re.sub(r'(await page\.goto\([^\n]+\)\n)', r'\1        await maybe_login(page)\n', code, count=1)

    return code

def generate_script(history_path: str, output_path: str = None, model_name: str = "qwen2.5:14b", provider: str = "ollama", objective: str = None, mandatory_history_path: str = None):
    """
    Generate script from history file.
    
    Args:
        history_path: Path to the process history JSON file
        output_path: Where to save the generated script
        model_name: LLM model to use
        provider: LLM provider (ollama, openai, etc.)
        objective: The task objective
        mandatory_history_path: Optional path to mandatory/login history to prepend
    """
    try:
        with open(history_path, 'r', encoding='utf-8') as f:
            history_data = json.load(f)
            
        run_id = Path(history_path).stem
        logger.info(f"Loaded history for {run_id}")
        
        # If mandatory history is provided, merge it with process history
        # This ensures non-mandatory process scripts can run standalone with login
        if mandatory_history_path and os.path.exists(mandatory_history_path):
            try:
                with open(mandatory_history_path, 'r', encoding='utf-8') as f:
                    mandatory_data = json.load(f)
                
                mandatory_steps = mandatory_data.get('history', [])
                process_steps = history_data.get('history', [])
                
                # Prepend mandatory steps to process steps
                history_data['history'] = mandatory_steps + process_steps
                
                # Include mandatory objective in the combined objective
                mandatory_objective = mandatory_data.get('task', '')
                if mandatory_objective:
                    objective = f"PREREQUISITES (Login/Session Setup):\n{mandatory_objective}\n\nMAIN TASK:\n{objective or history_data.get('task', '')}"
                
                logger.info(f"Merged {len(mandatory_steps)} mandatory steps with {len(process_steps)} process steps")
            except Exception as merge_err:
                logger.warning(f"Failed to merge mandatory history: {merge_err}. Proceeding with process-only history.")
        
        # Prepare content
        cleaned_history = clean_history_json(history_data, objective=objective)
        logger.info(f"Cleaned History Context (First 500 chars): {cleaned_history[:500]}...")
        
        # Save a debug copy of the cleaned history for inspection
        try:
            debug_path = history_path.replace('.json', '_cleaned.json')
            with open(debug_path, 'w', encoding='utf-8') as f:
                f.write(cleaned_history)
            logger.info(f"Debug cleaned history saved to: {debug_path}")
        except: pass

        # Prepare user prompt
        # Initialize LLM
        logger.info(f"Initializing LLM: {provider}/{model_name} (num_predict=4096)")
        llm = get_llm_model(
            provider=provider, 
            model_name=model_name, 
            temperature=0.1,
            num_predict=-1,   # Remove output limit to allow for very long scripts
            num_ctx=32000     # Maintain large context window
        )
        
        # Create Prompt
        messages = [
            SystemMessage(content=SYSTEM_PROMPT),
            HumanMessage(content=f"Task Objective: {objective or 'Complete the automation task'}\n\nGenerate a Playwright script for the following execution history (Run ID: {run_id}):\n\n{cleaned_history}")
        ]
        
        # Generate
        logger.info("Sending request to LLM (this may take a minute)...")
        response = llm.invoke(messages)
        code = response.content
        
        # Robust Extraction (Handle extra text, markdown blocks, etc.)
        import re
        code = response.content
        
        # Try to find the largest python code block
        code_blocks = re.findall(r'```python\n(.*?)\n```', code, re.DOTALL)
        if not code_blocks:
            code_blocks = re.findall(r'```\n(.*?)\n```', code, re.DOTALL)
            
        if code_blocks:
            # Use the longest block (highest probability of being the actual script)
            code = max(code_blocks, key=len).strip()
        else:
            # Fallback: strip markers manually but keep the whole content if no blocks found
            code = code.replace("```python", "").replace("```", "").strip()

        # Enforce required DOM scoping, columns, and login hook
        code = _postprocess_generated_code(code, cleaned_history)
        
        # Final safety check: if it STILL doesn't look like code, log it
        if "async def run" not in code:
            logger.warning("Generated content might not be a valid script! Checking for generic text...")
            # If there's an 'async filter' or similar, we might have a nested structure
            
        # Verify imports (Basic check)
        if "from playwright.async_api" not in code:
            logger.warning("Generated code might be missing imports!")

        # Determine output path
        if not output_path:
            output_path = str(Path(history_path).parent / f"{run_id}_LLM.py")
            
        with open(output_path, 'w', encoding='utf-8') as f:
            f.write(code)
            
        logger.info(f"✅ Script saved to: {output_path}")
        return output_path, code

    except Exception as e:
        logger.error(f"Failed to generate script: {e}")
        raise

if __name__ == "__main__":
    parser = argparse.ArgumentParser(description="Generate Playwright script via LLM")
    parser.add_argument("history_file", help="Path to agent history JSON file")
    parser.add_argument("--output", help="Output path for .py file")
    parser.add_argument("--provider", default="ollama", help="LLM Provider (default: ollama)")
    parser.add_argument("--model", default="qwen2.5:14b", help="Model name")
    
    args = parser.parse_args()
    
    generate_script(args.history_file, args.output, args.model, args.provider)
