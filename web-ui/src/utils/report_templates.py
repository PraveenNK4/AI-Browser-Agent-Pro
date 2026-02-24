"""
report_templates.py
===================
Generate Word (.docx) reports using the project's template files.

- generate_execution_report()  → "While Execution" template (agent runs via webui)
- generate_script_report()     → "While Running the script" template (generated script runs)

Templates are loaded from tmp/report_templates/ to inherit styling.
"""

import io
import os
import re
import base64
import logging
from pathlib import Path
from typing import Any, Dict, List, Optional
from datetime import datetime

from docx import Document
from docx.shared import Inches, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH

logger = logging.getLogger(__name__)

# Template paths — check src/utils/report_templates/ first (git-tracked), fall back to tmp/
_HERE = Path(__file__).resolve().parent
_WEBUI_ROOT = _HERE.parent.parent

# Primary location (in src, git-tracked)
_SRC_TEMPLATE_DIR = _HERE / "report_templates"
# Secondary location (in tmp, user's working copy)
_TMP_TEMPLATE_DIR = _WEBUI_ROOT / "tmp" / "report_templates"

TEMPLATE_DIR = _SRC_TEMPLATE_DIR if _SRC_TEMPLATE_DIR.exists() else _TMP_TEMPLATE_DIR
EXECUTION_TEMPLATE = TEMPLATE_DIR / "while_execution.docx"
SCRIPT_TEMPLATE = TEMPLATE_DIR / "while_running_script.docx"



def _new_doc_from_template(template_path: Path) -> Document:
    """Load template for styling, then clear all content to start fresh."""
    if template_path.exists():
        doc = Document(str(template_path))
        # Clear all existing content but preserve sectPr (section properties)
        from docx.oxml.ns import qn
        for element in list(doc.element.body):
            if element.tag != qn('w:sectPr'):
                doc.element.body.remove(element)
        return doc
    # Fallback: plain doc if template missing
    logger.warning(f"Template not found: {template_path}, using plain document")
    return Document()



def _add_bold_field(doc: Document, label: str, value: str):
    """Add a paragraph with a bold label followed by normal text."""
    p = doc.add_paragraph()
    p.add_run(label).bold = True
    p.add_run(str(value))
    return p


def _add_image_safe(doc: Document, image_source, width=Inches(6), caption: str = None):
    """Add an image from file path, bytes, or base64 string."""
    try:
        if isinstance(image_source, (str, Path)):
            path = Path(image_source)
            if path.exists() and path.stat().st_size > 0:
                doc.add_picture(str(path), width=width)
            else:
                return
        elif isinstance(image_source, bytes):
            doc.add_picture(io.BytesIO(image_source), width=width)
        elif isinstance(image_source, io.BytesIO):
            doc.add_picture(image_source, width=width)
        else:
            return
        if caption:
            cap = doc.add_paragraph(caption)
            cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
    except Exception as e:
        logger.warning(f"Failed to add image: {e}")


# ---------------------------------------------------------------------------
# 1. Agent Execution Report ("While Execution" template)
# ---------------------------------------------------------------------------

def generate_execution_report(
    history,
    task: str,
    task_id: str = None,
    output_path: str = None,
    step_failures: dict = None,
) -> Optional[str]:
    """Generate DOCX report from agent execution history.

    Args:
        history:        AgentHistoryList from browser-use agent.
        task:           The task description string.
        task_id:        Unique task identifier.
        output_path:    Full path for the output .docx file.
        step_failures:  Dict mapping step index to failure count.

    Returns:
        Path to generated .docx or None on failure.
    """
    try:
        doc = _new_doc_from_template(EXECUTION_TEMPLATE)

        # Title
        doc.add_heading("AI Browser Agent Execution Report", 0)

        # --- Task Summary ---
        doc.add_heading("Task Summary", level=1)
        _add_bold_field(doc, "Task: ", task)

        duration = history.total_duration_seconds()
        tokens = history.total_input_tokens()
        final_result = history.final_result()

        summary_table = doc.add_table(rows=5, cols=2)
        summary_table.style = "Table Grid"
        rows_data = [
            ("Duration", f"{duration:.2f}s"),
            ("Tokens", str(tokens)),
            ("Steps", str(len(history.history))),
            ("Status", "Completed" if final_result else "Failed"),
            ("Task ID", task_id or "Unknown"),
        ]
        for i, (label, value) in enumerate(rows_data):
            summary_table.rows[i].cells[0].text = label
            summary_table.rows[i].cells[1].text = value

        # --- Final Result ---
        doc.add_heading("Final Result", level=1)
        if final_result:
            _add_bold_field(doc, "Outcome: ", final_result)

        _step_failures = step_failures or {}
        total_failures = sum(_step_failures.values())
        if total_failures > 0:
            _add_bold_field(doc, "Status: ", f"Completed with {total_failures} failures")
        else:
            _add_bold_field(doc, "Status: ", "All steps successful")

        # --- Step-by-Step Execution ---
        doc.add_heading("Step-by-Step Execution", level=1)
        for i, step in enumerate(history.history, 1):
            step_status = "PASS" if _step_failures.get(i, 0) == 0 else f"FAIL ({_step_failures.get(i, 0)}x)"
            doc.add_heading(f"Step {i} - {step_status}", level=2)

            # Goal
            goal_text = "Execute action"
            if hasattr(step, "model_output") and step.model_output:
                if hasattr(step.model_output, "current_state") and step.model_output.current_state:
                    if hasattr(step.model_output.current_state, "next_goal") and step.model_output.current_state.next_goal:
                        goal_text = str(step.model_output.current_state.next_goal)
            _add_bold_field(doc, "Goal: ", goal_text)

            # Action
            action_desc = "No action"
            if hasattr(step, "action"):
                action = getattr(step, "action", None)
                if action:
                    action_desc = str(action[0]) if isinstance(action, (list, tuple)) else str(action)
            _add_bold_field(doc, "Action: ", action_desc)

            # Screenshot (from step state base64)
            if hasattr(step, "state") and hasattr(step.state, "screenshot") and step.state.screenshot:
                try:
                    screenshot_data = step.state.screenshot
                    if screenshot_data.startswith("data:image"):
                        screenshot_data = screenshot_data.split(",")[1]
                    from PIL import Image
                    image_data = base64.b64decode(screenshot_data)
                    image = Image.open(io.BytesIO(image_data))
                    if image.width > 900:
                        ratio = 900 / image.width
                        image = image.resize(
                            (int(image.width * ratio), int(image.height * ratio)),
                            Image.Resampling.LANCZOS,
                        )
                    if image.mode in ("RGBA", "LA", "P"):
                        image = image.convert("RGB")
                    img_buffer = io.BytesIO()
                    image.save(img_buffer, format="JPEG", quality=95)
                    img_buffer.seek(0)
                    _add_image_safe(doc, img_buffer, caption=f"Step {i} screenshot")
                except Exception as e:
                    logger.warning(f"Failed to add screenshot for step {i}: {e}")

            # Result
            result_text = "No result"
            if hasattr(step, "result") and step.result:
                if isinstance(step.result, list) and len(step.result) > 0:
                    first_result = step.result[0]
                    extracted = getattr(first_result, "extracted_content", None)
                    result_text = str(extracted) if extracted else str(first_result)
                else:
                    result_text = str(step.result)
            _add_bold_field(doc, "Result: ", result_text)
            doc.add_paragraph()

        # --- Errors ---
        errors = history.errors()
        if errors and any(errors):
            doc.add_heading("Errors", level=1)
            for error in errors:
                if error:
                    doc.add_paragraph(f"Error: {str(error)}")

        # Save
        if not output_path:
            output_path = str(Path.cwd() / f"execution_report_{datetime.now().strftime('%Y%m%d_%H%M%S')}.docx")
        doc.save(output_path)
        logger.info(f"Execution report saved: {output_path}")
        return output_path

    except Exception as e:
        logger.error(f"Execution report generation failed: {e}", exc_info=True)
        return None


# ---------------------------------------------------------------------------
# 2. Script Execution Report ("While Running the script" template)
# ---------------------------------------------------------------------------

def generate_script_report(
    script_name: str,
    steps: List[Dict[str, Any]],
    screenshots_dir: Optional[str] = None,
    output_path: Optional[str] = None,
    status: str = "Execution Complete",
    captured_outputs: Optional[str] = None,
) -> Optional[str]:
    """Generate DOCX report from a generated Playwright script execution.

    Args:
        script_name:     Name/ID of the script.
        steps:           List of step dicts with keys: action, output, screenshot_path (optional).
        screenshots_dir: Directory containing numbered screenshots from take_full_screenshot.
        output_path:     Full path for output .docx file.
        status:          Overall execution status string.
        captured_outputs: Combined text output to include in "Captured Outputs" section.

    Returns:
        Path to generated .docx or None on failure.
    """
    try:
        doc = _new_doc_from_template(SCRIPT_TEMPLATE)

        # Title
        doc.add_heading("Playwright Automation Report", 0)

        # --- Execution Summary ---
        doc.add_heading("Execution Summary", level=1)
        _add_bold_field(doc, "Script: ", script_name)
        _add_bold_field(doc, "Status: ", status)
        _add_bold_field(doc, "Total Steps: ", str(len(steps)))
        _add_bold_field(doc, "Generated: ", datetime.now().strftime("%Y-%m-%d %H:%M:%S"))

        # --- Captured Outputs ---
        if captured_outputs:
            doc.add_heading("Captured Outputs", level=1)
            doc.add_paragraph(captured_outputs)

        # --- Summary Table ---
        if steps:
            summary_table = doc.add_table(rows=len(steps) + 1, cols=3)
            summary_table.style = "Table Grid"
            summary_table.rows[0].cells[0].text = "Step"
            summary_table.rows[0].cells[1].text = "Action"
            summary_table.rows[0].cells[2].text = "Output"
            for i, step in enumerate(steps, 1):
                summary_table.rows[i].cells[0].text = str(i)
                summary_table.rows[i].cells[1].text = step.get("action", "N/A")
                output_text = step.get("output", "N/A")
                summary_table.rows[i].cells[2].text = str(output_text)[:200] if output_text else "N/A"

        # --- Automation Steps ---
        doc.add_heading("Automation Steps", level=1)
        for i, step in enumerate(steps, 1):
            action_name = step.get("action", "unknown")
            doc.add_heading(f"Step {i}: {action_name}", level=2)

            # Detail table
            detail_table = doc.add_table(rows=2, cols=2)
            detail_table.style = "Table Grid"
            detail_table.rows[0].cells[0].text = "Action"
            detail_table.rows[0].cells[1].text = action_name
            detail_table.rows[1].cells[0].text = "Captured Output"
            output_val = step.get("output", "N/A")
            detail_table.rows[1].cells[1].text = str(output_val)[:500] if output_val else "N/A"

            # Screenshot from step data or from screenshots directory
            screenshot_path = step.get("screenshot_path")
            if screenshot_path:
                _add_image_safe(doc, screenshot_path, caption=f"Step {i}: {action_name}")
            elif screenshots_dir:
                ss_dir = Path(screenshots_dir)
                if ss_dir.exists():
                    # Match by step number prefix
                    matches = sorted(ss_dir.glob(f"{i:02d}_*.png"))
                    if not matches:
                        matches = sorted(ss_dir.glob(f"{i-1}.png"))
                    if matches:
                        _add_image_safe(doc, matches[0], caption=f"Step {i}: {action_name}")

            doc.add_paragraph()

        # Save
        if not output_path:
            output_path = str(Path.cwd() / f"script_report_{datetime.now().strftime('%Y%m%d_%H%M%S')}.docx")
        doc.save(output_path)
        logger.info(f"Script report saved: {output_path}")
        return output_path

    except Exception as e:
        logger.error(f"Script report generation failed: {e}", exc_info=True)
        return None
