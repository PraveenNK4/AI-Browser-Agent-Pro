"""
script_helpers.py
=================
Shared runtime helpers imported by every generated Playwright script.

Generated scripts should NOT redefine these functions — they import them:

    from src.utils.script_helpers import (
        get_secret, find_in_frames,
        resilient_fill, resilient_click, click_and_upload_file,
        maybe_login, resilient_extract_table, generate_report,
    )

All tuneable values (timeouts, selectors, vault prefix) come from
src.utils.config so they can be changed via environment variables.
"""

import io
import json
import os
import re
import sys
import asyncio
from pathlib import Path
from typing import Union, Optional, Type, Callable, Dict, Any, Awaitable, TypeVar

# Ensure src is importable regardless of working directory
_here = Path(__file__).resolve().parent
for _candidate in (_here.parent.parent, _here.parent.parent.parent):
    if (_candidate / "src").exists() and str(_candidate) not in sys.path:
        sys.path.insert(0, str(_candidate))

from src.utils.config import (
    VAULT_CREDENTIAL_PREFIX,
    LOGIN_USER_SELECTORS,
    LOGIN_PASS_SELECTORS,
    LOGIN_SUBMIT_SELECTORS,
    TIMEOUT_ELEMENT_VISIBLE_MS,
    TIMEOUT_FILL_MS,
    TIMEOUT_CLICK_MS,
    TIMEOUT_FILE_CHOOSER_MS,
    TIMEOUT_UPLOAD_CLICK_MS,
    TIMEOUT_UPLOAD_FALLBACK_MS,
    TIMEOUT_TABLE_WAIT_MS,
)
from src.utils.vault import vault


# ---------------------------------------------------------------------------
# Credential helpers
# ---------------------------------------------------------------------------

async def get_secret(key: str, vault_prefix: str = None):
    """Look up a credential from the vault by composite key.

    The key is either a full composite like "PREFIX_USERNAME" (vault key = "PREFIX",
    field = "USERNAME") or a bare field name like "USERNAME" when vault_prefix is
    supplied separately.

    vault_prefix overrides the first component derived from key, allowing per-script
    vault keys without changing the global VAULT_CREDENTIAL_PREFIX config.

    Examples:
        get_secret("PREFIX_USERNAME")            → vault.get_credentials("PREFIX")["username"]
        get_secret("USERNAME", "PREFIX")         → vault.get_credentials("PREFIX")["username"]
        get_secret("APP_USERNAME")               → vault.get_credentials("APP")["username"]
    """
    if vault_prefix:
        # Prefix supplied explicitly — key is just the field name (or full composite)
        v_key = vault_prefix
        field_key = key
    else:
        parts = key.split("_")
        v_key = parts[0] if len(parts) > 1 else key
        field_key = key

    creds = vault.get_credentials(v_key) or vault.get_credentials(v_key.lower())
    if not creds:
        print(f"[-] get_secret: no vault entry for prefix '{v_key}' (key='{key}')")
        return None
    if any(k in field_key.upper() for k in ("USERNAME", "USER")):
        return creds.get("username") or creds.get("Username") or creds.get("user") or creds.get("User")
    if any(k in field_key.upper() for k in ("PASSWORD", "PWD")):
        return creds.get("password") or creds.get("Password") or creds.get("pwd") or creds.get("Pwd")
    return None


# ---------------------------------------------------------------------------
# DOM interaction helpers
# ---------------------------------------------------------------------------

async def find_in_frames(page, selector: str, timeout: int = None):
    """Search for *selector* in the main page and all its iframes.

    Returns the first visible element found, or None.
    """
    visible_timeout = timeout or TIMEOUT_ELEMENT_VISIBLE_MS
    try:
        el = page.locator(selector).first
        if await el.count() > 0 and await el.is_visible(timeout=visible_timeout):
            return el
    except Exception:
        pass
    for frame in page.frames:
        try:
            el = frame.locator(selector).first
            if await el.count() > 0 and await el.is_visible(timeout=visible_timeout):
                return el
        except Exception:
            continue
    return None


async def resilient_fill(page, selector, value: str) -> bool:
    """Fill *value* into the first matching element from *selector* (str or list)."""
    if value is None:
        return False
    selectors = selector if isinstance(selector, list) else [selector]
    for sel in selectors:
        try:
            print(f"[*] Filling {sel}...")
            el = await find_in_frames(page, sel)
            if not el:
                await page.wait_for_selector(sel, timeout=TIMEOUT_FILL_MS)
                el = page.locator(sel).first
            await el.fill(str(value))
            await el.dispatch_event("input")
            await el.dispatch_event("change")
            return True
        except Exception as e:
            print(f"[-] Fill failed on {sel}: {e}")
            continue
    print(f"[-] Fill failed for all selectors: {selectors}")
    return False


async def resilient_click(page, selectors, desc: str = "element") -> bool:
    """Click the first visible match from *selectors* (str or list)."""
    if isinstance(selectors, str):
        selectors = [selectors]
    for sel in selectors:
        try:
            print(f"[*] Clicking {desc} via {sel}...")
            el = await find_in_frames(page, sel)
            if el:
                await el.click(timeout=TIMEOUT_CLICK_MS)
                return True
        except Exception as e:
            print(f"[-] Click failed on {sel}: {e}")
            continue
    print(f"[-] Click failed for all selectors ({desc}): {selectors}")
    return False


async def click_and_upload_file(page, target_name: str, path: Union[str, list[str]]) -> bool:
    """Robustly trigger a file-chooser and upload one or more files.

    Supports:
        - Single file path (str)
        - List of file paths (list[str])
        - Directory path (automatically expands to all files in directory)

    Tries three strategies in order:
      1. Click a button/link whose text matches *target_name* exactly.
      2. Set files directly on a hidden <input type="file">.
      3. Click any element with that text and intercept the chooser.
    """
    print(f"[*] Starting robust upload for: {target_name} (path={path})")
    
    # --- 1. Expand paths (handle strings, lists, and directories) ---
    target_paths = []
    raw_paths = [path] if isinstance(path, str) else path
    
    for p in raw_paths:
        p_norm = os.path.normpath(p)
        if os.path.isdir(p_norm):
            print(f"[*] Expanding directory: {p_norm}")
            for root, _, files in os.walk(p_norm):
                for f in files:
                    if not f.startswith('.'):
                        target_paths.append(os.path.join(root, f))
        else:
            if os.path.exists(p_norm):
                target_paths.append(p_norm)
            else:
                print(f"[-] File not found: {p_norm}")
    
    if not target_paths:
        print(f"[-] No valid file paths found for upload.")
        return False

    # Use a single string if only one file (some apps prefer this), else list
    payload = target_paths[0] if len(target_paths) == 1 else target_paths
    print(f"[*] Final payload ({len(target_paths)} files): {payload}")

    # Strategy 1: native chooser via strict label match
    try:
        target = page.locator("a, button, [role='button']").filter(
            has_text=re.compile(f"^{re.escape(target_name)}$", re.I)
        ).first
        async with page.expect_file_chooser(timeout=TIMEOUT_FILE_CHOOSER_MS) as fc_info:
            await target.click(timeout=TIMEOUT_UPLOAD_CLICK_MS)
        (await fc_info.value).set_files(payload)
        return True
    except Exception:
        pass

    # Strategy 2: hidden file input
    try:
        await page.locator('input[type="file"]').first.set_input_files(payload)
        return True
    except Exception:
        pass

    # Strategy 3: global text search
    try:
        target = page.get_by_text(target_name, exact=True).first
        async with page.expect_file_chooser(timeout=TIMEOUT_UPLOAD_FALLBACK_MS) as fc_info:
            await target.click()
        (await fc_info.value).set_files(payload)
        return True
    except Exception:
        pass

    print(f"[-] All upload strategies failed for: {target_name}")
    return False


# ---------------------------------------------------------------------------
# Login helper
# ---------------------------------------------------------------------------

async def maybe_login(page, vault_prefix: str = None) -> None:
    """Detect and fill a login form if one is present.

    vault_prefix selects which vault entry to use for credentials.
    When omitted, falls back to VAULT_CREDENTIAL_PREFIX from config.
    This allows each generated script to specify its own vault key
    without changing the global config.

    Example:
        await maybe_login(page, vault_prefix="PREFIX")   # uses prefix credentials
        await maybe_login(page)                         # uses VAULT_CREDENTIAL_PREFIX
    """
    prefix = vault_prefix or VAULT_CREDENTIAL_PREFIX
    username = await get_secret(f"{prefix}_USERNAME", vault_prefix=vault_prefix)
    password = await get_secret(f"{prefix}_PASSWORD", vault_prefix=vault_prefix)
    if not username or not password:
        print(f"[-] maybe_login: no credentials found for vault prefix '{prefix}'. "
              "Check vault key matches the prefix used in the task (e.g. @vault.prefix).")
        return

    # Only proceed if a password field is actually visible
    pass_el = None
    for sel in LOGIN_PASS_SELECTORS:
        pass_el = await find_in_frames(page, sel)
        if pass_el:
            break
    if not pass_el:
        return

    print(f"[+] maybe_login: logging in with vault prefix '{prefix}' (user={username[:3]}***)")
    
    # Ensure the fields are visible before filling
    await page.wait_for_selector(LOGIN_USER_SELECTORS[0], timeout=TIMEOUT_FILL_MS)
    
    await resilient_fill(page, LOGIN_USER_SELECTORS, username)
    await resilient_fill(page, LOGIN_PASS_SELECTORS, password)
    await resilient_click(page, LOGIN_SUBMIT_SELECTORS, desc="login submit")

    # Wait for navigation after submit
    try:
        await page.wait_for_load_state("networkidle", timeout=TIMEOUT_TABLE_WAIT_MS)
    except Exception:
        pass


# ---------------------------------------------------------------------------
# Table extraction helper
# ---------------------------------------------------------------------------

async def resilient_extract_table(
    page,
    table_selector: str,
    header_keywords: list,
    skip_boilerplate: list = None,
) -> list:
    """Extract structured rows from the best-matching table on the page.

    IMPLEMENTATION STRATEGY — single JS round-trip:
    =================================================
    All previous approaches made one Playwright call per row/cell (100+ async
    round-trips for a 30-row table).  This version does everything in ONE
    page.evaluate() call:

      1. Collect candidate tables:
         - If table_selector is a CSS selector that matches a specific element
           (e.g. "#admin-servers table"), use that element + its descendants.
         - If selector is bare "table", scan ALL tables on the page via
           document.querySelectorAll('table') — finds the data table even when
           it's a sibling of layout tables, not a descendant.

      2. Score each candidate by header quality:
         - Find the first row where ALL requested keywords land on distinct
           columns (no two keywords share a column — that would be a blob row).
         - Score = keywords_matched × (th_bonus=10 or 1) + 50 if all matched.
         - Best-scoring table wins.

      3. Extract data rows from the winning table:
         - Read only :scope > td / :scope > th (direct children only) to avoid
           inheriting text from nested tables.
         - Images: use img[title] text.
         - Skip the header row, blank rows, and rows that only repeat header text.

      4. Identity column:
         - Auto-detect a column whose header contains server/name/host/node/agent.
         - If the identity keyword IS in header_keywords (e.g. "Name"), promote
           that column — it was previously skipped because it was already in
           col_map.values().
         - Write both row["Server"] and row["Name"] so the report chain
           `r.get('Server') or r.get('Name')` always finds a value.

    Args:
        page:             Playwright page object.
        table_selector:   CSS selector for the table or its container.
        header_keywords:  Column names to extract (matched case-insensitively).
        skip_boilerplate: Extra row text patterns to skip (pass [] to disable).

    Returns:
        List of dicts keyed by header_keywords plus "Server" and "Name".
    """
    # ── Build the JS extraction script ─────────────────────────────────────
    id_kws_js   = json.dumps(["server", "name", "host", "node", "entity", "admin", "agent"])
    keywords_js = json.dumps([k.lower() for k in header_keywords])
    skip_js     = json.dumps([s.lower() for s in (skip_boilerplate or [])])
    orig_kws_js = json.dumps(header_keywords)   # original casing for output keys

    JS = f"""
    (function() {{
        const keywords    = {keywords_js};
        const origKws     = {orig_kws_js};
        const idKws       = {id_kws_js};
        const skipPat     = {skip_js};
        const selector    = {json.dumps(table_selector)};

        // ── 1. Collect candidate tables ───────────────────────────────────
        let candidates = [];
        if (selector === 'table') {{
            candidates = Array.from(document.querySelectorAll('table'));
        }} else {{
            const root = document.querySelector(selector);
            if (root) {{
                const tables = Array.from(root.querySelectorAll('table'));
                candidates = (root.tagName === 'TABLE') ? [root, ...tables] : tables;
                if (candidates.length === 0) candidates = [root];
            }}
            
            // Safety: if selector didn't find anything, or finds empty results later,
            // we'll eventually allow a full-page retry. 
            // For now, if root is null, start with all tables immediately.
            if (!root) {{
                candidates = Array.from(document.querySelectorAll('table'));
            }}
        }}

        // ── Helper: direct-child cells only ──────────────────────────────
        function cells(tr) {{
            return Array.from(tr.querySelectorAll(':scope > td, :scope > th')).map(td => {{
                const img = td.querySelector('img[title]');
                return img ? img.getAttribute('title') : (td.innerText || '').trim();
            }});
        }}

        // ── 2. Score each candidate and pick the best ─────────────────────
        let bestScore = -1, bestTable = null, bestHdrIdx = -1, bestColMap = null;

        for (const tbl of candidates) {{
            const trs = Array.from(tbl.querySelectorAll(':scope > tbody > tr, :scope > tr'))
                         .concat(Array.from(tbl.querySelectorAll(':scope > thead > tr')));
            // deduplicate preserving order
            const seen = new Set(); const rows = [];
            for (const r of trs) {{ if (!seen.has(r)) {{ seen.add(r); rows.push(r); }} }}

            for (let i = 0; i < Math.min(rows.length, 25); i++) {{
                const txts = cells(rows[i]);
                if (!txts.length) continue;

                // Map each keyword to the first column whose text contains it
                const cm = {{}};
                for (const kw of keywords) {{
                    const idx = txts.findIndex(t => t.toLowerCase().includes(kw));
                    cm[kw] = idx;  // -1 if not found
                }}
                const resolved = Object.values(cm).filter(v => v >= 0);
                // Reject blob rows (two keywords on same col)
                if (new Set(resolved).size !== resolved.length) continue;

                const hasTh = rows[i].querySelector(':scope > th') !== null;
                let score = resolved.length * (hasTh ? 10 : 1);
                
                // Bonus for perfect match, penalty for missing keys
                if (resolved.length === keywords.length) {{
                    score += 500; // massive bonus so perfect matches always win
                }} else {{
                    score -= 50;  // penalty for partial matches
                }}

                if (score > bestScore) {{
                    bestScore = score; bestTable = tbl;
                    bestHdrIdx = i; bestColMap = cm;
                }}
            }}
        }}

        if (!bestTable || bestScore < 0) {{
            return {{ error: 'no matching header found', candidates: candidates.length }};
        }}

        // ── 3. Re-collect rows from winning table (same order as above) ───
        const allTrs = Array.from(bestTable.querySelectorAll(':scope > tbody > tr, :scope > tr'))
                        .concat(Array.from(bestTable.querySelectorAll(':scope > thead > tr')));
        const seenR = new Set(); const winRows = [];
        for (const r of allTrs) {{ if (!seenR.has(r)) {{ seenR.add(r); winRows.push(r); }} }}

        // ── 4. Determine identity column ──────────────────────────────────
        const hdrCells = cells(winRows[bestHdrIdx]);
        const usedCols = new Set(Object.values(bestColMap).filter(v => v >= 0));
        let identIdx = -1, identName = '';

        // First try: find an identity column not already claimed
        for (let j = 0; j < hdrCells.length; j++) {{
            const t = hdrCells[j].toLowerCase();
            if (idKws.some(k => t.includes(k)) && !usedCols.has(j)) {{
                identIdx = j; identName = hdrCells[j]; break;
            }}
        }}
        // Fallback: promote a requested keyword that IS an identity key
        if (identIdx < 0) {{
            for (const kw of keywords) {{
                if (idKws.some(k => kw.includes(k)) && bestColMap[kw] >= 0) {{
                    identIdx = bestColMap[kw]; identName = kw; break;
                }}
            }}
        }}

        const hdrLower = new Set(hdrCells.map(t => t.toLowerCase()).filter(Boolean));
        const maxCol = Math.max(...Object.values(bestColMap).filter(v => v >= 0),
                                  identIdx >= 0 ? identIdx : 0);

        // ── 5. Extract data rows ──────────────────────────────────────────
        const results = [];
        for (let ri = 0; ri < winRows.length; ri++) {{
            if (ri === bestHdrIdx) continue;
            const cv = cells(winRows[ri]);
            if (cv.length <= maxCol) continue;  // too few cells

            // Skip single-cell or all-same-value rows (dividers, section titles)
            const nonempty = cv.filter(Boolean);
            if (cv.length === 1 || (new Set(nonempty).size === 1 && nonempty.length > 1)) continue;

            // Skip rows that just repeat header labels
            const rowLower = new Set(cv.map(t => t.toLowerCase()).filter(Boolean));
            if (rowLower.size && [...rowLower].every(t => hdrLower.has(t))) continue;

            // Skip boilerplate
            const allText = cv.join(' ').toLowerCase();
            if (skipPat.some(p => allText.includes(p))) continue;

            const row = {{}};

            // Identity
            if (identIdx >= 0 && identIdx < cv.length) {{
                const idVal = cv[identIdx];
                if (idVal && !hdrLower.has(idVal.toLowerCase())) {{
                    row['Server'] = idVal;
                    row['Name']   = idVal;
                    if (identName) row[identName] = idVal;
                }}
            }}

            // Requested columns
            for (let ki = 0; ki < keywords.length; ki++) {{
                const kw = keywords[ki];
                const origKw = origKws[ki];
                const cidx = bestColMap[kw];
                row[origKw] = (cidx >= 0 && cidx < cv.length && cv[cidx] && cv[cidx] !== 'null')
                               ? cv[cidx] : '';
            }}

            if (Object.values(row).some(Boolean)) results.push(row);
        }}

        return {{
            rows: results,
            debug: {{
                tableCandidates: candidates.length,
                winningSelector: bestTable.id ? '#' + bestTable.id : bestTable.className || '(table)',
                headerRow: hdrCells,
                colMap: bestColMap,
                identityCol: identName + ' @ col ' + identIdx,
                score: bestScore,
            }}
        }};
    }})()
    """

    # ── Run the JS ─────────────────────────────────────────────────────────
    # Try main frame first; fall back to each iframe if the element lives there.
    result = None
    frames_tried = [page]
    for frame in page.frames:
        if frame != page.main_frame:
            frames_tried.append(frame)

    for frame in frames_tried:
        try:
            result = await frame.evaluate(JS)
        except Exception as e:
            logger.debug(f"[extract] frame evaluate error: {e}")
            continue
        if result and not result.get("error"):
            break
        if result and result.get("error") == f"selector not found: {table_selector}":
            continue  # try next frame
        break

    if not result:
        raise RuntimeError(f"[extract] JS evaluate returned nothing for selector {table_selector!r}")
    if result.get("error"):
        raise RuntimeError(f"[extract] {result['error']} (selector={table_selector!r})")

    dbg = result.get("debug", {})
    print(f"[*] Extracted {len(result['rows'])} rows from {table_selector!r} "
          f"[table:{dbg.get('winningSelector','?')} "
          f"score:{dbg.get('score','?')} "
          f"identity:{dbg.get('identityCol','?')}]")

    if skip_boilerplate is None:
        # Filter rows with no meaningful values
        return [r for r in result["rows"] if any(v for v in r.values())]
    return result["rows"]


async def resilient_menu_click(page, row_text: str) -> bool:
    """Ultra-robust click on a row's function menu (hotspot/dropdown).
    
    Finds the element containing row_text, locates its parent row (tr or similar),
    and triggers mousedown+click+mouseup on any hotspot/dropdown element.
    """
    found = await page.evaluate('''async (text) => {
        const candidates = [...document.querySelectorAll("a, span, div, td")].filter(el => {
            const t = el.textContent.trim();
            return t === text || t.includes(text);
        });
        
        for (const el of candidates) {
            const row = el.closest("tr, [role='row'], .bin_container_row");
            if (row) {
                const sel = [
                    "a[class*='hotspot']", "a[title*='Action']", "a[id*='hotspot']",
                    ".ot-menu-hotspot", ".dropdown-toggle", "[aria-haspopup='true']",
                    "button[class*='menu']", ".otcs-menu-launcher"
                ].join(', ');
                const h = row.querySelector(sel);
                
                if (h) {
                    h.scrollIntoView({ block: 'center' });
                    h.focus();
                    h.dispatchEvent(new MouseEvent('mouseover', { bubbles: true }));
                    h.dispatchEvent(new MouseEvent('mousedown', { bubbles: true }));
                    h.dispatchEvent(new MouseEvent('mouseup', { bubbles: true }));
                    h.click();
                    return true;
                }
            }
        }
        return false;
    }''', row_text)
    
    if not found:
        print(f"[!] resilient_menu_click: No hotspot found for {row_text!r}")
    else:
        await page.wait_for_timeout(1500)
    return found


# ---------------------------------------------------------------------------
# Certificate / Security helper
# ---------------------------------------------------------------------------

async def check_certificate(page) -> dict:
    """Check the SSL certificate of the current page without clicking browser UI.

    Two complementary approaches run in parallel:

    1. CDP (Chrome DevTools Protocol) — always available via Playwright, gives
       the browser's own security assessment (same data as the lock-icon popup).
    2. ssl module — direct TLS handshake to get the full certificate chain,
       valid-from/to dates, issuer, SANs, cipher suite.

    The results are rendered as a pixel-accurate screenshot of a panel that
    looks exactly like the Chrome/Edge 'Connection is secure' popup, with no
    need to click browser-chrome buttons or hardcode any coordinates.

    Returns
    -------
    dict with keys:
        screenshot  : bytes  — PNG of the rendered security panel
        cert_info   : dict   — parsed certificate fields
        is_valid    : bool   — True if cert is present and not expired
        is_secure   : bool   — True if browser considers connection secure
        hostname    : str
    """
    import ssl as _ssl
    import socket
    import datetime
    from urllib.parse import urlparse

    url   = page.url
    parsed  = urlparse(url)
    hostname = parsed.hostname or ""
    port     = parsed.port or (443 if parsed.scheme == "https" else 80)
    is_https = parsed.scheme == "https"

    # ── 1. CDP security state ────────────────────────────────────────────────
    cdp_state = {}
    is_secure = False
    cdp_failed = False
    try:
        cdp = await page.context.new_cdp_session(page)
        await cdp.send("Security.enable")
        result = await cdp.send("Security.getSecurityState")
        cdp_state = result.get("visibleSecurityState", result)
        sec = cdp_state.get("securityState", "")
        is_secure = sec == "secure"
        await cdp.detach()
    except Exception as e:
        cdp_failed = True
        print(f"[cert] CDP security check failed: {e}")

    # ── 2. Direct TLS handshake for full cert details ────────────────────────
    cert_info: dict = {}
    is_valid = False
    if is_https and hostname:
        try:
            import hashlib
            ctx = _ssl.create_default_context()
            with socket.create_connection((hostname, port), timeout=10) as raw_sock:
                with ctx.wrap_socket(raw_sock, server_hostname=hostname) as tls_sock:
                    cert    = tls_sock.getpeercert()
                    cert_der = tls_sock.getpeercert(binary_form=True)
                    cipher  = tls_sock.cipher()           # (name, protocol, bits)
                    version = tls_sock.version()           # 'TLSv1.3' etc.

            subject = dict(x[0] for x in cert.get("subject",  []))
            issuer  = dict(x[0] for x in cert.get("issuer",   []))

            def _parse_dt(s):
                for fmt in ("%b %d %H:%M:%S %Y %Z", "%b  %d %H:%M:%S %Y %Z"):
                    try:
                        return datetime.datetime.strptime(s, fmt)
                    except ValueError:
                        continue
                return datetime.datetime.utcnow()

            not_after  = _parse_dt(cert.get("notAfter",  ""))
            not_before = _parse_dt(cert.get("notBefore", ""))
            now         = datetime.datetime.utcnow()
            days_left   = (not_after - now).days
            is_valid    = not_before <= now <= not_after

            # If CDP failed but TLS handshake succeeded and cert is valid,
            # treat the connection as secure (HTTPS + valid cert = secure)
            if cdp_failed and is_valid and is_https:
                is_secure = True

            san_list = [v for _, v in cert.get("subjectAltName", [])
                        if not v.startswith("*")][:6]

            # SHA-256 fingerprint of the DER-encoded certificate
            fp_sha256 = hashlib.sha256(cert_der).hexdigest() if cert_der else "—"
            
            # SHA-256 fingerprint of the Public Key (SPKI)
            try:
                from cryptography import x509
                from cryptography.hazmat.primitives import serialization
                cert_obj = x509.load_der_x509_certificate(cert_der)
                pubkey = cert_obj.public_key()
                pubkey_der = pubkey.public_bytes(
                    encoding=serialization.Encoding.DER,
                    format=serialization.PublicFormat.SubjectPublicKeyInfo
                )
                pubkey_fp = hashlib.sha256(pubkey_der).hexdigest()
            except Exception as e:
                print(f"[cert] Could not compute public key fingerprint: {e}")
                pubkey_fp = "—"

            # Chrome-style date format
            import os
            fmt = "%A, %B %d, %Y at %#I:%M:%S %p" if os.name == 'nt' else "%A, %B %-d, %Y at %-I:%M:%S %p"
            
            # Convert naive datetime (which represents UTC) to local timezone to match Chrome's behavior
            import datetime as dt_mod
            utc_not_before = not_before.replace(tzinfo=dt_mod.timezone.utc)
            utc_not_after = not_after.replace(tzinfo=dt_mod.timezone.utc)
            local_valid_from = utc_not_before.astimezone()
            local_valid_to = utc_not_after.astimezone()

            valid_from_fmt = local_valid_from.strftime(fmt)
            valid_to_fmt   = local_valid_to.strftime(fmt)

            cert_info = {
                "hostname":      hostname,
                "subject_cn":    subject.get("commonName", hostname),
                "subject_org":   subject.get("organizationName", ""),
                "issuer_cn":     issuer.get("commonName",  "Unknown CA"),
                "issuer_org":    issuer.get("organizationName", ""),
                "valid_from":    valid_from_fmt,
                "valid_to":      valid_to_fmt,
                "days_remaining": days_left,
                "is_expired":    days_left < 0,
                "cipher":        cipher[0] if cipher else "—",
                "protocol":      version or (cipher[1] if cipher else "—"),
                "key_bits":      cipher[2] if cipher else "—",
                "san":           san_list,
                "serial":        str(cert.get("serialNumber", "—")),
                "fingerprint_sha256": fp_sha256,
                "pubkey_fingerprint": pubkey_fp,
            }
            print(f"[cert] [OK] {hostname} - valid for {days_left} more days")
        except _ssl.SSLCertVerificationError as e:
            print(f"[cert] [WARN] Certificate verification FAILED: {e}")
            cert_info = {"hostname": hostname, "error": str(e),
                         "days_remaining": 0, "is_expired": True}
        except Exception as e:
            print(f"[cert] TLS handshake failed: {e}")
            cert_info = {"hostname": hostname, "error": str(e)}

    # ── 3. Gather all live page permissions + cookies via browser APIs ────────
    # Use navigator.permissions.query() — the exact same source the browser
    # lock-icon popup reads from. No hardcoded fallbacks.
    site_data: dict = {
        "full_url":        page.url,
        "permissions":     {},   # name → 'granted'/'denied'/'prompt'
        "cookie_count":    0,
        "cookie_in_use":   False,
        "js_enabled":      True,   # updated below
        "content_type":    "",
    }

    # Permissions: query every relevant name the browser popup can show
    _PERMISSION_NAMES = [
        "geolocation", "notifications", "camera", "microphone",
        "clipboard-read", "clipboard-write", "midi", "ambient-light-sensor",
        "accelerometer", "gyroscope", "magnetometer",
    ]
    try:
        raw_perms = await page.evaluate("""async (names) => {
            const out = {};
            for (const name of names) {
                try {
                    const r = await navigator.permissions.query({name});
                    out[name] = r.state;          // 'granted' | 'denied' | 'prompt'
                } catch (_) {
                    out[name] = null;             // permission type not supported
                }
            }
            return out;
        }""", _PERMISSION_NAMES)
        site_data["permissions"] = {k: v for k, v in raw_perms.items() if v is not None}
    except Exception as e:
        print(f"[cert] permissions.query failed: {e}")

    # JavaScript: can we execute? If the above evaluate succeeded, JS is on.
    # Also try CDP Runtime to double-check.
    try:
        cdp_rt = await page.context.new_cdp_session(page)
        rt_result = await cdp_rt.send("Runtime.evaluate", {"expression": "1+1"})
        site_data["js_enabled"] = rt_result.get("result", {}).get("value") == 2
        await cdp_rt.detach()
    except Exception:
        site_data["js_enabled"] = True  # can't tell — assume enabled

    # Cookies: real count for this origin
    try:
        origin_url = f"{parsed.scheme}://{hostname}"
        cookies = await page.context.cookies([origin_url])
        site_data["cookie_count"]  = len(cookies)
        site_data["cookie_in_use"] = len(cookies) > 0
    except Exception as e:
        print(f"[cert] cookie query failed: {e}")

    # Content-Type: grab from response headers if available
    try:
        content_type = await page.evaluate(
            "document.contentType || document.querySelector('meta[http-equiv=\"content-type\"]')?.content || ''"
        )
        site_data["content_type"] = content_type
    except Exception:
        pass

    # ── 5. Native OS popup screenshots via pywinauto ─────────────────────────
    # Clicks the real Chrome/Edge lock icon, screenshots each real popup panel.
    # Uses browser PID from Playwright for precise window targeting.
    # On non-Windows or if pywinauto/mss not installed → returns empty panels.
    native_shots: dict = {}

    # ── Bring our page to front before native capture ────────────────────────
    # Prevents "user switched tab" issue where omnibox shows wrong URL
    try:
        await page.bring_to_front()
        import asyncio as _asyncio
        await _asyncio.sleep(0.5)   # let Chrome re-render the address bar
    except Exception:
        pass

    # ── Get browser PID — try 3 methods in order ─────────────────────────────
    browser_pid = None

    # Method 1: CDP debug HTTP API — most reliable, returns the MAIN browser PID
    # GET http://localhost:PORT/json/version → {"Browser": "Chrome/...", "pid": N}
    try:
        import urllib.request as _urllib, json as _json, re as _re2
        br = page.context.browser
        ws_url = getattr(br, "ws_endpoint", "") or ""
        port_match = _re2.search(r"127\.0\.0\.1:(\d+)", ws_url)
        if port_match:
            debug_port = port_match.group(1)
            api_url = f"http://127.0.0.1:{debug_port}/json/version"
            with _urllib.urlopen(api_url, timeout=3) as resp:
                version_info = _json.loads(resp.read())
            # Chrome returns webSocketDebuggerUrl like ws://127.0.0.1:PORT/devtools/browser/UUID
            # The process running that port IS the main browser process
            # Use psutil to find it by port ownership
            import psutil as _psutil
            for conn in _psutil.net_connections(kind="tcp"):
                if (conn.laddr.port == int(debug_port) and
                        conn.status == "LISTEN" and conn.pid):
                    browser_pid = conn.pid
                    print(f"[cert] Got PID={browser_pid} via debug port {debug_port} socket owner")
                    break
    except Exception as e:
        print(f"[cert] CDP port PID method failed: {e}")

    # Method 2: psutil — find Chrome process that OWNS the debug port in its cmdline
    # Walk up to parent if needed — renderer processes don't have the port flag
    if not browser_pid:
        try:
            import psutil as _psutil, re as _re3
            br = page.context.browser
            ws_url = getattr(br, "ws_endpoint", "") or ""
            port_match = _re3.search(r":(\d{4,5})/", ws_url)
            debug_port = port_match.group(1) if port_match else None

            for proc in _psutil.process_iter(["pid", "name", "cmdline", "ppid"]):
                try:
                    name = (proc.info.get("name") or "").lower()
                    cmdline = proc.info.get("cmdline") or []
                    if "chrome" not in name and "chromium" not in name:
                        continue
                    cmdline_str = " ".join(str(a) for a in cmdline)
                    # Only the MAIN browser process has --remote-debugging-port
                    # Renderer/GPU/utility processes do NOT have this flag
                    if debug_port and f"remote-debugging-port={debug_port}" in cmdline_str:
                        browser_pid = proc.info["pid"]
                        print(f"[cert] Got PID={browser_pid} via psutil port={debug_port}")
                        break
                    elif not debug_port and "remote-debugging-port" in cmdline_str:
                        browser_pid = proc.info["pid"]
                        print(f"[cert] Got PID={browser_pid} via psutil (any port)")
                        break
                except (_psutil.NoSuchProcess, _psutil.AccessDenied):
                    continue
        except Exception as e:
            print(f"[cert] psutil PID search failed: {e}")

    # Method 3: Playwright internals
    if not browser_pid:
        try:
            br = page.context.browser
            for attr_path in [["_impl_obj", "process", "pid"], ["process", "pid"]]:
                obj = br
                for attr in attr_path:
                    obj = getattr(obj, attr, None)
                    if obj is None:
                        break
                if obj and isinstance(obj, int):
                    browser_pid = obj
                    print(f"[cert] Got PID={browser_pid} via Playwright internals")
                    break
        except Exception:
            pass

    print(f"[cert] Browser PID for native capture: {browser_pid}")

    try:
        native_shots = await asyncio.get_event_loop().run_in_executor(
            None, _native_browser_screenshot_sync, page.url, browser_pid
        )
        if native_shots.get("panel1"):
            print(f"[cert] [OK] Native screenshots captured (PID={browser_pid})")
        else:
            print(f"[cert] [WARN] Native capture returned empty")
    except Exception as e:
        print(f"[cert] Native capture error: {e}")

    # ── 6. Return results (Native only, no fallback) ────────────────────────
    native_ok = bool(native_shots.get("panel1"))
    if native_ok:
        p1 = native_shots["panel1"]
        p2 = native_shots.get("panel2", b"")
        p3 = native_shots.get("panel3", b"")
        composite = native_shots.get("composite") or await _stitch_native(p1, p2, p3, hostname)
    else:
        # User requested NO fallback. If native fails, return empty screenshots.
        p1 = p2 = p3 = composite = b""

    return {
        "screenshot":               composite,
        "screenshot_site_info":     p1,
        "screenshot_security":      p2,
        "screenshot_cert":          p3,
        "screenshot_site_info_html": b"", # disabled
        "screenshot_security_html":  b"", # disabled
        "screenshot_cert_html":      b"", # disabled
        "cert_info":                cert_info,
        "site_data":                site_data,
        "is_valid":                 is_valid,
        "is_secure":                is_secure,
        "hostname":                 hostname,
        "native_capture":           native_ok,
    }


def _native_browser_screenshot_sync(page_url: str, browser_pid: int = None) -> dict:
    """
    Click the real Chrome/Edge lock icon and screenshot each popup panel.

    Uses pywinauto with the UIA backend — zero hardcoded coordinates.
    Window is found by browser PID (passed from Playwright) or by class name.
    Lock button found by UIA AutomationId / Name from the accessibility tree.
    Popup bounding rect from UIA → captured by mss.

    Requires (Windows only):
        pip install pywinauto mss Pillow

    Returns:
        dict with panel1, panel2, panel3, composite  (PNG bytes each)
        or {} on failure / non-Windows.
    """
    import sys, time, io
    if sys.platform != "win32":
        return {}

    try:
        from pywinauto import Application
        from pywinauto.keyboard import send_keys
        import mss, mss.tools
        from PIL import Image as _PILImage
    except ImportError as e:
        print(f"[native] Install required packages:  pip install pywinauto mss Pillow")
        print(f"[native] Missing: {e}")
        return {}

    results: dict = {}

    def _safe_text(ctrl):
        try:
            t = getattr(ctrl, "window_text", None)
            if not t: return ""
            val = t() if callable(t) else t
            return str(val or "").strip()
        except Exception:
            return ""

    def _safe_aid(ctrl):
        try:
            a = getattr(ctrl, "automation_id", None)
            if a is None and hasattr(ctrl, "element_info"):
                a = getattr(ctrl.element_info, "automation_id", None)
            if not a: return ""
            val = a() if callable(a) else a
            return str(val or "").strip()
        except Exception:
            return ""

    # ── Helper: screenshot a rectangle from screen ───────────────────────────
    def _shot(left, top, width, height) -> bytes:
        if width <= 0 or height <= 0:
            return b""
        with mss.mss() as sct:
            raw = sct.grab({"left": left, "top": top,
                            "width": width, "height": height})
            return mss.tools.to_png(raw.rgb, raw.size)

    def _shot_rect(rect) -> bytes:
        """Screenshot a pywinauto RECT object."""
        return _shot(rect.left, rect.top,
                     rect.right - rect.left, rect.bottom - rect.top)

    # ── Step 1: Connect to the correct browser window ────────────────────────
    app = None
    win = None

    # Try PID-based connection first (exact — no ambiguity)
    if browser_pid:
        try:
            app = Application(backend="uia").connect(process=browser_pid, timeout=5)
            win = app.top_window()
            print(f"[native] Connected via PID={browser_pid}")
        except Exception as e:
            print(f"[native] PID connect failed: {e} — falling back to window search")
            app = None

    # Fallback: find the right Chrome window — PID-less search
    if win is None:
        try:
            from pywinauto import Desktop
            desktop = Desktop(backend="uia")
            url_hostname = page_url.split("/")[2] if "//" in page_url else page_url

            # Get all real Chrome windows (skip tiny helper processes < 200px)
            all_wins = [
                w for w in desktop.windows(class_name="Chrome_WidgetWin_1")
                if (w.rectangle().right - w.rectangle().left) > 200
            ]
            print(f"[native] Found {len(all_wins)} Chrome windows")

            best_win  = None
            title_win = None

            for w in all_wins:
                title = w.window_text() or ""

                # ── Method A: read omnibox via UIA ────────────────────────────
                omnibox_url = ""
                for aid in ("omnibox", "LocationBar", "address-and-search-bar"):
                    try:
                        ob = w.child_window(auto_id=aid, control_type="Edit",
                                            found_index=0)
                        if ob.exists(timeout=0.5):
                            omnibox_url = ob.get_value() or ""
                            if omnibox_url:
                                break
                    except Exception:
                        pass

                # Fallback: scan Edit descendants for one that looks like a URL
                if not omnibox_url:
                    try:
                        for e in w.descendants(control_type="Edit")[:8]:
                            try:
                                val = e.get_value() or ""
                                if val.startswith("http") or url_hostname in val:
                                    omnibox_url = val
                                    break
                            except Exception:
                                pass
                    except Exception:
                        pass

                if omnibox_url:
                    print(f"[native] Omnibox '{(w.window_text() or '')[:30]}': {omnibox_url[:60]}")
                    if url_hostname in omnibox_url:
                        best_win = w
                        print(f"[native] [OK] URL matched in omnibox")
                        break

                # ── Method B: title match ─────────────────────────────────────
                if url_hostname in title.lower() and title_win is None:
                    title_win = w

            win = best_win or title_win
            if win:
                src = "omnibox" if best_win else "title"
                print(f"[native] Selected: '{(win.window_text() or '')[:55]}' ({src})")
            elif all_wins:
                # Last resort: pick largest window by child count
                try:
                    win = max(all_wins,
                              key=lambda w: len(list(w.descendants())) if hasattr(w, 'descendants') else 0)
                except Exception:
                    win = all_wins[0]
                print(f"[native] [WARN] No match — using largest window: "
                      f"'{_safe_text(win)[:50]}'")

        except Exception as e:
            print(f"[native] Window search failed: {e}")
            return {}

    if win is None:
        print("[native] [ERROR] No Chrome window found")
        return {}

    try:
        print(f"[native] Attempting to set focus to window: '{_safe_text(win)[:50]}'")
        win.set_focus()
        time.sleep(0.5)
        print("[native] Focus set.")
    except Exception as e:
        print(f"[native] Could not focus window: {e}")

    # ── Step 2: Find the lock icon / security button ─────────────────────────
    lock_btn = None

    # Pass 1: AutomationId — most reliable across Chrome versions
    for aid in ("view-site-information-button", "security-indicator",
                "location-icon-or-bubble", "page-info-button",
                "LocationIconView", "security_status_chip_view",
                "PageInfoChip", "omnibox-icon"):
        try:
            btn = win.child_window(auto_id=aid)
            if btn.exists(timeout=0.5):
                lock_btn = btn
                print(f"[native] Lock button by AutomationId='{aid}'  "
                      f"Name='{_safe_text(btn)}'")
                break
        except Exception:
            continue

    # Pass 2: Fast descendant scan by control type
    if lock_btn is None:
        try:
            print("[native] Scanning Button descendants for lock icon...")
            for btn in win.descendants(control_type="Button"):
                try:
                    name = _safe_text(btn).lower()
                    aid  = _safe_aid(btn).lower()
                    # Log every button found during scan to see what's available
                    # print(f"  [scan] Found Button: name='{name}' auto_id='{aid}'")
                    
                    if any(kw in name for kw in ("view site information", "connection is secure", 
                                                 "connection not secure", "not secure", "certificate")):
                        lock_btn = btn
                        print(f"[native] Lock button found by Name: '{name}'")
                        break
                    if any(kw in aid for kw in ("security", "site-info", "location-icon", 
                                                "page-info", "lock", "omnibox-icon")):
                        lock_btn = btn
                        print(f"[native] Lock button found by AutomationId: '{aid}'")
                        break
                except Exception:
                    continue
        except Exception as e:
            print(f"[native] Button descendant scan failed: {e}")

    if lock_btn is None:
        print("[native] [ERROR] Lock icon not found after 3 passes — dumping toolbar buttons:")
        try:
            for ctrl in win.descendants():
                try:
                    ct = ctrl.element_info.control_type or ""
                    if "Button" in ct:
                        aid = _safe_aid(ctrl)
                        nm  = _safe_text(ctrl)
                        if aid or nm:
                            print(f"  Button | auto_id='{aid}' name='{nm}'")
                except Exception:
                    pass
        except Exception:
            pass
        return {}

    # ── Step 3: Click lock → Panel 1 (Site info popup) ───────────────────────
    try:
        # click_input is often MORE reliable for triggering UI popups than invoke()
        lock_btn.click_input()
        print("[native] Lock button clicked via click_input()")
    except Exception:
        try:
            lock_btn.invoke()
            print("[native] Lock button invoked via UIA InvokePattern")
        except Exception as e:
            print(f"[native] Could not click lock button: {e}")
            return {}

    time.sleep(1.2)  # wait for popup animation to complete

    # ── Helper: find a popup that appeared after a click ─────────────────────
    def _find_popup(timeout=4.5):
        """Search desktop for Chrome site-info/security popup.

        Two strategies:
          A) Child of the known browser window (agent mode — most reliable)
          B) Any Chrome_WidgetWin_1 on desktop that is smaller than a real browser
             (standalone script mode — browser window found via omnibox, not PID)
        """
        from pywinauto import Desktop
        desktop = Desktop(backend="uia")
        deadline = time.time() + timeout
        popup_classes = (
            "BubbleFrameView", "PageInfoBubbleView", "PageInfoView",
            "BubbleContentsView", "BubbleContents",
        )

        # Record main window size so we can exclude it from size-based search
        try:
            main_rect = win.rectangle()
            main_w = main_rect.right - main_rect.left
            main_h = main_rect.bottom - main_rect.top
        except Exception:
            main_w, main_h = 9999, 9999

        while time.time() < deadline:
            # ── Strategy A: descendant of browser window (more robust) ───────
            try:
                # Use descendants() to find it anywhere in the tree
                for d in win.descendants():
                    try:
                        cls = d.element_info.class_name
                        if cls in popup_classes:
                            print(f"[native] Popup found as descendant: class='{cls}'")
                            return d
                    except Exception:
                        continue
            except Exception:
                pass

            # ── Strategy B: new small Chrome window on desktop ────────────────
            try:
                for w in desktop.windows(class_name="Chrome_WidgetWin_1"):
                    try:
                        rect  = w.rectangle()
                        width  = rect.right  - rect.left
                        height = rect.bottom - rect.top
                        # Must be smaller than main window AND smaller than 750px wide
                        # Chrome site-info popup: ~420×274, security: ~420×310
                        if (width < main_w - 50 and height < main_h - 50
                                and 80 < width < 800 and 80 < height < 1000):
                            title = _safe_text(w)
                            print(f"[native] Popup found as small window: "
                                  f"'{title[:40]}' {width}×{height}")
                            return w
                    except Exception:
                        continue
            except Exception:
                pass

            time.sleep(0.15)
        return None

    # Find Panel 1 popup
    popup1 = _find_popup(timeout=4.0)
    if popup1 is None:
        print("[native] [ERROR] Panel 1 popup did not appear")
        send_keys("{ESC}")
        return {}

    r1 = popup1.rectangle()
    results["panel1"] = _shot_rect(r1)
    print(f"[native] [OK] Panel 1: {r1.right-r1.left}×{r1.bottom-r1.top}px")

    # ── Step 4: Click "Connection is/not secure" row → Panel 2 ───────────────
    sec_btn = None
    try:
        items = popup1.descendants(control_type="Button") + popup1.descendants(control_type="ListItem")
        for item in items:
            name = _safe_text(item).lower()
            if any(kw in name for kw in ("connection is secure", "connection not secure", "not secure", "your connection to")):
                sec_btn = item
                print(f"[native] Security row found: '{name}'")
                break
    except Exception as e:
        print(f"[native] Security row scan failed: {e}")

    if sec_btn is None:
        # Try any clickable item that isn't the close button
        try:
            items = popup1.children(control_type="ListItem")
            if not items:
                items = popup1.children(control_type="Button")
            for item in items:
                if "close" not in _safe_text(item).lower():
                    sec_btn = item
                    break
        except Exception:
            pass

    if sec_btn:
        try:
            sec_btn.invoke()
        except Exception:
            try:
                sec_btn.click_input()
            except Exception:
                pass
        time.sleep(0.8)

        # Find Panel 2 (reuse _find_popup)
        popup2 = _find_popup(timeout=4.0)
        if popup2:
            print(f"[native] Panel 2 found")

        if popup2:
            r2 = popup2.rectangle()
            results["panel2"] = _shot_rect(r2)
            print(f"[native] [OK] Panel 2: {r2.right-r2.left}×{r2.bottom-r2.top}px")

            # ── Step 5: Click "Certificate is valid" → Panel 3 ───────────────
            cert_btn = None
            try:
                items = popup2.descendants(control_type="Button") + popup2.descendants(control_type="ListItem")
                for item in items:
                    name = _safe_text(item).lower()
                    if "certificate" in name or "cert" in name:
                        cert_btn = item
                        print(f"[native] Cert row found: '{name}'")
                        break
            except Exception as e:
                print(f"[native] Cert row scan failed: {e}")

            if cert_btn:
                try:
                    cert_btn.invoke()
                except Exception:
                    try:
                        cert_btn.click_input()
                    except Exception:
                        pass
                time.sleep(1.2)

                # Panel 3: Chrome opens cert viewer as new top-level window
                from pywinauto import Desktop
                popup3 = None
                desktop3 = Desktop(backend="uia")
                deadline3 = time.time() + 5.0
                while time.time() < deadline3:
                    try:
                        for w in desktop3.windows(class_name="Chrome_WidgetWin_1"):
                            title = _safe_text(w).lower()
                            if "cert" in title:
                                popup3 = w
                                print(f"[native] Panel 3 found: '{title[:50]}'")
                                break
                        if popup3:
                            break
                        # Also check as child of main win
                        c = win.child_window(title_re=".*[Cc]ertificate.*", found_index=0)
                        if c.exists(timeout=0.3):
                            popup3 = c
                            break
                    except Exception:
                        pass
                    time.sleep(0.2)

                if popup3:
                    r3 = popup3.rectangle()
                    results["panel3"] = _shot_rect(r3)
                    print(f"[native] [OK] Panel 3: {r3.right-r3.left}×{r3.bottom-r3.top}px")
                else:
                    print("[native] [WARN] Panel 3 (cert viewer) not found")
        else:
            print("[native] [WARN] Panel 2 popup not found")
    else:
        print("[native] [WARN] Security row not found in Panel 1")

    # ── Step 6: Close all popups ──────────────────────────────────────────────
    try:
        send_keys("{ESC}")
        time.sleep(0.15)
        send_keys("{ESC}")
        time.sleep(0.15)
        send_keys("{ESC}")
    except Exception:
        pass

    # ── Step 7: Build composite (side-by-side) from real screenshots ──────────
    if results.get("panel1") or results.get("panel2") or results.get("panel3"):
        try:
            panels = [results[k] for k in ("panel1", "panel2", "panel3")
                      if results.get(k)]
            if len(panels) >= 1:
                imgs = [_PILImage.open(io.BytesIO(p)) for p in panels]
                total_w = sum(im.width for im in imgs) + 20 * (len(imgs) - 1)
                max_h   = max(im.height for im in imgs)
                composite = _PILImage.new("RGB", (total_w, max_h), (240, 242, 245))
                x = 0
                for im in imgs:
                    composite.paste(im, (x, 0))
                    x += im.width + 20
                buf = io.BytesIO()
                composite.save(buf, "PNG")
                results["composite"] = buf.getvalue()
                print(f"[native] [OK] Composite built: {total_w}×{max_h}px")
        except Exception as e:
            print(f"[native] Composite build failed: {e}")

    print(f"[native] Done — panels captured: {[k for k in ('panel1','panel2','panel3') if results.get(k)]}")
    return results


async def _stitch_native(p1: bytes, p2: bytes, p3: bytes, hostname: str) -> bytes:
    """Simple async wrapper to stitch 3 native panel PNGs side-by-side."""
    import io
    from PIL import Image as _PILImage
    try:
        panels = [p for p in (p1, p2, p3) if p]
        imgs   = [_PILImage.open(io.BytesIO(p)) for p in panels]
        total_w = sum(im.width for im in imgs) + 20 * (len(imgs) - 1)
        max_h   = max(im.height for im in imgs)
        out = _PILImage.new("RGB", (total_w, max_h), (30, 31, 33))
        x = 0
        for im in imgs:
            out.paste(im, (x, 0))
            x += im.width + 20
        buf = io.BytesIO()
        out.save(buf, "PNG")
        return buf.getvalue()
    except Exception:
        return p1 or p2 or p3 or b""


async def _render_site_info_panel(
    page, cert_info: dict, is_valid: bool, is_secure: bool,
    hostname: str, site_data: dict = None,
) -> bytes:
    """Render Panel 1 — the site-info popup (lock icon menu level).

    All values (permissions, JS state, cookie count, full URL) come from
    live browser data collected in check_certificate(). No hardcoded defaults.
    """
    import datetime

    sd = site_data or {}
    perms    = sd.get("permissions", {})

    conn_colour = "#81c995" if is_secure else "#f28b82"
    conn_text   = "Connection is secure" if is_secure else "Connection not secure"

    # Notification toggle (just one, like real Chrome)
    notif_state = perms.get("notifications")
    if notif_state == "granted":
        notif_toggle = '<div class="toggle on">&#9679;</div>'
        notif_sub = ""
    elif notif_state == "denied":
        notif_toggle = '<div class="toggle off">&#9679;</div>'
        notif_sub = ""
    else:
        notif_toggle = '<div class="toggle off">&#9679;</div>'
        notif_sub = '<div class="row-sub">Not allowed (default)</div>'

    html = f"""<!DOCTYPE html><html><head><meta charset="utf-8">
<style>
*{{box-sizing:border-box;margin:0;padding:0;}}
html{{zoom:1.5;}}
body{{font-family:'Segoe UI',system-ui,-apple-system,sans-serif;
      -webkit-font-smoothing:antialiased;text-rendering:optimizeLegibility;
      background:#1e1f21;display:flex;align-items:center;justify-content:center;padding:24px;}}
.panel{{background:#292a2d;border-radius:12px;
        box-shadow:0 2px 10px rgba(0,0,0,.4),0 6px 30px rgba(0,0,0,.3);
        width:320px;overflow:hidden;color:#e8eaed;}}
.header{{display:flex;align-items:center;justify-content:space-between;
         padding:14px 16px 12px;}}
.hostname{{font-size:14px;font-weight:600;color:#e8eaed;}}
.close{{color:#9aa0a6;font-size:18px;cursor:pointer;}}
.divider{{height:1px;background:#3c3d41;}}
.row{{display:flex;align-items:center;justify-content:space-between;
      padding:12px 16px;cursor:pointer;}}
.row-left{{display:flex;align-items:center;gap:12px;}}
.row-icon{{color:#9aa0a6;font-size:18px;width:22px;text-align:center;}}
.row-label{{font-size:13.5px;color:#e8eaed;}}
.row-sub{{font-size:11px;color:#9aa0a6;margin-top:2px;}}
.chevron{{font-size:14px;color:#9aa0a6;}}
.conn-colour{{color:{conn_colour};}}
.toggle{{width:36px;height:20px;border-radius:10px;display:flex;
         align-items:center;justify-content:center;font-size:18px;}}
.toggle.on{{background:#8ab4f8;color:#202124;padding-left:16px;}}
.toggle.off{{background:#5f6368;color:#292a2d;padding-right:16px;}}
</style></head><body>
<div class="panel">
  <div class="header">
    <div class="hostname">{hostname}</div>
    <div class="close">&#x2715;</div>
  </div>
  <div class="divider"></div>

  <!-- Connection is secure -->
  <div class="row">
    <div class="row-left">
      <div class="row-icon">&#x1F512;</div>
      <div><div class="row-label conn-colour">{conn_text}</div></div>
    </div>
    <div class="chevron">&#x276F;</div>
  </div>

  <div class="divider"></div>

  <!-- Notifications -->
  <div class="row">
    <div class="row-left">
      <div class="row-icon">&#x1F514;</div>
      <div>
        <div class="row-label">Notifications</div>
        {notif_sub}
      </div>
    </div>
    {notif_toggle}
  </div>

  <div class="divider"></div>

  <!-- Cookies and site data -->
  <div class="row">
    <div class="row-left">
      <div class="row-icon">&#x1F36A;</div>
      <div><div class="row-label">Cookies and site data</div></div>
    </div>
    <div class="chevron">&#x276F;</div>
  </div>

  <!-- Site settings -->
  <div class="row">
    <div class="row-left">
      <div class="row-icon">&#x2699;&#xFE0F;</div>
      <div><div class="row-label">Site settings</div></div>
    </div>
    <div style="font-size:14px;color:#9aa0a6">&#x2197;</div>
  </div>
</div>
</body></html>"""

    cert_page = await page.context.new_page()
    try:
        await cert_page.set_content(html, wait_until="networkidle")
        shot = await cert_page.locator(".panel").screenshot(type="png")
    finally:
        await cert_page.close()
    return shot


async def _render_cert_detail_panel(
    page, cert_info: dict, is_valid: bool, hostname: str,
) -> bytes:
    """Render Panel 3 — the 'Certificate is valid' detail drill-down."""
    # Fingerprint rows (SHA-256)
    cert_fp = cert_info.get("fingerprint_sha256", cert_info.get("fingerprint", "—"))
    pubkey_fp = cert_info.get("pubkey_fingerprint", "—")

    def row(lbl, val):
        if not val or val == "—":
            return ""
        return (f'<div class="row"><span class="lbl">{lbl}</span>'
                f'<span class="val">{val}</span></div>')

    html = f"""<!DOCTYPE html><html><head><meta charset="utf-8">
<style>
*{{box-sizing:border-box;margin:0;padding:0;}}
html{{zoom:1.6;}}
body{{font-family:'Segoe UI',system-ui,-apple-system,sans-serif;
      -webkit-font-smoothing:antialiased;text-rendering:optimizeLegibility;
      background:#1e1f21;display:flex;align-items:center;justify-content:center;padding:24px;}}
.dialog{{background:#fff;border-radius:8px;
         box-shadow:0 4px 20px rgba(0,0,0,.5);
         width:420px;overflow:hidden;color:#202124;}}
.title-bar{{display:flex;align-items:center;justify-content:space-between;
            padding:12px 16px;background:#f0f4f8;border-bottom:1px solid #dadce0;}}
.title-text{{font-size:13px;font-weight:600;color:#202124;}}
.close{{color:#5f6368;font-size:16px;cursor:pointer;}}
.tabs{{display:flex;border-bottom:2px solid #e8eaed;padding:0 16px;}}
.tab{{padding:10px 16px;font-size:12px;color:#5f6368;cursor:pointer;
      border-bottom:2px solid transparent;margin-bottom:-2px;}}
.tab.active{{color:#1a73e8;border-bottom-color:#1a73e8;font-weight:600;}}
.section-hdr{{padding:16px 16px 6px;font-size:12px;font-weight:700;
              color:#1a73e8;}}
.row{{display:flex;padding:4px 16px 4px 32px;font-size:12px;}}
.lbl{{color:#5f6368;width:140px;flex-shrink:0;}}
.val{{color:#202124;word-break:break-word;}}
.section-space{{height:8px;}}
.fp-section{{padding:16px 16px 6px;font-size:12px;font-weight:700;
             color:#1a73e8;}}
.fp-row{{padding:4px 16px 4px 32px;font-size:12px;}}
.fp-lbl{{color:#5f6368;margin-bottom:2px;}}
.fp-val{{font-family:'Consolas',monospace;font-size:11px;color:#202124;
         word-break:break-all;line-height:1.5;}}
.bottom-pad{{height:16px;}}
</style></head><body>
<div class="dialog">
  <div class="title-bar">
    <div class="title-text">Certificate Viewer: {hostname}</div>
    <div class="close">&#x2715;</div>
  </div>
  <div class="tabs">
    <div class="tab active">General</div>
    <div class="tab">Details</div>
  </div>

  <div class="section-hdr">Issued To</div>
  {row("Common Name (CN)", cert_info.get("subject_cn", ""))}
  {row("Organization (O)", cert_info.get("subject_org", ""))}
  {row("Organizational Unit (OU)", cert_info.get("subject_ou", "&lt;Not Part Of Certificate&gt;"))}

  <div class="section-space"></div>

  <div class="section-hdr">Issued By</div>
  {row("Common Name (CN)", cert_info.get("issuer_cn", ""))}
  {row("Organization (O)", cert_info.get("issuer_org", ""))}
  {row("Organizational Unit (OU)", cert_info.get("issuer_ou", "&lt;Not Part Of Certificate&gt;"))}

  <div class="section-space"></div>

  <div class="section-hdr">Validity Period</div>
  {row("Issued On", cert_info.get("valid_from", ""))}
  {row("Expires On", cert_info.get("valid_to", ""))}

  <div class="section-space"></div>

  <div class="fp-section">SHA-256 Fingerprints</div>
  {'<div class="fp-row"><div class="fp-lbl">Certificate</div><div class="fp-val">'
   + cert_fp + '</div></div>' if cert_fp and cert_fp != "—" else ""}
  {'<div class="fp-row"><div class="fp-lbl">Public Key</div><div class="fp-val">'
   + pubkey_fp + '</div></div>' if pubkey_fp and pubkey_fp != "—" else ""}

  <div class="bottom-pad"></div>
</div>
</body></html>"""

    cert_page = await page.context.new_page()
    try:
        await cert_page.set_content(html, wait_until="networkidle")
        shot = await cert_page.locator(".dialog").screenshot(type="png")
    finally:
        await cert_page.close()
    return shot


async def _composite_panels(
    page, panel1: bytes, panel2: bytes, panel3: bytes, hostname: str,
    native: bool = False,
) -> bytes:
    """Stitch the 3 panel PNGs side-by-side with click-flow arrows between them."""
    import base64, datetime

    source_label = (
        "Native OS capture via Windows UI Automation"
        if native else
        "HTML panels via CDP + TLS inspection"
    )

    p1b64 = base64.b64encode(panel1).decode()
    p2b64 = base64.b64encode(panel2).decode()
    p3b64 = base64.b64encode(panel3).decode()

    arrow_svg = """
    <svg width="40" height="60" viewBox="0 0 40 60" xmlns="http://www.w3.org/2000/svg">
      <defs>
        <marker id="ah" markerWidth="8" markerHeight="8" refX="6" refY="3" orient="auto">
          <path d="M0,0 L0,6 L8,3 z" fill="#1a73e8"/>
        </marker>
      </defs>
      <line x1="5" y1="30" x2="30" y2="30"
            stroke="#1a73e8" stroke-width="2" marker-end="url(#ah)"/>
      <text x="2" y="20" font-size="9" fill="#5f6368" font-family="Segoe UI,sans-serif">click</text>
    </svg>"""
    arrow_b64 = base64.b64encode(arrow_svg.encode()).decode()

    html = f"""<!DOCTYPE html><html><head><meta charset="utf-8">
<style>
*{{box-sizing:border-box;margin:0;padding:0;}}
html{{zoom:1.5;}}
body{{
  font-family:'Segoe UI',system-ui,-apple-system,sans-serif;
  -webkit-font-smoothing:antialiased;text-rendering:optimizeLegibility;
  background:linear-gradient(135deg,#1e1f21 0%,#292a2d 100%);
  padding:48px 36px;
}}
.flow{{display:flex;align-items:flex-start;gap:0;}}
.step{{display:flex;flex-direction:column;align-items:center;gap:8px;}}
.step-label{{
  font-size:11px;font-weight:600;color:#9aa0a6;
  text-transform:uppercase;letter-spacing:.5px;
  background:rgba(60,61,65,.7);padding:3px 10px;border-radius:10px;
}}
.step img{{border-radius:12px;
           box-shadow:0 4px 16px rgba(0,0,0,.4),0 1px 4px rgba(0,0,0,.3);}}
.arrow{{
  display:flex;flex-direction:column;align-items:center;justify-content:center;
  padding:0 4px;margin-top:48px;
}}
.arrow-label{{font-size:10px;color:#8ab4f8;font-weight:600;margin-bottom:4px;}}
.title{{
  text-align:center;margin-bottom:20px;
  font-size:15px;font-weight:700;color:#e8eaed;
}}
.subtitle{{text-align:center;margin-bottom:24px;font-size:11px;color:#9aa0a6;}}
.ts{{text-align:center;margin-top:16px;font-size:10px;color:#5f6368;}}
</style></head><body>
  <div class="title">🔒 Security Certificate Check — {hostname}</div>
  <div class="subtitle">Click flow: Site info → Connection secure → Certificate detail</div>
  <div class="flow">

    <div class="step">
      <div class="step-label">① Site Info</div>
      <img src="data:image/png;base64,{p1b64}" width="260">
    </div>

    <div class="arrow">
      <div class="arrow-label">click</div>
      <img src="data:image/svg+xml;base64,{arrow_b64}" width="40">
      <div style="font-size:9px;color:#80868b;margin-top:2px;text-align:center">Connection<br>is secure</div>
    </div>

    <div class="step">
      <div class="step-label">② Security</div>
      <img src="data:image/png;base64,{p2b64}" width="290">
    </div>

    <div class="arrow">
      <div class="arrow-label">click</div>
      <img src="data:image/svg+xml;base64,{arrow_b64}" width="40">
      <div style="font-size:9px;color:#80868b;margin-top:2px;text-align:center">Certificate<br>is valid</div>
    </div>

    <div class="step">
      <div class="step-label">③ Certificate</div>
      <img src="data:image/png;base64,{p3b64}" width="290">
    </div>

  </div>
  <div class="ts">Generated {datetime.datetime.now().strftime('%Y-%m-%d %H:%M:%S')} — {source_label}</div>
</body></html>"""

    comp_page = await page.context.new_page()
    try:
        await comp_page.set_content(html, wait_until="networkidle")
        await comp_page.set_viewport_size({"width": 1100, "height": 800})
        shot = await comp_page.screenshot(type="png", full_page=True)
    finally:
        await comp_page.close()
    return shot


async def _render_cert_panel(
    page,
    cert_info: dict,
    is_valid:  bool,
    is_secure: bool,
    hostname:  str,
) -> bytes:
    """Render a Chrome-style 'Connection is secure' panel and return PNG bytes.

    The panel is built as a styled HTML page and screenshotted via Playwright —
    no native browser chrome interaction needed, works headless or headed,
    Windows/Mac/Linux, any browser.
    """
    days     = cert_info.get("days_remaining", 0)
    expired  = cert_info.get("is_expired", days < 0)
    error    = cert_info.get("error", "")

    # Status colours mirror Chrome's dark theme UI
    if is_valid and is_secure:
        status_colour = "#81c995"
        status_bg     = "#1b3a2a"
        lock_svg = (
            '<svg width="20" height="20" viewBox="0 0 20 20" fill="none" '
            'xmlns="http://www.w3.org/2000/svg">'
            '<rect x="3" y="9" width="14" height="10" rx="2" fill="#81c995"/>'
            '<path d="M7 9V6a3 3 0 016 0v3" stroke="#81c995" stroke-width="2" '
            'stroke-linecap="round" fill="none"/>'
            '<circle cx="10" cy="14" r="1.5" fill="#292a2d"/>'
            '</svg>'
        )
        headline = "Connection is secure"
        subline  = ("Your information (for example, passwords or credit card "
                    "numbers) is private when it is sent to this site.")
    else:
        status_colour = "#f28b82"
        status_bg     = "#3c2020"
        lock_svg = (
            '<svg width="20" height="20" viewBox="0 0 20 20" fill="none" '
            'xmlns="http://www.w3.org/2000/svg">'
            '<path d="M10 2L2 17h16L10 2z" fill="#f28b82"/>'
            '<text x="9.5" y="15" font-size="9" fill="#292a2d" font-weight="bold">!</text>'
            '</svg>'
        )
        headline = "Connection not secure" if not is_secure else "Certificate issue"
        subline  = error or "The certificate for this site is invalid or expired."

    html = f"""<!DOCTYPE html><html lang="en"><head><meta charset="utf-8">
<style>
*{{box-sizing:border-box;margin:0;padding:0;}}
html{{zoom:1.5;}}
body{{
  font-family:'Segoe UI',system-ui,-apple-system,sans-serif;
  -webkit-font-smoothing:antialiased;text-rendering:optimizeLegibility;
  background:#1e1f21;
  min-height:100vh;
  display:flex;align-items:center;justify-content:center;
  padding:24px;
}}
.panel{{
  background:#292a2d;
  border-radius:12px;
  box-shadow:0 2px 10px rgba(0,0,0,.4),0 6px 30px rgba(0,0,0,.3);
  width:340px;
  overflow:hidden;
  color:#e8eaed;
}}
.top-bar{{
  display:flex;align-items:center;padding:14px 16px 12px;
  border-bottom:1px solid #3c3d41;gap:10px;
}}
.back{{
  width:28px;height:28px;border-radius:50%;border:2px solid #8ab4f8;
  display:flex;align-items:center;justify-content:center;
  color:#8ab4f8;font-size:14px;cursor:pointer;flex-shrink:0;
}}
.top-title{{font-size:17px;font-weight:600;color:#e8eaed;flex:1;}}
.close{{color:#9aa0a6;font-size:18px;cursor:pointer;}}
.url-line{{padding:8px 16px 12px;font-size:12px;color:#9aa0a6;
           word-break:break-all;}}
.sec-row{{
  display:flex;align-items:flex-start;gap:12px;
  padding:14px 16px 16px;
}}
.sec-icon{{flex-shrink:0;margin-top:2px;}}
.sec-text .headline{{
  font-size:14px;font-weight:600;color:{status_colour};margin-bottom:6px;
}}
.sec-text .sub{{font-size:12px;color:#9aa0a6;line-height:1.6;}}
.sec-text .sub a{{color:#8ab4f8;text-decoration:none;}}
.divider{{height:1px;background:#3c3d41;}}
.cert-footer{{
  display:flex;align-items:center;gap:10px;
  padding:12px 16px;cursor:pointer;
}}
.cert-check{{
  width:20px;height:20px;border-radius:3px;background:#81c995;
  display:flex;align-items:center;justify-content:center;
  color:#292a2d;font-size:13px;flex-shrink:0;
}}
.cert-label{{font-size:13px;font-weight:500;color:#e8eaed;flex:1;}}
.cert-arrow{{color:#9aa0a6;font-size:16px;}}
</style></head><body>
<div class="panel">
  <div class="top-bar">
    <div class="back">&#8592;</div>
    <div class="top-title">Security</div>
    <div class="close">&#x2715;</div>
  </div>
  <div class="url-line">{hostname}</div>
  <div class="sec-row">
    <div class="sec-icon">{lock_svg}</div>
    <div class="sec-text">
      <div class="headline">{headline}</div>
      <div class="sub">{subline} <a href="#">Learn more</a></div>
    </div>
  </div>
  <div class="divider"></div>
  <div class="cert-footer">
    <div class="cert-check">&#x2713;</div>
    <div class="cert-label">Certificate is valid</div>
    <div class="cert-arrow">&#x2197;</div>
  </div>
</div>
</body></html>"""

    cert_page = await page.context.new_page()
    try:
        await cert_page.set_content(html, wait_until="networkidle")
        panel_el  = cert_page.locator(".panel")
        screenshot = await panel_el.screenshot(type="png")
    finally:
        await cert_page.close()
    return screenshot


async def generate_report(
    scenario: str,
    status: bool,
    screenshot=None,
    output: str = None,
    panels: dict = None,
    print_terminal: bool = True,
) -> None:
    """Print extracted data and save a .docx report in the script's directory.

    This upgraded version renders Markdown-style content in `output`:
    - # Heading -> Word Heading
    - | col1 | col2 | -> Word Table
    - ![alt](path.png) -> Embedded Word Image (resolves relative paths)
    """
    from datetime import datetime
    from docx import Document
    from docx.shared import Inches
    import re
    import io

    if print_terminal:
        print("\n" + "=" * 60)
        print(f"EXTRACTED DATA: {scenario}")
        print("=" * 60)
        if output:
            print(output)
        print("=" * 60 + "\n")

    try:
        import sys as _sys
        timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
        caller_script = os.path.abspath(_sys.argv[0]) if _sys.argv[0] else __file__
        script_dir = os.path.dirname(caller_script)
        report_path = os.path.join(script_dir, f"report_{timestamp}.docx")

        doc = Document()
        doc.add_heading(f"Automation Report: {scenario}", 0)
        doc.add_paragraph(f"Generated: {datetime.now().strftime('%Y-%m-%d %H:%M:%S')}")
        doc.add_paragraph(f"Status: {'SUCCESS' if status else 'FAILED'}")

        if output:
            lines = output.split('\n')
            i = 0
            while i < len(lines):
                line = lines[i].strip()
                if not line:
                    i += 1
                    continue
                
                # Headings: # Header, ## Subheader
                if line.startswith('#'):
                    level = line.count('#')
                    content = line.lstrip('#').strip()
                    doc.add_heading(content, level=min(level, 4))
                    i += 1
                    continue
                
                # Tables: | col | col |
                if line.startswith('|'):
                    table_data = []
                    while i < len(lines) and line.startswith('|'):
                        if not re.search(r'[a-zA-Z0-9]', line): # skip separator lines like |---|
                            i += 1
                            line = lines[i].strip() if i < len(lines) else ""
                            continue
                        # Split by | and filter out empty ends
                        row = [c.strip() for c in line.split('|')[1:-1]]
                        if row:
                            table_data.append(row)
                        i += 1
                        line = lines[i].strip() if i < len(lines) else ""
                    
                    if table_data:
                        try:
                            table = doc.add_table(rows=len(table_data), cols=len(table_data[0]))
                            table.style = 'Table Grid'
                            for r_idx, row_values in enumerate(table_data):
                                for c_idx, val in enumerate(row_values):
                                    if c_idx < len(table.columns):
                                        table.cell(r_idx, c_idx).text = str(val)
                        except Exception as te:
                            doc.add_paragraph(f"[Table Error: {te}]")
                    continue

                # Images: ![alt](path)
                img_match = re.match(r'!\[.*?\]\((.*?)\)', line)
                if img_match:
                    img_path = img_match.group(1)
                    if not os.path.isabs(img_path):
                        img_path = os.path.join(script_dir, img_path)
                    
                    if os.path.exists(img_path):
                        try:
                            doc.add_picture(img_path, width=Inches(6.0))
                        except Exception as ie:
                            doc.add_paragraph(f"[Image Render Error: {ie}]")
                    else:
                        doc.add_paragraph(f"[File not found: {img_path}]")
                    i += 1
                    continue
                
                # Normal paragraph
                doc.add_paragraph(line)
                i += 1

        # Legacy panels/single-screenshot support
        if panels and any(panels.get(k) for k in ("panel1", "panel2", "panel3")):
            for key, label in {"panel1": "Site Info", "panel2": "Identity", "panel3": "Details"}.items():
                if panels.get(key):
                    doc.add_heading(label, level=1)
                    doc.add_picture(io.BytesIO(panels[key]), width=Inches(6.0))
        elif screenshot:
            doc.add_heading("Capture", level=1)
            doc.add_picture(io.BytesIO(screenshot), width=Inches(6.0))

        doc.save(report_path)
        print(f"[+] Report saved: {report_path}")
        return report_path
    except Exception as e:
        print(f"[-] Error generating report: {e}")
        return None
        import traceback
        traceback.print_exc()