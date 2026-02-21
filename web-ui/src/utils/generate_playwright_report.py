"""
Generate DOCX report from Playwright script execution with screenshots and step summaries.
"""

import json
import base64
from pathlib import Path
from typing import Any, Dict, List, Optional
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Inches


def generate_playwright_report(
    history_json_path: Path,
    screenshots_dir: Optional[Path] = None,
    execution_log_path: Optional[Path] = None,
) -> Optional[Path]:
    """
    Generate DOCX report from Playwright execution history.
    
    Args:
        history_json_path: Path to agent history JSON
        screenshots_dir: Directory containing numbered screenshots (0.png, 1.png, etc.)
        execution_log_path: Optional path to execution log JSON with step details
    
    Returns:
        Path to generated DOCX report or None if generation failed
    """
    try:
        with history_json_path.open('r', encoding='utf-8') as f:
            data = json.load(f)
        
        steps = data.get('history', []) if isinstance(data, dict) else data
        
        doc = Document()
        doc.add_heading('Playwright Automation Report', 0)
        
        # Summary section
        doc.add_heading('Execution Summary', level=1)
        summary_para = doc.add_paragraph()
        summary_para.add_run('Script: ').bold = True
        summary_para.add_run(f"{history_json_path.stem}.py")
        
        summary_para = doc.add_paragraph()
        summary_para.add_run('Source: ').bold = True
        summary_para.add_run('AI Browser Agent')
        
        summary_para = doc.add_paragraph()
        summary_para.add_run('Status: ').bold = True
        summary_para.add_run('Execution Complete')
        
        # Steps section
        doc.add_heading('Automation Steps', level=1)
        
        step_counter = 0
        for step_idx, step in enumerate(steps, 1):
            if not isinstance(step, dict):
                continue
            
            model_output = step.get('model_output', {})
            actions = model_output.get('action', [])
            
            if not actions:
                continue
            
            for action_idx, action in enumerate(actions):
                if not isinstance(action, dict) or not action:
                    continue
                
                action_name = list(action.keys())[0]
                params = action.get(action_name, {})
                dom_ctx = action.get('dom_context', {})
                
                step_counter += 1
                
                # Step heading
                doc.add_heading(f'Step {step_counter}: {action_name}', level=2)
                
                # Action details
                details_table = doc.add_table(rows=1, cols=2)
                details_table.style = 'Table Grid'
                
                # Action type
                row = details_table.rows[0]
                row.cells[0].text = 'Action'
                row.cells[1].text = action_name
                
                # Action parameters
                if params:
                    row = details_table.add_row().cells
                    row[0].text = 'Parameters'
                    
                    param_lines = []
                    for key, value in params.items():
                        if key != 'index':
                            param_lines.append(f"{key}: {str(value)[:50]}")
                    row[1].text = '\n'.join(param_lines) if param_lines else 'N/A'
                
                # DOM context (selector used)
                if dom_ctx:
                    row = details_table.add_row().cells
                    row[0].text = 'Selector'
                    attrs = dom_ctx.get('attributes', {})
                    selectors = dom_ctx.get('selectors', {})
                    selector = selectors.get('css', '') or f"#{attrs.get('id', 'N/A')}"
                    row[1].text = selector
                
                # Screenshot if available
                if screenshots_dir and screenshots_dir.exists():
                    screenshot_path = screenshots_dir / f"{step_counter - 1}.png"
                    if screenshot_path.exists():
                        try:
                            doc.add_paragraph()
                            doc.add_paragraph('Screenshot:')
                            doc.add_picture(str(screenshot_path), width=Inches(6))
                        except Exception as e:
                            doc.add_paragraph(f"[Screenshot unavailable: {e}]")
                
                doc.add_paragraph()
        
        # Save report
        report_path = history_json_path.parent / f"{history_json_path.stem}_playwright_report.docx"
        doc.save(str(report_path))
        
        return report_path
    
    except Exception as e:
        print(f"Error generating Playwright report: {e}")
        return None
