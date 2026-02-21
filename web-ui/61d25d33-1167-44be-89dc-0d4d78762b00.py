"""
Standalone Playwright Automation Script
Generated from: 61d25d33-1167-44be-89dc-0d4d78762b00.json

This script is completely standalone and only depends on environment variables
for vault credentials. Creates only a final report with embedded screenshots.

Usage:
  set OTCS_USERNAME=value
  set OTCS_PASSWORD=value
  python 61d25d33-1167-44be-89dc-0d4d78762b00.py
"""

import asyncio
import os
import io
import re
from pathlib import Path
from playwright.async_api import async_playwright
from docx import Document
from docx.shared import Inches
import time

def clean_text(raw: str) -> str:
    """Remove script/style blocks and collapse whitespace for readability."""
    if not raw:
        return ""
    text = re.sub(r"<script.*?>.*?</script>", " ", raw, flags=re.IGNORECASE | re.DOTALL)
    text = re.sub(r"<style.*?>.*?</style>", " ", text, flags=re.IGNORECASE | re.DOTALL)
    text = re.sub(r"\s+", " ", text).strip()
    return text

async def parse_table_data(element) -> str:
    """Parse table data intelligently - works with any table structure."""
    try:
        # First, try to get structured data using DOM queries
        rows = await element.query_selector_all("tr")
        if not rows or len(rows) < 2:
            return await element.inner_text()
        
        result = []
        
        # Process each row
        row_data_list = []
        for row in rows:
            cells = await row.query_selector_all("td, th")
            if not cells:
                continue
            
            # Extract text from cells
            cell_texts = []
            for cell in cells:
                text = (await cell.inner_text()).strip()
                # Clean up excessive whitespace
                text = ' '.join(text.split())
                if text:
                    cell_texts.append(text)
            
            if cell_texts:
                row_data_list.append(cell_texts)
        
        if not row_data_list:
            return await element.inner_text()
        
        # Identify header row and data rows
        # Headers typically have keywords or are all short strings
        headers = None
        header_idx = -1
        header_keywords = {'name', 'id', 'status', 'type', 'description', 'processes', 'running', 'errors', 
                          'count', 'state', 'size', 'date', 'user', 'created', 'modified', 'action'}
        
        for idx, row in enumerate(row_data_list):
            # Count how many header-like keywords are in this row
            keyword_count = sum(1 for cell in row for kw in header_keywords if kw in cell.lower())
            # Check if row has many short strings (typical of headers)
            short_count = sum(1 for cell in row if len(cell) < 20)
            
            # Header rows have multiple keywords AND mostly short strings
            if keyword_count >= 2 and short_count >= len(row) * 0.6:
                headers = row
                header_idx = idx
                break
        
        # If no header found by keywords, assume first row is header if it's different from second row
        if headers is None and len(row_data_list) >= 2:
            first = row_data_list[0]
            second = row_data_list[1]
            # First row is header if it has same length and different content
            if len(first) == len(second):
                # Check if first row looks more like labels
                first_has_text = all(isinstance(c, str) and any(w in c.lower() for w in header_keywords) for c in first if c)
                if first_has_text:
                    headers = first
                    header_idx = 0
                else:
                    headers = first  # Default to first row
                    header_idx = 0
            else:
                headers = first
                header_idx = 0
        
        if headers is None:
            headers = row_data_list[0]
            header_idx = 0
        
        # Extract data rows (rows after header)
        data_rows = row_data_list[header_idx + 1:]
        
        # Format results
        for data_row in data_rows:
            if not data_row or all(not cell for cell in data_row):
                continue
            
            # Match cells to headers by position
            row_dict = {}
            for i, header in enumerate(headers):
                if i < len(data_row):
                    value = data_row[i]
                    if value:  # Only include non-empty values
                        row_dict[header] = value
            
            # Also add any extra columns that don't have headers
            if len(data_row) > len(headers):
                for i in range(len(headers), len(data_row)):
                    row_dict[f"Column_{i+1}"] = data_row[i]
            
            if row_dict:
                formatted = '\n'.join([f"  {k}: {v}" for k, v in row_dict.items()])
                result.append(formatted)
        
        return '\n\n'.join(result) if result else await element.inner_text()
        
    except Exception as e:
        # Fallback to plain text if anything fails
        return await element.inner_text()

async def run_automation():
    """Execute the automation workflow with hardcoded steps and generate report."""
    # Setup
    run_id = "61d25d33-1167-44be-89dc-0d4d78762b00"
    script_dir = Path(__file__).parent
    report_path = script_dir / f"{run_id}_report.docx"
    
    # In-memory storage for steps and screenshots
    steps_data = []
    
    async with async_playwright() as p:
        # Launch browser in headed mode to see actions in real-time
        browser = await p.chromium.launch(headless=False)
        page = await browser.new_page()
        step_counter = 0
        try:

            # Step execution
            step_counter += 1
            print(f"Step {step_counter}: go_to_url")
            time.sleep(1)  # Pause to see action in browser
            await page.goto("http://otcsvm.otxlab.net:8080/OTCS/cs.exe")

            # Capture screenshot and data from live execution
            screenshot_bytes = await page.screenshot()
            step_output = f"URL: {page.url}"
            steps_data.append({"step": step_counter, "action": "go_to_url", "screenshot": screenshot_bytes, "output": step_output})

            # Step execution
            step_counter += 1
            print(f"Step {step_counter}: input_text")
            time.sleep(1)  # Pause to see action in browser
            cred = os.getenv("OTCS_USERNAME")
            if cred:
                await page.fill("#otds_username", cred)
            else:
                print(f"⚠️ Warning: OTCS_USERNAME not set in environment")

            # Capture screenshot and data from live execution
            screenshot_bytes = await page.screenshot()
            step_output = f"URL: {page.url}"
            steps_data.append({"step": step_counter, "action": "input_text", "screenshot": screenshot_bytes, "output": step_output})

            # Step execution
            step_counter += 1
            print(f"Step {step_counter}: input_text")
            time.sleep(1)  # Pause to see action in browser
            cred = os.getenv("OTCS_PASSWORD")
            if cred:
                await page.fill("#otds_password", cred)
            else:
                print(f"⚠️ Warning: OTCS_PASSWORD not set in environment")

            # Capture screenshot and data from live execution
            screenshot_bytes = await page.screenshot()
            step_output = f"URL: {page.url}"
            steps_data.append({"step": step_counter, "action": "input_text", "screenshot": screenshot_bytes, "output": step_output})

            # Step execution
            step_counter += 1
            print(f"Step {step_counter}: click_element_by_index")
            time.sleep(1)  # Pause to see action in browser
            await page.click("#otds_password")
            await page.keyboard.press("Enter")
            await page.wait_for_load_state("networkidle")

            # Capture screenshot and data from live execution
            screenshot_bytes = await page.screenshot()
            step_output = f"URL: {page.url}"
            steps_data.append({"step": step_counter, "action": "click_element_by_index", "screenshot": screenshot_bytes, "output": step_output})

            # Step execution
            step_counter += 1
            print(f"Step {step_counter}: go_to_url")
            time.sleep(1)  # Pause to see action in browser
            await page.goto("http://otcsvm.otxlab.net:8080/OTCS/cs.exe?func=ll&objtype=148&objaction=browse")

            # Capture screenshot and data from live execution
            screenshot_bytes = await page.screenshot()
            step_output = f"URL: {page.url}"
            steps_data.append({"step": step_counter, "action": "go_to_url", "screenshot": screenshot_bytes, "output": step_output})

            # Step execution
            step_counter += 1
            print(f"Step {step_counter}: extract_content")
            time.sleep(1)  # Pause to see action in browser
            # Extract content from element
            element = await page.query_selector("table#LLInnerContainer > tbody:nth-child(1) > tr:nth-child(1) > td.cs-valign-top:nth-child(1) > header:nth-child(1)")
            if element:
                extracted_content = await element.text_content()
            else:
                extracted_content = "Element not found"

            # Capture screenshot and data from live execution
            screenshot_bytes = await page.screenshot()
            if isinstance(extracted_content, dict):
                raw_text = extracted_content.get('text', '')
                cleaned = clean_text(raw_text)
                step_output = f"Text: {cleaned[:500]}" if cleaned else "Text: [empty]"
            elif isinstance(extracted_content, str):
                cleaned = clean_text(extracted_content)
                step_output = f"Extracted: {cleaned[:500]}" if cleaned else "Extracted: [empty]"
            else:
                step_output = str(extracted_content)[:500] if extracted_content else None
            steps_data.append({"step": step_counter, "action": "extract_content", "screenshot": screenshot_bytes, "output": step_output})

            # Step execution
            step_counter += 1
            print(f"Step {step_counter}: done")
            time.sleep(1)  # Pause to see action in browser
            pass  # Task complete

            # Capture screenshot and data from live execution
            screenshot_bytes = await page.screenshot()
            step_output = f"URL: {page.url}"
            steps_data.append({"step": step_counter, "action": "done", "screenshot": screenshot_bytes, "output": step_output})

        except Exception as e:
            print(f"Error during automation: {e}")
            await browser.close()
            raise
        finally:
            await browser.close()

        # Generate report with embedded screenshots
        print("\nGenerating report...")
        doc = Document()
        doc.add_heading("Playwright Automation Report", 0)
        
        # Summary section
        doc.add_heading("Execution Summary", level=1)
        summary_para = doc.add_paragraph()
        summary_para.add_run("Script: ").bold = True
        summary_para.add_run(f"{run_id}.py")
        
        summary_para = doc.add_paragraph()
        summary_para.add_run("Status: ").bold = True
        summary_para.add_run("Execution Complete")
        
        summary_para = doc.add_paragraph()
        summary_para.add_run("Total Steps: ").bold = True
        summary_para.add_run(str(len(steps_data)))
        
        # Captured outputs summary
        doc.add_heading("Captured Outputs", level=1)
        outputs = [entry for entry in steps_data if entry.get("output")]
        if outputs:
            table = doc.add_table(rows=1, cols=3)
            table.style = "Table Grid"
            hdr = table.rows[0].cells
            hdr[0].text = "Step"
            hdr[1].text = "Action"
            hdr[2].text = "Output"
            for o in outputs:
                row = table.add_row().cells
                row[0].text = str(o.get("step"))
                row[1].text = str(o.get("action"))
                row[2].text = str(o.get("output"))[:500]
        else:
            doc.add_paragraph("No outputs captured.")
        
        # Steps section with embedded screenshots
        doc.add_heading("Automation Steps", level=1)
        
        for entry in steps_data:
            step_num = entry.get("step")
            action_name = entry.get("action")
            screenshot_bytes = entry.get("screenshot")
            output_data = entry.get("output")
            
            doc.add_heading(f"Step {step_num}: {action_name}", level=2)
            
            # Action details table
            details_table = doc.add_table(rows=2 if output_data else 1, cols=2)
            details_table.style = "Table Grid"
            row = details_table.rows[0]
            row.cells[0].text = "Action"
            row.cells[1].text = action_name
            
            # Add captured output if available
            if output_data:
                row = details_table.rows[1]
                row.cells[0].text = "Captured Output"
                row.cells[1].text = str(output_data)
            
            # Add screenshot if captured
            if screenshot_bytes:
                try:
                    doc.add_paragraph("Screenshot:")
                    screenshot_io = io.BytesIO(screenshot_bytes)
                    doc.add_picture(screenshot_io, width=Inches(6))
                except Exception as e:
                    doc.add_paragraph(f"[Screenshot unavailable: {e}]")
            
            doc.add_paragraph()
        
        # Save report with fallback locations
        save_attempts = [
            report_path,  # Primary: script directory
            script_dir.parent / "tmp" / f"{run_id}_report.docx",  # Fallback: tmp directory
            Path.home() / "Downloads" / f"{run_id}_report.docx",  # Last resort: Downloads
        ]
        
        saved = False
        for attempt_path in save_attempts:
            try:
                attempt_path.parent.mkdir(parents=True, exist_ok=True)
                doc.save(str(attempt_path))
                print(f"\n✓ Report generated: {attempt_path}")
                print(f"✓ Report size: {attempt_path.stat().st_size / 1024:.1f} KB")
                saved = True
                break
            except Exception as e:
                print(f"Failed to save to {attempt_path}: {e}")
        
        if not saved:
            print("\n⚠️ Warning: Could not save report to any location")


if __name__ == "__main__":
    asyncio.run(run_automation())