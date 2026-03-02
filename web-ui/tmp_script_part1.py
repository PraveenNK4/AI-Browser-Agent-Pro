# Part 1 script
import os
import re

content_part1 = r"""from __future__ import annotations

import asyncio
import base64
import io
import json
import logging
import os
import pathlib
import platform
import shutil
import time
import uuid
import builtins
from datetime import datetime
from src.utils.utils import slugify
from typing import Any, AsyncGenerator, Dict, Optional

from browser_use.agent.views import AgentHistoryList, AgentOutput
from browser_use.browser.browser import BrowserConfig
from browser_use.browser.context import BrowserContext, BrowserContextConfig
from browser_use.browser.views import BrowserState
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Inches
from gradio.components import Component
from langchain_core.language_models.chat_models import BaseChatModel
from PIL import Image

# NOTE: Security wait bypass monkeypatch removed for compatibility with browser-use v0.1.48+

from src.agent.browser_use.browser_use_agent import BrowserUseAgent
from src.orchestrator.process_orchestrator import run_all_processes
from src.orchestrator.process_parser import parse_processes
from src.browser.custom_browser import CustomBrowser
from src.controller.custom_controller import CustomController
from src.utils import llm_provider
from src.utils.dom_snapshot import set_run_id
from src.utils.strip_and_enrich_history import strip_and_enrich
from src.utils.llm_script_generator import generate_script as generate_llm_script
from src.webui.webui_manager import WebuiManager
import gradio as gr

logger = logging.getLogger(__name__)


def _detect_browser():
    \"""Detect available browsers and prefer Chrome for automation.\"""
    system = platform.system()
    if system == "Windows":
        chrome_paths = [
            r"C:\Program Files\Google\Chrome\Application\chrome.exe",
            r"C:\Program Files (x86)\Google\Chrome\Application\chrome.exe",
        ]
        edge_paths = [
            r"C:\Program Files\Microsoft\Edge\Application\msedge.exe",
            r"C:\Program Files (x86)\Microsoft\Edge\Application\msedge.exe"
        ]
    elif system == "Darwin":
        chrome_paths = ["/Applications/Google Chrome.app/Contents/MacOS/Google Chrome"]
        edge_paths = ["/Applications/Microsoft Edge.app/Contents/MacOS/Microsoft Edge"]
    else:
        chrome_paths = ["/usr/bin/google-chrome", "/usr/bin/google-chrome-stable"]
        edge_paths = ["/usr/bin/microsoft-edge", "/usr/bin/microsoft-edge-stable"]

    for path in edge_paths + chrome_paths:
        if os.path.exists(path):
            logger.info(f"Detected browser: {path}")
            return path
    logger.warning("No supported browser found.")
    return None

class _SensitiveDataSecurityFilter(logging.Filter):
    def filter(self, record):
        msg = record.getMessage()
        if any(term in msg for term in ["not locked down", "sensitive_data", "malicious website", "Waiting 10 seconds", "insecure settings"]):
            return False
        return True

logging.getLogger("browser_use.agent.service").addFilter(_SensitiveDataSecurityFilter())
logging.getLogger("browser_use.agent").addFilter(_SensitiveDataSecurityFilter())
logging.getLogger("browser_use").addFilter(_SensitiveDataSecurityFilter())

async def _initialize_llm(
    provider: Optional[str],
    model_name: Optional[str],
    temperature: float,
    base_url: Optional[str],
    api_key: Optional[str],
    num_ctx: Optional[int] = None,
) -> Optional[BaseChatModel]:
    \"""Initialize LLM based on settings.\"""
    if not provider or not model_name:
        logger.info("LLM Provider/Model not specified.")
        return None
    try:
        logger.info(f"Initializing LLM: {provider}/{model_name} (temp={temperature})")
        llm = llm_provider.get_llm_model(
            provider=provider,
            model_name=model_name,
            temperature=temperature,
            base_url=base_url or None,
            api_key=api_key or None,
            num_ctx=num_ctx if provider == "ollama" else None,
        )
        return llm
    except Exception as e:
        logger.error(f"Failed to initialize LLM: {e}", exc_info=True)
        gr.Warning(f"Failed to initialize LLM '{model_name}': {e}")
        return None

def _get_config_value(
    webui_manager: WebuiManager,
    comp_dict: Dict[Component, Any],
    comp_id_suffix: str,
    default: Any = None,
) -> Any:
    \"""Get value from component dictionary using ID suffix.\"""
    for prefix in ["browser_use_agent", "agent_settings", "browser_settings"]:
        try:
            comp_id = f"{prefix}.{comp_id_suffix}"
            comp = webui_manager.get_component_by_id(comp_id)
            return comp_dict.get(comp, default)
        except KeyError:
            continue
    logger.warning(f"Component '{comp_id_suffix}' not found.")
    return default

def _format_agent_output(model_output: AgentOutput) -> str:
    \"""Format AgentOutput for chatbot display.\"""
    content = ""
    if model_output:
        try:
            llm_response = ""
            if hasattr(model_output, 'response'):
                response = getattr(model_output, 'response', None)
                if response:
                    llm_response = str(response)
            elif hasattr(model_output, 'text'):
                text = getattr(model_output, 'text', None)
                if text:
                    llm_response = str(text)

            if llm_response:
                content += f"**LLM Reasoning:**\n{llm_response}\n\n"

            action_dump = [action.model_dump(exclude_none=True) for action in model_output.action]
            state_dump = model_output.current_state.model_dump(exclude_none=True)
            output_dump = {"current_state": state_dump, "action": action_dump}
            json_str = json.dumps(output_dump, indent=4, ensure_ascii=False)
            content += f"**Actions:**\n<pre><code class='language-json'>{json_str}</code></pre>"
        except Exception as e:
            logger.error(f"Error formatting output: {e}", exc_info=True)
            content = f"<pre><code>Error formatting output: {e}\nRaw: {str(model_output)}</code></pre>"
    return content.strip()

async def _handle_new_step(
    webui_manager: WebuiManager, state: BrowserState, output: AgentOutput, step_num: int
):
    \"""Callback for each agent step.\"""
    if not hasattr(webui_manager, "bu_chat_history"):
        webui_manager.bu_chat_history = []
    
    step_num -= 1
    step_end_time = time.perf_counter()
    step_duration = 0.0
    if webui_manager.bu_step_start_time:
        step_duration = round(step_end_time - webui_manager.bu_step_start_time, 3)
    
    logger.info(f"Step {step_num} completed.")

    if webui_manager.check_step_progression(step_num):
        logger.error(f"Step {step_num}: Excessive backward progression.")
    
    step_failed = False
    if hasattr(output, 'error'):
        error = getattr(output, 'error', None)
        if error:
            step_failed = True
            logger.warning(f"Step {step_num} failed: {error}")
    
    if webui_manager.check_step_failure(step_num, step_failed):
        logger.error(f"Step {step_num} exceeded 3 failures. Stopping.")
        webui_manager.bu_should_stop_agent = True
        if webui_manager.bu_agent:
            webui_manager.bu_agent.state.stopped = True
    
    current_action = None
    if hasattr(output, 'action') and output.action:
        current_action = str(output.action[0]) if isinstance(output.action, (list, tuple)) else str(output.action)
    
    if webui_manager.check_hallucination(current_action):
        logger.error(f"HALLUCINATION: Action repeated {webui_manager.bu_repeated_action_count} times. Stopping.")
        webui_manager.bu_should_stop_agent = True
        webui_manager.bu_hallucination_triggered = True
        if webui_manager.bu_agent:
            webui_manager.bu_agent.state.stopped = True

    current_tokens = 0
    if webui_manager.bu_agent:
        current_tokens = webui_manager.bu_agent.state.history.total_input_tokens()
    if not hasattr(webui_manager, 'bu_previous_tokens'):
        webui_manager.bu_previous_tokens = 0
    step_tokens = current_tokens - webui_manager.bu_previous_tokens
    webui_manager.bu_previous_tokens = current_tokens

    screenshot_html = ""
    screenshot_data = getattr(state, "screenshot", None)
    if screenshot_data and isinstance(screenshot_data, str) and len(screenshot_data) > 100:
        img_tag = f'<img src="data:image/jpeg;base64,{screenshot_data}" style="max-width:800px;max-height:600px;object-fit:contain;" />'
        screenshot_html = img_tag + "<br/>"

    formatted_output = _format_agent_output(output)

    failure_warning = ""
    if step_failed:
        failure_warning += "⚠️ **STEP FAILED** - Retrying...<br/>"
    if webui_manager.bu_should_stop_agent:
        if webui_manager.bu_hallucination_triggered:
            failure_warning += "🔴 **HALLUCINATION DETECTED** - Stopped.<br/>"
        else:
            failure_warning += "🔴 **MAX FAILURES** - Stopped.<br/>"

    step_header = f"--- **Step {step_num}** (Tokens: {step_tokens}, Time: {step_duration:.3f}s) ---"
    final_content = step_header + "<br/>" + screenshot_html + failure_warning + formatted_output

    webui_manager.bu_chat_history.append({"role": "assistant", "content": final_content.strip()})
    webui_manager.bu_step_start_time = time.perf_counter()
    await asyncio.sleep(0.05)


def generate_docx_report(webui_manager: WebuiManager, history: AgentHistoryList, task: str):
    \"""Generate comprehensive DOCX report with screenshots and findings.\"""
    try:
        logger.info("Generating DOCX report...")
        if not webui_manager or not hasattr(webui_manager, 'bu_agent_task_id'):
            logger.error("Missing webui_manager or task_id")
            return None
        
        doc = Document()
        doc.add_heading('AI Browser Agent Execution Report', 0)

        # Task summary
        doc.add_heading('Task Summary', level=1)
        task_para = doc.add_paragraph()
        task_para.add_run('Task: ').bold = True
        task_para.add_run(task)

        duration = history.total_duration_seconds()
        tokens = history.total_input_tokens()
        final_result = history.final_result()

        summary_table = doc.add_table(rows=5, cols=2)
        summary_table.style = 'Table Grid'
        summary_table.rows[0].cells[0].text = 'Duration'
        summary_table.rows[0].cells[1].text = f"{duration:.2f}s"
        summary_table.rows[1].cells[0].text = 'Tokens'
        summary_table.rows[1].cells[1].text = str(tokens)
        summary_table.rows[2].cells[0].text = 'Steps'
        summary_table.rows[2].cells[1].text = str(len(history.history))
        summary_table.rows[3].cells[0].text = 'Status'
        summary_table.rows[3].cells[1].text = 'Completed' if final_result else 'Failed'
        summary_table.rows[4].cells[0].text = 'Task ID'
        summary_table.rows[4].cells[1].text = webui_manager.bu_agent_task_id or 'Unknown'

        # Final result
        doc.add_heading('Final Result', level=1)
        if final_result:
            result_para = doc.add_paragraph()
            result_para.add_run('Outcome: ').bold = True
            result_para.add_run(final_result)
        
        step_failures = getattr(webui_manager, 'bu_step_failures', {})
        total_failures = sum(step_failures.values())
        status_para = doc.add_paragraph()
        status_para.add_run('Status: ').bold = True
        if total_failures > 0:
            status_para.add_run(f"Completed with {total_failures} failures")
        else:
            status_para.add_run("All steps successful")

        # Step details
        doc.add_heading('Step-by-Step Execution', level=1)
        for i, step in enumerate(history.history, 1):
            step_status = "PASS" if step_failures.get(i, 0) == 0 else f"FAIL ({step_failures.get(i, 0)}x)"
            doc.add_heading(f'Step {i} - {step_status}', level=2)

            goal_para = doc.add_paragraph()
            goal_para.add_run('Goal: ').bold = True
            goal_text = "Execute action"
            if hasattr(step, 'model_output') and step.model_output:
                if hasattr(step.model_output, 'current_state') and step.model_output.current_state:
                    if hasattr(step.model_output.current_state, 'next_goal') and step.model_output.current_state.next_goal:
                        goal_text = str(step.model_output.current_state.next_goal)
            goal_para.add_run(goal_text)

            action_para = doc.add_paragraph()
            action_para.add_run('Action: ').bold = True
            action_desc = "No action"
            if hasattr(step, 'action'):
                action = getattr(step, 'action', None)
                if action:
                    action_desc = str(action[0]) if isinstance(action, (list, tuple)) else str(action)
            action_para.add_run(action_desc)

            if hasattr(step, 'state') and hasattr(step.state, 'screenshot') and step.state.screenshot:
                try:
                    screenshot_data = step.state.screenshot
                    if screenshot_data.startswith('data:image'):
                        screenshot_data = screenshot_data.split(',')[1]
                    image_data = base64.b64decode(screenshot_data)
                    image = Image.open(io.BytesIO(image_data))
                    
                    if image.width > 900:
                        ratio = 900 / image.width
                        image = image.resize((int(image.width * ratio), int(image.height * ratio)), Image.Resampling.LANCZOS)
                    
                    if image.mode in ("RGBA", "LA", "P"):
                        image = image.convert("RGB")
                    
                    img_buffer = io.BytesIO()
                    image.save(img_buffer, format='JPEG', quality=95)
                    img_buffer.seek(0)
                    
                    doc.add_picture(img_buffer, width=Inches(6))
                    caption = doc.add_paragraph(f'Step {i} screenshot')
                    caption.alignment = WD_ALIGN_PARAGRAPH.CENTER
                except Exception as e:
                    logger.warning(f"Failed to add screenshot for step {i}: {e}")

            result_para = doc.add_paragraph()
            result_para.add_run('Result: ').bold = True
            result_text = "No result"
            custom_images = []
            if hasattr(step, 'result') and step.result:
                if isinstance(step.result, list) and len(step.result) > 0:
                    first_result = step.result[0]
                    if hasattr(first_result, 'extracted_content'):
                        extracted = getattr(first_result, 'extracted_content', None)
                        if extracted:
                            result_text = str(extracted)
                        else:
                            result_text = str(first_result)
                    else:
                        result_text = str(first_result)
                else:
                    result_text = str(step.result)
            
            import re
            img_matches = re.findall(r'\[IMAGE_PATH\]\s+(.*?)(?=\n|$)', result_text)
            for img_path in img_matches:
                img_path = img_path.strip()
                if os.path.exists(img_path):
                    custom_images.append(img_path)
                result_text = result_text.replace(f'[IMAGE_PATH] {img_path}', '').strip()

            result_para.add_run(result_text)

            for img_path in custom_images:
                try:
                    doc.add_picture(img_path, width=Inches(6))
                    cap = doc.add_paragraph(f'Action Artifact: {os.path.basename(img_path)}')
                    cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
                except Exception as e:
                    logger.warning(f"Failed to add custom result image {img_path}: {e}")

            doc.add_paragraph()

        errors = history.errors()
        if errors and any(errors):
            doc.add_heading('Errors', level=1)
            for error in errors:
                if error:
                    doc.add_paragraph(f'Error: {str(error)}')

        task_id = webui_manager.bu_agent_task_id
        if not task_id:
            logger.error("Task ID is None")
            return None
        agent_history_path = os.getenv("AGENT_HISTORY_PATH", "./tmp/agent_history")
        task_dir = os.path.join(agent_history_path, task_id)
        os.makedirs(task_dir, exist_ok=True)
        
        safe_filename = task_id.replace("/", "_").replace("\\", "_")
        docx_path = os.path.join(task_dir, f"{safe_filename}_report.docx")
        doc.save(docx_path)
        logger.info(f"DOCX report saved: {docx_path}")
        return docx_path
    except Exception as e:
        logger.error(f"DOCX generation failed: {e}", exc_info=True)
        return None
"""

with open("c:\\Users\\pnandank\\Downloads\\Dynamic_Scrapper_4008 1\\Dynamic_Scrapper_4008\\Dynamic_Scrapper\\web-ui\\tmp_script_part1.py", "w", encoding="utf-8") as f:
    f.write(content_part1)
