# ── Vault credential reader ───────────────────────────────────────────────────
import sys as _sys
from pathlib import Path as _Path

def vault_creds(prefix: str) -> dict:
    """Return {'username': ..., 'password': ...} from the encrypted vault."""
    # Walk up from this script to find the web-ui root
    base = _Path(__file__).resolve()
    for _ in range(8):
        base = base.parent
        candidate = base / "web-ui"
        if candidate.exists():
            _sys.path.insert(0, str(candidate))
            break
    try:
        from src.utils.vault import vault
        creds = vault.get_credentials(prefix) or vault.get_credentials(prefix.lower())
        if not creds:
            raise RuntimeError(f"No vault entry found for prefix '{prefix}'")
        lower = {k.lower(): v for k, v in creds.items()}
        def _pick(cs):
            for c in cs:
                if c in lower: return lower[c]
            return list(lower.values())[0] if lower else ""
        return {
            "username": _pick(("username","user","login","email")),
            "password": _pick(("password","passwd","pwd","pass")),
        }
    except Exception as e:
        raise RuntimeError(f"Vault error for '{prefix}': {e}")
# ─────────────────────────────────────────────────────────────────────────────

import os, sys, logging, asyncio, io, re, base64
from datetime import datetime
from playwright.async_api import async_playwright
from docx import Document
from docx.shared import Inches
from urllib.parse import urljoin

if sys.platform == 'win32':
    asyncio.set_event_loop_policy(asyncio.WindowsProactorEventLoopPolicy())

async def run():
    # INITIALIZE ALL DATA VARIABLES TO EMPTY LISTS/DICTS/STRINGS HERE
    steps = []
    p = None
    browser = None
    new_folder_name = ""
    expected_files = []
    files_in_folder_table = []

    SCRIPT_CONFIG = {
        "base_url": "http://otcsvm.otxlab.net:8080/OTCS/cs.exe/app/nodes/2004",
        "username_selector": "input[name*='user' i], input[id*='user' i], input[name*='login' i]",
        "password_selector": "input[type='password']",
        "login_button_selector": "button[type='submit'], input[type='submit'], #loginbutton, #login_button, .submit-btn",
        "add_item_button_selector": "a[aria-label='Add item'], .csui-add-item.binf-dropdown-toggle",
        "folder_option_selector": "li[data-csui-addtype='0'] a",
        "save_button_name": "Save",
        "frame_width": 1280,
        "frame_height": 900,
        "folder_option_text": "Folder",
        "folder_name_input_selector": "input[placeholder*='name' i], input[aria-label*='name' i], .csui-create-folder-inline input",
        "modified_column_header_selector": "th:has-text('Modified'), div:has-text('Modified'), span:has-text('Modified')",
        "modified_column_header_text": "Modified",
        "document_option_selector": "li[data-csui-addtype='144'] a",
        "document_option_text": "Document",
        "upload_directory_path": r"C:\Users\pnandank\Downloads\OTCS_Upload_Test",
        "files_table_selector": "table[role='grid'], .csui-table-container table",
        "name_column_header_text": "Name",
        "report_filename": f"OTCS_Upload_Report_{datetime.now().strftime('%Y%m%d_%H%M%S')}.docx",
    }

    report_path = os.path.join(os.path.dirname(os.path.abspath(__file__)), SCRIPT_CONFIG["report_filename"])
    doc = Document()

    def clean_id(raw):
        # Normalize whitespace, take first line, truncate to 100 chars
        return re.sub(r'\s+', ' ', (raw or "").split('\n')[0].strip())[:100]

    ## ── PASTE THE CAPTURE_STEP_CODE BLOCK HERE ────────────────────────────────
    import base64
    from urllib.parse import urlparse

    def _build_chrome_frame_html(page_png_b64, url, title, favicon_data_uri="", page_width=1280, page_height=900):
        is_secure = url.startswith("https://")
        parsed = urlparse(url)
        # Show full address in the bar: netloc + path + query
        addr_text = parsed.netloc + parsed.path
        if parsed.query:
            addr_text += "?" + parsed.query
        
        lock_col  = "#188038" if is_secure else "#c5221f"
        tab_title = title[:30] + "…" if len(title) > 32 else title
        
        import time as _t
        _tzn = _t.tzname[_t.daylight] if _t.daylight else _t.tzname[0]
        _tz  = "".join(w[0] for w in _tzn.split()) if len(_tzn) > 5 else _tzn
        ts_label = datetime.now().strftime(f"%Y-%m-%d %H:%M:%S {_tz}")

        if is_secure:
            lock_svg = f'<svg viewBox="0 0 16 16" width="14" height="14" fill="{lock_col}"><path d="M8 1a3.5 3.5 0 0 0-3.5 3.5V6H4a1 1 0 0 0-1 1v6a1 1 0 0 0 1 1h8a1 1 0 0 0 1-1V7a1 1 0 0 0-1-1h-.5V4.5A3.5 3.5 0 0 0 8 1zm2 5H6V4.5a2 2 0 1 1 4 0V6z"/></svg>'
            secure_label_text = 'Secure'
        else:
            lock_svg = f'<svg viewBox="0 0 20 20" width="14" height="14" fill="{lock_col}"><path d="M3.7 3.7a.75.75 0 0 0-1.1 1.1l2.2 2.2H4a1 1 0 0 0-1 1v6a1 1 0 0 0 1 1h8a1 1 0 0 1 .6.2l1.7 1.7a.75.75 0 1 0 1.1-1.1L3.7 3.7zM5.1 9l-.1-.1V9h.1zm1 0H13v5H7.1L6 12.9V9zM7.5 5a2 2 0 0 1 3.9-.5l1.5 1.5H8V4.5A2 2 0 0 0 7.5 5zm2.5-2A3.5 3.5 0 0 0 6.5 6H4.9L3.5 4.6A3.5 3.5 0 0 1 10 3z"/></svg>'
            secure_label_text = 'Not secure'

        if favicon_data_uri:
            favicon_html = f'<img src="{favicon_data_uri}" width="16" height="16" style="flex-shrink:0;border-radius:2px" onerror="this.style.display=\'none\'">'
        else:
            favicon_html = '<svg viewBox="0 0 16 16" width="16" height="16" fill="#5f6368" style="flex-shrink:0"><circle cx="8" cy="8" r="7" fill="none" stroke="#5f6368" stroke-width="1.2"/><ellipse cx="8" cy="8" rx="3" ry="7" fill="none" stroke="#5f6368" stroke-width="1.2"/><line x1="1" y1="8" x2="15" y2="8" stroke="#5f6368" stroke-width="1.2"/></svg>'

        return f"""<!DOCTYPE html><html><head><meta charset="utf-8">
<style>
  html, body {{ overflow: hidden; }}
  * {{ box-sizing: border-box; margin: 0; padding: 0; font-family: 'Segoe UI', -apple-system, system-ui, sans-serif; font-size: 13px; }}
  body {{ background: #e8eaed; display: inline-flex; flex-direction: column; }}
  .window {{ display: flex; flex-direction: column; width: {page_width + 2}px; box-shadow: 0 2px 16px rgba(0,0,0,.25); border-radius: 8px 8px 0 0; overflow: hidden; border: 1px solid #b0b0b0; }}
  .tabbar {{ background: #dee1e6; display: flex; align-items: flex-end; height: 40px; padding-left: 10px; flex-shrink: 0; position: relative; }}
  .win-ctrls {{ position: absolute; right: 0; top: 0; display: flex; }}
  .win-btn {{ width: 44px; height: 32px; display: flex; align-items: center; justify-content: center; color: #5f6368; }}
  .win-btn:hover {{ background: #dadce0; }}
  .win-btn.close:hover {{ background: #e81123; color: #fff; }}
  .tab {{ background: #fff; height: 34px; padding: 0 12px; min-width: 180px; max-width: 240px; display: flex; align-items: center; gap: 8px; border-radius: 8px 8px 0 0; color: #202124; position: relative; }}
  .tab-title {{ flex: 1; overflow: hidden; white-space: nowrap; text-overflow: ellipsis; }}
  .tab-x {{ color: #5f6368; font-size: 16px; width: 18px; text-align: center; }}
  .toolbar {{ background: #fff; height: 40px; display: flex; align-items: center; padding: 0 8px; gap: 4px; border-bottom: 1px solid #dadce0; }}
  .tbtn {{ width: 28px; height: 28px; display: flex; align-items: center; justify-content: center; color: #5f6368; border-radius: 50%; flex-shrink:0; }}
  .tbtn:hover {{ background: #f1f3f4; }}
  .omni {{ flex: 1; height: 32px; background: #f1f3f4; border-radius: 24px; display: flex; align-items: center; gap: 8px; padding: 0 12px; margin: 0 4px; overflow: hidden; }}
  .secure-grp {{ display: flex; align-items: center; gap: 6px; flex-shrink: 0; }}
  .secure-label {{ font-size: 11px; color: {lock_col}; font-weight: 600; white-space: nowrap; line-height: 1; }}
  .addr {{ font-size: 14px; color: #202124; flex: 1; white-space: nowrap; overflow: hidden; text-overflow: ellipsis; }}
  .tools-right {{ display: flex; align-items: center; gap: 2px; padding-right: 4px; }}
  .page img {{ display: block; }}
  .status {{ background: #f8f9fa; height: 22px; display: flex; align-items: center; padding: 0 10px; font-size: 11px; color: #80868b; border-top: 1px solid #e8eaed; }}
  .status-url {{ flex: 1; white-space: nowrap; overflow: hidden; text-overflow: ellipsis; }}
  .status-ts {{ flex-shrink: 0; padding-left: 10px; }}
</style></head><body>
<div class="window">
  <div class="tabbar">
    <div class="tab">{favicon_html}<span class="tab-title">{tab_title}</span><span class="tab-x">×</span></div>
    <div class="win-ctrls">
      <div class="win-btn"><svg viewBox="0 0 10 1" width="10"><path d="M0 0h10v1H0z" fill="currentColor"/></svg></div>
      <div class="win-btn"><svg viewBox="0 0 10 10" width="10"><path d="M0 0h10v10H0V0zm1 1v8h8V1H1z" fill="currentColor"/></svg></div>
      <div class="win-btn close"><svg viewBox="0 0 10 10" width="10"><path d="M0 0l10 10m0-10L0 10" stroke="currentColor" stroke-width="1.2"/></svg></div>
    </div>
  </div>
  <div class="toolbar">
    <div class="tbtn"><svg viewBox="0 0 24 24" width="16" height="16" fill="currentColor"><path d="M20 11H7.8l5.6-5.6L12 4l-8 8 8 8 1.4-1.4L7.8 13H20v-2z"/></svg></div>
    <div class="tbtn" style="opacity:0.5"><svg viewBox="0 0 24 24" width="16" height="16" fill="currentColor"><path d="M12 4l-1.4 1.4L16.2 11H4v2h12.2l-5.6 5.6L12 20l8-8-8-8z"/></svg></div>
    <div class="tbtn"><svg viewBox="0 0 24 24" width="16" height="16" fill="currentColor"><path d="M17.65 6.35A7.958 7.958 0 0 0 12 4c-4.42 0-7.99 3.58-7.99 8s3.57 8 7.99 8c3.73 0 6.84-2.55 7.73-6h-2.08A5.99 5.99 0 0 1 12 18c-3.31 0-6-2.69-6-6s2.69-6 6-6c1.66 0 3.14.69 4.22 1.78L13 11h7V4l-2.35 2.35z"/></svg></div>
    <div class="omni">
      <div class="secure-grp">{lock_svg}<span class="secure-label">{secure_label_text}</span></div>
      <div class="addr">{addr_text}</div>
    </div>
    <div class="tools-right">
      <div class="tbtn" title="Downloads"><svg viewBox="0 0 24 24" width="16" height="16" fill="currentColor"><path d="M19 9h-4V3H9v6H5l7 7 7-7zM5 18v2h14v-2H5z"/></svg></div>
      <div class="tbtn" title="Extensions"><svg viewBox="0 0 24 24" width="16" height="16" fill="currentColor"><path d="M20.5 11H19V7c0-1.1-.9-2-2-2h-4V3.5a2.5 2.5 0 0 0-5 0V5H4c-1.1 0-2 .9-2 2v3.8h1.5c1.1 0 2.7 1.2 2.7 2.7s-1.2 2.7-2.7 2.7H2V20c0 1.1.9 2 2 2h3.8v-1.5c0-1.5 1.2-2.7 2.7-2.7s2.7 1.2 2.7 2.7V22H17c1.1 0 2-.9 2-2v-4h1.5a2.5 2.5 0 0 0 0-5z"/></svg></div>
      <div class="tbtn" title="Menu"><svg viewBox="0 0 24 24" width="16" height="16" fill="currentColor"><circle cx="12" cy="5" r="2"/><circle cx="12" cy="12" r="2"/><circle cx="12" cy="19" r="2"/></svg></div>
    </div>
  </div>
  <div class="page"><img src="data:image/png;base64,{page_png_b64}" width="{page_width}"></div>
  <div class="status">
    <div class="status-url">{url}</div>
    <div class="status-ts">Captured {ts_label}</div>
  </div>
</div></body></html>"""

    async def _get_favicon_b64(page) -> str:
        try:
            return await page.evaluate("""async () => {
                const el = document.querySelector("link[rel*='icon']");
                const url = el ? el.href : (location.origin + '/favicon.ico');
                try {
                    const r = await fetch(url);
                    if (!r.ok) return '';
                    const b = await r.arrayBuffer();
                    let bin = ''; new Uint8Array(b).forEach(v => bin += String.fromCharCode(v));
                    return 'data:image/png;base64,' + btoa(bin);
                } catch { return ''; }
            }""") or ""
        except: return ""

    async def _render_all_frames(context, steps_list):
        """Render all raw screenshots in the steps list into decorative frames."""
        if not steps_list: return
        logging.info(f"Rendering {len(steps_list)} screenshots into decorative frames...")
        temp_page = await context.new_page()
        try:
            for step in steps_list:
                if not step.get("raw_image"): continue
                try:
                    page_b64 = base64.b64encode(step["raw_image"]).decode()
                    html = _build_chrome_frame_html(
                        page_png_b64=page_b64,
                        url=step["url"],
                        title=step["title"],
                        favicon_data_uri=step["favicon_uri"],
                        page_width=SCRIPT_CONFIG["frame_width"],
                        page_height=SCRIPT_CONFIG["frame_height"]
                    )
                    await temp_page.set_content(html)
                    await temp_page.wait_for_load_state("networkidle")
                    # Determine height dynamically
                    h = await temp_page.evaluate("document.querySelector('.window').getBoundingClientRect().height")
                    await temp_page.set_viewport_size({"width": SCRIPT_CONFIG["frame_width"] + 20, "height": int(h) + 20})
                    step["image"] = await temp_page.locator(".window").screenshot(timeout=10000)
                except Exception as e:
                    logging.error(f"Failed to render frame for step '{step['name']}': {e}")
                    step["image"] = step["raw_image"] # Fallback to raw image
        finally:
            await temp_page.close()

    # --- Step Actions ---
    steps = []
    async def capture_step(page, action_name, result_text):
        try:
            # Optimize: only capture raw data, don't render frames yet
            raw_image = await page.screenshot(full_page=False, timeout=5000)
            url = page.url
            title = await page.title()
            favicon_uri = await _get_favicon_b64(page)
            
            steps.append({
                "name": action_name, 
                "result": result_text, 
                "raw_image": raw_image,
                "url": url,
                "title": title,
                "favicon_uri": favicon_uri,
                "image": None # Placeholder for framed image
            })
            logging.info(f"Captured step (raw): {action_name}")
        except Exception as e:
            logging.error(f"Failed to capture step '{action_name}': {e}")
            steps.append({"name": action_name, "result": result_text + f" (Error: {e})", "image": None})

    ## ─────────────────────────────────────────────────────────────────────────

    ## MANDATORY CODE BLOCK — RESILIENT HELPERS (copy verbatim inside run(), before automation steps)
    # The code below provides robust navigation and bulk upload. You MUST use these helpers for ALL interactions.
    async def find_resilient(page, selector, text=None, timeout=10000):
        deadline = asyncio.get_event_loop().time() + (timeout / 1000.0)
        if text:
            if isinstance(text, (re.Pattern, type(re.compile("")))):
                search_target = text
            else:
                search_target = re.compile(re.escape(str(text).strip()), re.I)
        else:
            search_target = None
        
        # Support for comma-separated multiple candidate selectors
        selectors = [s.strip() for s in selector.split(",")]
        
        while asyncio.get_event_loop().time() < deadline:
            for target in [page] + page.frames:
                for sel in selectors:
                    try:
                        locs = target.locator(sel)
                        if search_target: locs = locs.filter(has_text=search_target)
                        if await locs.count() > 0:
                            for i in range(await locs.count()):
                                l = locs.nth(i)
                                if await l.is_visible(): return l
                    except: continue
            await asyncio.sleep(0.5)
        raise TimeoutError(f"Element '{selector}' (text='{text}') not found.")

    async def resilient_click(page, selector, text=None):
        el = await find_resilient(page, selector, text)
        await el.click()

    async def resilient_fill(page, selector, value):
        el = await find_resilient(page, selector)
        await el.fill(value)

    async def resilient_upload(page, target_name, local_path):
        """Hardened 'Native FileChooser' bulk upload strategy."""
        paths = []
        if os.path.isdir(local_path):
            paths = [os.path.abspath(os.path.join(local_path, f)) for f in os.listdir(local_path) if os.path.isfile(os.path.join(local_path, f)) and not f.startswith('.')]
        elif os.path.isfile(local_path):
            paths = [os.path.abspath(local_path)]
        
        if not paths: 
            print(f"[!] ERROR: No files found at {local_path}")
            return 0
        
        print(f"[+] Found {len(paths)} files to upload:")
        for p in paths: print(f"  - {os.path.basename(p)}")

        # 1. Select the Document option and wait for the file chooser to appear
        print(f"[+] Expecting FileChooser from '{target_name}' click...")
        try:
            # Deterministic selector is key
            target_btn = await page.wait_for_selector(SCRIPT_CONFIG["document_option_selector"], state="attached", timeout=10000)
            
            # Using async with context manager is the recommended Playwright way
            async with page.expect_file_chooser() as fc_info:
                # Trigger via dispatch_event or click (dispatch tends to bypass menu occlusion)
                try:
                    await target_btn.dispatch_event("click")
                except:
                    await target_btn.click(force=True)
            
            file_chooser = await fc_info.value
            print("[+] Intercepted native file chooser. Setting files...")
            await file_chooser.set_files(paths)
            
            print(f"[+] Files set. Monitoring upload progress...")
            # Wait for upload indicator to appear and then disappear
            try:
                # OTCS Smart View often shows a progress bar or a specific loader overlay
                await page.wait_for_selector(".csui-progress-bar, .csui-uploading, .binf-progress, .csui-upload-status", state="visible", timeout=7000)
                print("[+] Upload progress detected...")
                await page.wait_for_selector(".csui-progress-bar, .csui-uploading, .binf-progress, .csui-upload-status", state="hidden", timeout=45000)
                print("[+] Upload complete.")
            except:
                print("[+] No explicit progress indicator found, waiting for UI settle...")
                await page.wait_for_timeout(5000)
            
            await page.wait_for_load_state("networkidle")
            return len(paths)
        except Exception as e:
            print(f"[!] FileChooser upload failed: {e}")
            raise e

    try:
        logging.getLogger("asyncio").setLevel(logging.ERROR)
        creds = vault_creds("OTCS")
        username = creds["username"]
        password = creds["password"]

        p = await async_playwright().start()
        browser = await p.chromium.launch(headless=False, args=["--start-maximized"])
        context = await browser.new_context(no_viewport=True)
        page = await context.new_page()

        new_folder_name = f"Upload_Test_{datetime.now().strftime('%Y%m%d_%H%M%S')}"

        # 1. Navigate to URL
        await page.goto(SCRIPT_CONFIG["base_url"])
        await page.wait_for_load_state("networkidle")
        await capture_step(page, "Navigate to OTCS Base URL", f"Navigated to {page.url}")

        # Login
        await page.wait_for_selector(SCRIPT_CONFIG["password_selector"], timeout=15000)
        await resilient_fill(page, SCRIPT_CONFIG["username_selector"], username)
        await resilient_fill(page, SCRIPT_CONFIG["password_selector"], password)
        await resilient_click(page, SCRIPT_CONFIG["login_button_selector"], text=re.compile(r"Sign in|Login", re.I))
        await page.wait_for_load_state("networkidle")
        await page.wait_for_selector(SCRIPT_CONFIG["password_selector"], state="hidden", timeout=15000)
        await capture_step(page, "Login to OTCS", "Successfully logged in.")

        # 2. Click the '+' (Add Item) button
        await resilient_click(page, SCRIPT_CONFIG["add_item_button_selector"])
        await page.wait_for_timeout(2000) # Buffer for menu animation
        await capture_step(page, "Click Add Item Button", "Add item menu opened.")

        # 3. Select 'Folder'
        # Wait for any dropdown menu container to appear
        try: await page.wait_for_selector(".binf-dropdown-menu, .csui-menu-container, [role='menu']", state="visible", timeout=5000)
        except: pass

        # Use deterministic selector instead of text search
        folder_btn = await page.wait_for_selector(SCRIPT_CONFIG["folder_option_selector"], state="attached", timeout=10000)

        if not folder_btn:
            raise TimeoutError(f"Could not find deterministic folder item: '{SCRIPT_CONFIG['folder_option_text']}'")

        print(f"[+] Found unique menu item: '{SCRIPT_CONFIG['folder_option_text']}'. Dispatching click event...")
        try:
            await folder_btn.dispatch_event("click")
        except:
            await folder_btn.click(force=True)
            
        await capture_step(page, "Select 'Folder'", "Folder creation form displayed.")

        # 4. Type 'Upload_Test_{timestamp}' into the 'Enter name' field and save by pressing 'Enter'
        await resilient_fill(page, SCRIPT_CONFIG["folder_name_input_selector"], new_folder_name)
        await page.keyboard.press("Enter")
        
        # Wait for the folder name to appear in the table. 
        # Use find_resilient and a reload fallback to handle SPA synchronization delays.
        print(f"[+] Waiting for folder '{new_folder_name}' to appear in list...")
        try:
            # Look for name in <a>, <span>, or <div> for maximum resilience
            await find_resilient(page, "a, span, div", text=new_folder_name, timeout=10000)
        except Exception:
            print("[!] Folder didn't appear immediately. Reloading page to force synchronization...")
            await page.reload()
            await page.wait_for_load_state("networkidle")
            await find_resilient(page, "a, span, div", text=new_folder_name, timeout=15000)

        await capture_step(page, f"Create Folder '{new_folder_name}'", "Folder created and verified in list.")

        # --- 5. Created folder check & Sort ---
        await page.wait_for_load_state("networkidle")
        await page.wait_for_timeout(2000) # Wait for SPA indexing to catch up
        
        logging.info(f"[+] Locating 'Modified' column header for sorting...")
        modified_header = await find_resilient(page, SCRIPT_CONFIG["modified_column_header_selector"], text=re.compile(SCRIPT_CONFIG["modified_column_header_text"], re.I))
        
        # Check if already sorted descending, if not, click once or twice
        current_url = page.url
        if "order_by=modify_date_desc" not in current_url:
            print("[+] Clicking Modified header to sort...")
            await modified_header.click()
            await page.wait_for_load_state("networkidle")
            await page.wait_for_timeout(2000)
            # Check again, if still not descending, click once more
            if "order_by=modify_date_desc" not in page.url:
                print("[+] Clicking Modified header again for descending sort...")
                await modified_header.click()
                await page.wait_for_load_state("networkidle")
                await page.wait_for_timeout(2000)
        
        await capture_step(page, "Sort by 'Modified' Column", f"Table sorted. Current URL: {page.url}")

        # 6. Enter the new folder (top row)
        # Use find_resilient and a reload fallback if navigation fails
        print(f"[+] Searching for folder: {new_folder_name}")
        try:
            folder_link = await find_resilient(page, "a", text=new_folder_name, timeout=10000)
        except Exception:
            print("[!] Folder not found after sorting. Attempting page reload...")
            await page.reload()
            await page.wait_for_load_state("networkidle")
            folder_link = await find_resilient(page, "a", text=new_folder_name, timeout=15000)

        href = await folder_link.get_attribute("href")
        if href:
            target_url = urljoin(page.url, href)
            print(f"[+] Navigating to folder URL: {target_url}")
            await page.goto(target_url)
        else:
            await folder_link.click()
        
        await page.wait_for_load_state("networkidle")
        await capture_step(page, f"Enter Folder '{new_folder_name}'", f"Navigated into folder: {page.url}")

        # 7. Inside the new folder, click '+' (Add Item) button and select 'Document'
        await resilient_click(page, SCRIPT_CONFIG["add_item_button_selector"])
        await page.wait_for_timeout(2000) # Buffer for menu animation
        await capture_step(page, "Click Add Item Button (inside new folder)", "Add item menu opened.")

        # Resilient match for 'Document' (allowing icons)
        await resilient_click(page, SCRIPT_CONFIG["document_option_selector"], text=re.compile(re.escape(SCRIPT_CONFIG['document_option_text']), re.I))
        await capture_step(page, "Select 'Document'", "Document upload dialog/file chooser initiated.")

        # 8. Bulk upload the entire folder 'C:\Users\pnandank\Downloads\OTCS_Upload_Test'
        # 8. Bulk upload the entire folder 'C:\Users\pnandank\Downloads\OTCS_Upload_Test'
        if not os.path.exists(SCRIPT_CONFIG["upload_directory_path"]):
            print(f"[!] ERROR: Source directory '{SCRIPT_CONFIG['upload_directory_path']}' does not exist!")
            await capture_step(page, "Validate Source Directory", f"FAILURE: Directory '{SCRIPT_CONFIG['upload_directory_path']}' missing.")
            uploaded_count = 0
        else:
            uploaded_count = await resilient_upload(page, SCRIPT_CONFIG['document_option_text'], SCRIPT_CONFIG["upload_directory_path"])
            await page.wait_for_timeout(2000)
            await capture_step(page, "Upload Files Triggered", f"Files set. Count: {uploaded_count}")
            await capture_step(page, "Bulk Upload Files", f"Successfully uploaded {uploaded_count} files.")

        # 9. Finalized 
        pass

    except Exception as e:
        logging.error(f"Execution failed: {e}")
        import traceback; traceback.print_exc()
    finally:
        # CLOSING BROWSER AND GENERATING REPORT MUST BE IN FINALLY
        try:
            if context and steps:
                await _render_all_frames(context, steps)
            if browser: await browser.close()
            if p: await p.stop()
        except Exception as e:
            logging.error(f"Error closing browser/Playwright: {e}")
        await asyncio.sleep(0.5) 

        # --- GENERATE REPORT HERE ---
        doc.add_heading("OTCS Folder Creation and Bulk Upload Report", level=1)
        doc.add_paragraph(f"Report generated on: {datetime.now().strftime('%Y-%m-%d %H:%M:%S')}")
        doc.add_paragraph(f"Target URL: {SCRIPT_CONFIG['base_url']}")
        doc.add_paragraph(f"Created Folder Name: {new_folder_name}")
        doc.add_paragraph(f"Uploaded from Directory: {SCRIPT_CONFIG['upload_directory_path']}")
        
        doc.add_heading("Execution Summary", level=2)
        table = doc.add_table(rows=1, cols=3)
        table.style = 'Table Grid'
        hdr_cells = table.rows[0].cells
        hdr_cells[0].text = 'Step'
        hdr_cells[1].text = 'Action'
        hdr_cells[2].text = 'Result'

        for i, step in enumerate(steps):
            row_cells = table.add_row().cells
            row_cells[0].text = str(i + 1)
            row_cells[1].text = step["name"]
            row_cells[2].text = step["result"]
        
        doc.add_page_break()
        doc.add_heading("Step Screenshots", level=2)
        import tempfile
        for step in steps:
            if step.get("image"):
                doc.add_heading(f"Step {steps.index(step) + 1}: {step['name']}", level=3)
                doc.add_paragraph(f"Result: {step['result']}")
                
                # Use a temp file for the screenshot to insert into Word
                tfile = tempfile.NamedTemporaryFile(delete=False, suffix=".png")
                try:
                    tfile.write(step["image"])
                    tfile.close()
                    doc.add_picture(tfile.name, width=Inches(6.0))
                finally:
                    # Clean up the temp file
                    if os.path.exists(tfile.name):
                        os.remove(tfile.name)
                doc.add_paragraph()

        try:
            doc.save(report_path)
            logging.info(f"Report saved to {report_path}")
        except Exception as e:
            logging.error(f"Failed to save report: {e}")

if __name__ == "__main__":
    asyncio.run(run())