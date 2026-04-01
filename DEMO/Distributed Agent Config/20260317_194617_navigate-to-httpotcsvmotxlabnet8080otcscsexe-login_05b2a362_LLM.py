# ── Vault credential reader ───────────────────────────────────────────────────
import sys as _sys
from pathlib import Path as _Path

def vault_creds(prefix: str) -> dict:
    """Return {'username': ..., 'password': ...} from the encrypted vault."""
    # Walk up from this script to find the web-ui root
    base = _Path(__file__).resolve()
    for _ in range(8):
        base = base.parent
        candidate = base / "web-ui"
        if candidate.exists():
            _sys.path.insert(0, str(candidate))
            break
    try:
        from src.utils.vault import vault
        creds = vault.get_credentials(prefix) or vault.get_credentials(prefix.lower())
        if not creds:
            raise RuntimeError(f"No vault entry found for prefix '{prefix}'")
        lower = {k.lower(): v for k, v in creds.items()}
        def _pick(cs):
            for c in cs:
                if c in lower: return lower[c]
            return list(lower.values())[0] if lower else ""
        return {
            "username": _pick(("username","user","login","email")),
            "password": _pick(("password","passwd","pwd","pass")),
        }
    except Exception as e:
        raise RuntimeError(f"Vault error for '{prefix}': {e}")
# ─────────────────────────────────────────────────────────────────────────────

import os, sys, logging, asyncio, io, re
from datetime import datetime
from playwright.async_api import async_playwright
from docx import Document
from docx.shared import Inches

# Set Windows event loop policy for stability
if sys.platform == 'win32':
    asyncio.set_event_loop_policy(asyncio.WindowsProactorEventLoopPolicy())

async def run():
    logging.getLogger("asyncio").setLevel(logging.ERROR)

    SCRIPT_CONFIG = {
        "base_url": "http://otcsvm.otxlab.net:8080/OTCS/cs.exe",
        "agent_status_url": "http://otcsvm.otxlab.net:8080/OTCS/cs.exe?func=distributedagent.AgentStatus",
        "admin_servers_url": "http://otcsvm.otxlab.net:8080/OTCS/cs.exe?func=ll&objtype=148&objaction=browse",
        "username_selector": "#otds_username",
        "password_selector": "#otds_password",
        "login_button_selector": "#loginbutton",
        "distributed_agent_id_col": "Distributed Agent ID",
        "status_col": "Status",
        "expected_status": "Running",
        "configure_agent_system_link_text": "Configure the Distributed Agent System",
        "save_button_name": "Save",
        "admin_server_name": "AdminServer-01",
        "admin_server_action_menu_selectors": ["#x1"],
        "report_filename_prefix": "OTCS_Agent_Report",
        "frame_width": 1280,
        "frame_height": 900
    }

    import base64
    from urllib.parse import urlparse

    def _build_chrome_frame_html(page_png_b64, url, title, favicon_data_uri="", page_width=1280, page_height=900):
        is_secure = url.startswith("https://")
        parsed = urlparse(url)
        # Show full address in the bar: netloc + path + query
        addr_text = parsed.netloc + parsed.path
        if parsed.query:
            addr_text += "?" + parsed.query

        lock_col  = "#188038" if is_secure else "#c5221f"
        tab_title = title[:30] + "…" if len(title) > 32 else title
        
        import time as _t
        _tzn = _t.tzname[_t.daylight] if _t.daylight else _t.tzname[0]
        _tz  = "".join(w[0] for w in _tzn.split()) if len(_tzn) > 5 else _tzn
        ts_label = datetime.now().strftime(f"%Y-%m-%d %H:%M:%S {_tz}")

        if is_secure:
            lock_svg = f'<svg viewBox="0 0 16 16" width="14" height="14" fill="{lock_col}"><path d="M8 1a3.5 3.5 0 0 0-3.5 3.5V6H4a1 1 0 0 0-1 1v6a1 1 0 0 0 1 1h8a1 1 0 0 0 1-1V7a1 1 0 0 0-1-1h-.5V4.5A3.5 3.5 0 0 0 8 1zm2 5H6V4.5a2 2 0 1 1 4 0V6z"/></svg>'
            secure_label_text = 'Secure'
        else:
            lock_svg = f'<svg viewBox="0 0 20 20" width="14" height="14" fill="{lock_col}"><path d="M3.7 3.7a.75.75 0 0 0-1.1 1.1l2.2 2.2H4a1 1 0 0 0-1 1v6a1 1 0 0 0 1 1h8a1 1 0 0 1 .6.2l1.7 1.7a.75.75 0 1 0 1.1-1.1L3.7 3.7zM5.1 9l-.1-.1V9h.1zm1 0H13v5H7.1L6 12.9V9zM7.5 5a2 2 0 0 1 3.9-.5l1.5 1.5H8V4.5A2 2 0 0 0 7.5 5zm2.5-2A3.5 3.5 0 0 0 6.5 6H4.9L3.5 4.6A3.5 3.5 0 0 1 10 3z"/></svg>'
            secure_label_text = 'Not secure'

        if favicon_data_uri:
            favicon_html = f'<img src="{favicon_data_uri}" width="16" height="16" style="flex-shrink:0;border-radius:2px" onerror="this.style.display=\'none\'">'
        else:
            favicon_html = '<svg viewBox="0 0 16 16" width="16" height="16" fill="#5f6368" style="flex-shrink:0"><circle cx="8" cy="8" r="7" fill="none" stroke="#5f6368" stroke-width="1.2"/><ellipse cx="8" cy="8" rx="3" ry="7" fill="none" stroke="#5f6368" stroke-width="1.2"/><line x1="1" y1="8" x2="15" y2="8" stroke="#5f6368" stroke-width="1.2"/></svg>'

        return f"""<!DOCTYPE html><html><head><meta charset="utf-8">
<style>
  html, body {{ overflow: hidden; }}
  * {{ box-sizing: border-box; margin: 0; padding: 0; font-family: 'Segoe UI', -apple-system, system-ui, sans-serif; font-size: 13px; }}
  body {{ background: #e8eaed; display: inline-flex; flex-direction: column; }}
  .window {{ display: flex; flex-direction: column; width: {page_width + 2}px; box-shadow: 0 2px 16px rgba(0,0,0,.25); border-radius: 8px 8px 0 0; overflow: hidden; border: 1px solid #b0b0b0; }}
  .tabbar {{ background: #dee1e6; display: flex; align-items: flex-end; height: 40px; padding-left: 10px; flex-shrink: 0; position: relative; }}
  .win-ctrls {{ position: absolute; right: 0; top: 0; display: flex; }}
  .win-btn {{ width: 44px; height: 32px; display: flex; align-items: center; justify-content: center; color: #5f6368; }}
  .win-btn:hover {{ background: #dadce0; }}
  .win-btn.close:hover {{ background: #e81123; color: #fff; }}
  .tab {{ background: #fff; height: 34px; padding: 0 12px; min-width: 180px; max-width: 240px; display: flex; align-items: center; gap: 8px; border-radius: 8px 8px 0 0; color: #202124; position: relative; }}
  .tab-title {{ flex: 1; overflow: hidden; white-space: nowrap; text-overflow: ellipsis; }}
  .tab-x {{ color: #5f6368; font-size: 16px; width: 18px; text-align: center; }}
  .toolbar {{ background: #fff; height: 40px; display: flex; align-items: center; padding: 0 8px; gap: 4px; border-bottom: 1px solid #dadce0; }}
  .tbtn {{ width: 28px; height: 28px; display: flex; align-items: center; justify-content: center; color: #5f6368; border-radius: 50%; flex-shrink:0; }}
  .tbtn:hover {{ background: #f1f3f4; }}
  .omni {{ flex: 1; height: 32px; background: #f1f3f4; border-radius: 24px; display: flex; align-items: center; gap: 8px; padding: 0 12px; margin: 0 4px; overflow: hidden; }}
  .secure-grp {{ display: flex; align-items: center; gap: 6px; flex-shrink: 0; }}
  .secure-label {{ font-size: 11px; color: {lock_col}; font-weight: 600; white-space: nowrap; line-height: 1; }}
  .addr {{ font-size: 14px; color: #202124; flex: 1; white-space: nowrap; overflow: hidden; text-overflow: ellipsis; }}
  .tools-right {{ display: flex; align-items: center; gap: 2px; padding-right: 4px; }}
  .page img {{ display: block; }}
  .status {{ background: #f8f9fa; height: 22px; display: flex; align-items: center; padding: 0 10px; font-size: 11px; color: #80868b; border-top: 1px solid #e8eaed; }}
  .status-url {{ flex: 1; white-space: nowrap; overflow: hidden; text-overflow: ellipsis; }}
  .status-ts {{ flex-shrink: 0; padding-left: 10px; }}
</style></head><body>
<div class="window">
  <div class="tabbar">
    <div class="tab">{favicon_html}<span class="tab-title">{tab_title}</span><span class="tab-x">×</span></div>
    <div class="win-ctrls">
      <div class="win-btn"><svg viewBox="0 0 10 1" width="10"><path d="M0 0h10v1H0z" fill="currentColor"/></svg></div>
      <div class="win-btn"><svg viewBox="0 0 10 10" width="10"><path d="M0 0h10v10H0V0zm1 1v8h8V1H1z" fill="currentColor"/></svg></div>
      <div class="win-btn close"><svg viewBox="0 0 10 10" width="10"><path d="M0 0l10 10m0-10L0 10" stroke="currentColor" stroke-width="1.2"/></svg></div>
    </div>
  </div>
  <div class="toolbar">
    <div class="tbtn"><svg viewBox="0 0 24 24" width="16" height="16" fill="currentColor"><path d="M20 11H7.8l5.6-5.6L12 4l-8 8 8 8 1.4-1.4L7.8 13H20v-2z"/></svg></div>
    <div class="tbtn" style="opacity:0.5"><svg viewBox="0 0 24 24" width="16" height="16" fill="currentColor"><path d="M12 4l-1.4 1.4L16.2 11H4v2h12.2l-5.6 5.6L12 20l8-8-8-8z"/></svg></div>
    <div class="tbtn"><svg viewBox="0 0 24 24" width="16" height="16" fill="currentColor"><path d="M17.65 6.35A7.958 7.958 0 0 0 12 4c-4.42 0-7.99 3.58-7.99 8s3.57 8 7.99 8c3.73 0 6.84-2.55 7.73-6h-2.08A5.99 5.99 0 0 1 12 18c-3.31 0-6-2.69-6-6s2.69-6 6-6c1.66 0 3.14.69 4.22 1.78L13 11h7V4l-2.35 2.35z"/></svg></div>
    <div class="omni">
      <div class="secure-grp">{lock_svg}<span class="secure-label">{secure_label_text}</span></div>
      <div class="addr">{addr_text}</div>
    </div>
    <div class="tools-right">
      <div class="tbtn" title="Downloads"><svg viewBox="0 0 24 24" width="16" height="16" fill="currentColor"><path d="M19 9h-4V3H9v6H5l7 7 7-7zM5 18v2h14v-2H5z"/></svg></div>
      <div class="tbtn" title="Extensions"><svg viewBox="0 0 24 24" width="16" height="16" fill="currentColor"><path d="M20.5 11H19V7c0-1.1-.9-2-2-2h-4V3.5a2.5 2.5 0 0 0-5 0V5H4c-1.1 0-2 .9-2 2v3.8h1.5c1.1 0 2.7 1.2 2.7 2.7s-1.2 2.7-2.7 2.7H2V20c0 1.1.9 2 2 2h3.8v-1.5c0-1.5 1.2-2.7 2.7-2.7s2.7 1.2 2.7 2.7V22H17c1.1 0 2-.9 2-2v-4h1.5a2.5 2.5 0 0 0 0-5z"/></svg></div>
      <div class="tbtn" title="Menu"><svg viewBox="0 0 24 24" width="16" height="16" fill="currentColor"><circle cx="12" cy="5" r="2"/><circle cx="12" cy="12" r="2"/><circle cx="12" cy="19" r="2"/></svg></div>
    </div>
  </div>
  <div class="page"><img src="data:image/png;base64,{page_png_b64}" width="{page_width}"></div>
  <div class="status">
    <div class="status-url">{url}</div>
    <div class="status-ts">Captured {ts_label}</div>
  </div>
</div></body></html>"""

    async def _get_favicon_b64(page) -> str:
        try:
            return await page.evaluate("""async () => {
                const el = document.querySelector("link[rel*='icon']");
                const url = el ? el.href : (location.origin + '/favicon.ico');
                try {
                    const r = await fetch(url);
                    if (!r.ok) return '';
                    const b = await r.arrayBuffer();
                    let bin = ''; new Uint8Array(b).forEach(v => bin += String.fromCharCode(v));
                    return 'data:image/png;base64,' + btoa(bin);
                } catch { return ''; }
            }""") or ""
        except: return ""

    async def _render_all_frames(context, steps_list):
        """Render all raw screenshots in the steps list into decorative frames."""
        if not steps_list: return
        logging.info(f"Rendering {len(steps_list)} screenshots into decorative frames...")
        temp_page = await context.new_page()
        try:
            for step in steps_list:
                if not step.get("raw_image"): continue
                try:
                    page_b64 = base64.b64encode(step["raw_image"]).decode()
                    html = _build_chrome_frame_html(
                        page_png_b64=page_b64,
                        url=step["url"],
                        title=step["title"],
                        favicon_data_uri=step["favicon_uri"],
                        page_width=SCRIPT_CONFIG["frame_width"],
                        page_height=SCRIPT_CONFIG["frame_height"]
                    )
                    await temp_page.set_content(html)
                    await temp_page.wait_for_load_state("networkidle")
                    # Determine height dynamically
                    h = await temp_page.evaluate("document.querySelector('.window').getBoundingClientRect().height")
                    await temp_page.set_viewport_size({"width": SCRIPT_CONFIG["frame_width"] + 20, "height": int(h) + 20})
                    step["image"] = await temp_page.locator(".window").screenshot(timeout=10000)
                except Exception as e:
                    logging.error(f"Failed to render frame for step '{step['name']}': {e}")
                    step["image"] = step["raw_image"] # Fallback to raw image
        finally:
            await temp_page.close()

    # --- Step Actions ---
    steps = []
    async def capture_step(page, action_name, result_text):
        try:
            # Optimize: only capture raw data, don't render frames yet
            raw_image = await page.screenshot(full_page=False, timeout=5000)
            url = page.url
            title = await page.title()
            favicon_uri = await _get_favicon_b64(page)
            
            steps.append({
                "name": action_name, 
                "result": result_text, 
                "raw_image": raw_image,
                "url": url,
                "title": title,
                "favicon_uri": favicon_uri,
                "image": None # Placeholder for framed image
            })
            logging.info(f"Captured step (raw): {action_name}")
        except Exception as e:
            logging.error(f"Failed to capture step '{action_name}': {e}")
            steps.append({"name": action_name, "result": result_text + f" (Error: {e})", "image": None})

    def clean_id(raw):
        # Strip all leading/trailing whitespace, take the first non-empty line, then normalize internal spaces
        if not raw: return ""
        lines = [line.strip() for line in raw.split('\n') if line.strip()]
        if not lines: return ""
        return re.sub(r'\s+', ' ', lines[0])[:100]

    creds = vault_creds("OTCS")
    username = creds["username"]
    password = creds["password"]

    p = None
    browser = None
    try:
        p = await async_playwright().start()
        browser = await p.chromium.launch(headless=False, args=["--start-maximized"])
        context = await browser.new_context(no_viewport=True)
        page = await context.new_page()

        # --- Login ---
        await page.goto(SCRIPT_CONFIG["base_url"])
        await page.wait_for_selector('input[type="password"]', timeout=15000)
        await capture_step(page, "Navigate to Login Page", f"Navigated to {SCRIPT_CONFIG['base_url']}")

        await page.locator(SCRIPT_CONFIG["username_selector"]).fill(username)
        await page.locator(SCRIPT_CONFIG["password_selector"]).fill(password)
        await capture_step(page, "Fill Login Credentials", "Username and password filled.")

        await page.locator(SCRIPT_CONFIG["login_button_selector"]).click()
        await page.wait_for_load_state("networkidle")
        await page.wait_for_selector('input[type="password"]', state="hidden", timeout=15000)
        await capture_step(page, "Click Login Button", "Login successful, redirected to home page.")

        # --- Step 1: Navigate to Distributed Agent Dashboard ---
        await page.goto(SCRIPT_CONFIG["agent_status_url"])
        await page.wait_for_load_state("networkidle")
        await capture_step(page, "Navigate to Distributed Agent Dashboard", f"Navigated to {SCRIPT_CONFIG['agent_status_url']}")

        # --- Step 2: Read the Distributed Agent Status table ---
        distributed_agents_data = []
        all_agents_running = True
        agents_not_running = []

        logging.info("Searching for Distributed Agent Status table...")
        
        # More robust: find the header text first, then its containing table
        header_th_locator = page.get_by_role("columnheader", name=re.compile(r"Distributed Agent ID", re.I)).first
        await header_th_locator.wait_for(state="visible", timeout=20000)
        
        header_row_locator = header_th_locator.locator("xpath=./ancestor::tr[1]")
        table_locator = header_th_locator.locator("xpath=./ancestor::table[1]")
        
        # Get all cells in the header row
        header_cells = await header_row_locator.locator("th, td").all()
        header_texts = [clean_id(await cell.inner_text()) for cell in header_cells]
        logging.info(f"Found table headers: {header_texts}")

        agent_id_col_idx = -1
        status_col_idx = -1

        for i, text in enumerate(header_texts):
            if SCRIPT_CONFIG["distributed_agent_id_col"].lower() in text.lower():
                agent_id_col_idx = i
            if SCRIPT_CONFIG["status_col"].lower() in text.lower() and "description" not in text.lower():
                status_col_idx = i
        
        logging.info(f"Column Mapping -> Agent ID: {agent_id_col_idx}, Status: {status_col_idx}")

        if agent_id_col_idx == -1 or status_col_idx == -1:
            raise RuntimeError(f"Could not identify columns in {header_texts}")

        # Get all data rows
        all_rows = await table_locator.locator("tr").all()
        logging.info(f"Total rows in table: {len(all_rows)}")
        
        for row_locator in all_rows:
            cells = await row_locator.locator("td").all()
            # Data rows in this table typically have 4+ cells and start with "DAAgent" or similar
            if len(cells) > max(agent_id_col_idx, status_col_idx):
                agent_id = clean_id(await cells[agent_id_col_idx].inner_text())
                status = clean_id(await cells[status_col_idx].inner_text())
                
                # Check for "DAAgent" to ensure it's a data row and not a spacer/header
                if agent_id and "DAAgent" in agent_id:
                    logging.info(f"Processing Agent: {agent_id}, Status: {status}")
                    distributed_agents_data.append({
                        SCRIPT_CONFIG["distributed_agent_id_col"]: agent_id,
                        SCRIPT_CONFIG["status_col"]: status
                    })
                    if status.lower() != SCRIPT_CONFIG["expected_status"].lower():
                        all_agents_running = False
                        agents_not_running.append(agent_id)

        await capture_step(page, "Read Distributed Agent Status Table", f"Extracted {len(distributed_agents_data)} agents. All running: {all_agents_running}")

        # --- Step 3 & 4: Verify Status and Conditional Configuration ---
        if not all_agents_running:
            logging.info(f"Agents not running: {agents_not_running}. Proceeding to configure.")
            # Step 4a.a: Click "Configure the Distributed Agent System"
            configure_link = page.get_by_text(SCRIPT_CONFIG["configure_agent_system_link_text"], exact=True).filter(visible=True).first
            await configure_link.wait_for(state="visible", timeout=15000)
            await configure_link.click()
            await page.wait_for_load_state("networkidle")
            await capture_step(page, "Click Configure Agent System", "Navigated to configuration page.")

            # Step 4a.b: Select radio buttons for agents not running
            for agent_id in agents_not_running:
                # Find the row for the specific agent
                agent_row_locator = page.locator("tr").filter(has=page.locator("td, th").filter(has_text=re.compile(r"\b" + re.escape(agent_id) + r"\b", re.I))).first
                await agent_row_locator.wait_for(state="visible", timeout=15000)
                
                # Locate and click the radio button within that row
                radio_button = agent_row_locator.locator("input[type='radio']").first
                await radio_button.wait_for(state="visible", timeout=15000)
                await radio_button.click()
                await capture_step(page, f"Select Agent '{agent_id}' for Configuration", f"Radio button selected for {agent_id}.")
            
            # Step 4a.c: Click Save
            save_button = page.get_by_role("button", name=re.compile(r"^" + re.escape(SCRIPT_CONFIG["save_button_name"]) + r"$", re.I)).filter(visible=True).first
            await save_button.wait_for(state="visible", timeout=15000)
            await save_button.click()
            await page.wait_for_load_state("networkidle")
            await capture_step(page, "Click Save Configuration", "Configuration saved.")

            # Step 4a.d: Navigate back to Distributed Agent Dashboard
            await page.goto(SCRIPT_CONFIG["agent_status_url"])
            await page.wait_for_load_state("networkidle")
            await capture_step(page, "Navigate back to Agent Dashboard", "Returned to Distributed Agent Status page after configuration.")
            
            action_taken_summary = f"Configured agents: {', '.join(agents_not_running)}"

        else:
            logging.info("All agents are running. Skipping configuration.")
            await capture_step(page, "Verify Agent Status", "All agents are already running. Configuration skipped.")
            action_taken_summary = "All agents were running. No configuration needed."

        # --- Step 5: Navigate to Admin Servers ---
        await page.goto(SCRIPT_CONFIG["admin_servers_url"])
        await page.wait_for_load_state("networkidle")
        await capture_step(page, "Navigate to Admin Servers", f"Navigated to {SCRIPT_CONFIG['admin_servers_url']}")

        # --- Step 6: Open the Admin Server function menu ---
        admin_server_name = SCRIPT_CONFIG["admin_server_name"]
        
        # Find the row for "AdminServer-01"
        admin_server_row = page.locator("tr").filter(has=page.locator("td, th").filter(has_text=re.compile(r"\b" + re.escape(admin_server_name) + r"\b", re.I))).first
        await admin_server_row.wait_for(state="visible", timeout=15000)
        
        # Locate the action menu within that row
        action_menu_locator = None
        for selector in SCRIPT_CONFIG["admin_server_action_menu_selectors"]:
            candidate = admin_server_row.locator(selector).first
            if await candidate.is_visible():
                action_menu_locator = candidate
                break
        
        if not action_menu_locator:
            raise RuntimeError(f"Could not find action menu for '{admin_server_name}' using provided selectors.")

        # Vigorous Menu Interaction
        await action_menu_locator.hover()
        await action_menu_locator.dispatch_event('mousedown')
        await action_menu_locator.dispatch_event('mouseup')
        await action_menu_locator.click()
        await page.wait_for_timeout(1000) # Mandatory wait for menu to animate
        await capture_step(page, f"Open Function Menu for '{admin_server_name}'", f"Opened action menu for {admin_server_name}.")

        # --- Post-process: Render all frames ---
        await _render_all_frames(context, steps)

        # --- Reporting ---
        report_filename = f"{SCRIPT_CONFIG['report_filename_prefix']}_{datetime.now().strftime('%Y%m%d_%H%M%S')}.docx"
        report_path = os.path.join(os.path.dirname(os.path.abspath(__file__)), report_filename)
        doc = Document()
        doc.add_heading("OTCS Distributed Agent and Admin Server Report", level=1)

        doc.add_heading("Execution Summary", level=2)
        summary_table = doc.add_table(rows=1, cols=3)
        summary_table.autofit = True
        hdr_cells = summary_table.rows[0].cells
        hdr_cells[0].text = "Step"
        hdr_cells[1].text = "Action"
        hdr_cells[2].text = "Result"

        for i, step in enumerate(steps):
            row_cells = summary_table.add_row().cells
            row_cells[0].text = str(i + 1)
            row_cells[1].text = step["name"]
            row_cells[2].text = step["result"]

        doc.add_heading("Distributed Agent Status Table", level=2)
        if distributed_agents_data:
            table = doc.add_table(rows=1, cols=len(distributed_agents_data[0]))
            table.autofit = True
            # Add header row
            hdr_cells = table.rows[0].cells
            for i, col_name in enumerate(distributed_agents_data[0].keys()):
                hdr_cells[i].text = col_name
            # Add data rows
            for agent in distributed_agents_data:
                row_cells = table.add_row().cells
                for i, value in enumerate(agent.values()):
                    row_cells[i].text = value
        else:
            doc.add_paragraph("No distributed agent data was extracted.")

        doc.add_heading("Distributed Agent Status Verification", level=2)
        doc.add_paragraph(f"All agents running: {'Yes' if all_agents_running else 'No'}")
        if not all_agents_running:
            doc.add_paragraph(f"Agents not running: {', '.join(agents_not_running)}")
        doc.add_paragraph(f"Action taken: {action_taken_summary}")

        doc.add_heading("Step Screenshots", level=2)
        for i, step in enumerate(steps):
            doc.add_heading(f"Step {i+1}: {step['name']}", level=3)
            doc.add_paragraph(f"Result: {step['result']}")
            if step["image"]:
                image_stream = io.BytesIO(step["image"])
                doc.add_picture(image_stream, width=Inches(6))
            else:
                doc.add_paragraph("No screenshot available for this step.")
            doc.add_page_break()

        doc.save(report_path)
        logging.info(f"Report saved to: {report_path}")

    except Exception as e:
        logging.error(f"An error occurred: {e}", exc_info=True)
        import traceback; traceback.print_exc()
        raise
    finally:
        if browser:
            await browser.close()
        if p:
            await p.stop()
        if sys.platform == 'win32':
            await asyncio.sleep(0.5) # Give time for browser to close on Windows

if __name__ == "__main__":
    asyncio.run(run())